"""
Blueprint Generator – PyroguardIAn & SolarEyes (ESPAÑOL)
Proyecto:  IT - BTP - PyroguardIAn - SolarEyes
Unidad:    EGP TGX (Enel Green Power – Technology and Grid eXpansion)
Fuente:    Bluedraft/Blueprint_PyroguardIAn_SolarEyes.docx  (versión italiana)
Plantilla: Blueprint_1.4_vuoto.docx
Idioma:    Español

PREGUNTAS ABIERTAS — secciones con datos a integrar antes de publicación:
  1. roles            — Nombres y unidades organizativas de roles del proyecto
                        (Business Owner, IT Owner, Data Owner, Product Owner EGP TGX)
  2. sistemas_asis    — Nombres exactos de sistemas SCADA/DCS, CMMS, ERP en uso
                        en instalaciones solares EGP [A DEFINIR con IT Owner]
  3. kpi_cuantitativos — Métricas baseline AS-IS del histórico operativo EGP
                        (eventos incendio/año, MTTR real, coste/MW inspección)
                        [A CONFIRMAR con datos EGP]
  4. data_mapping     — Detalle formatos API y protocolos de integración disponibles
                        [A DEFINIR con IT Owner OT]
  5. arquitectura     — Elección de stack tecnológico final (cloud provider, MLOps,
                        proveedor de drones) [A VALIDAR con IT Owner]

Ejecución:
    python generate_blueprint_PyroguardIAn_ES.py [RUTA_PLANTILLA] [RUTA_SALIDA]
"""

import copy
import os
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── resolución de rutas ───────────────────────────────────────────────────────

def _resolve_path(cli_index, env_name, default_value):
    if len(sys.argv) > cli_index and sys.argv[cli_index]:
        return sys.argv[cli_index]
    return os.environ.get(env_name, default_value)


TEMPLATE = _resolve_path(1, "BLUEPRINT_TEMPLATE", "Blueprint_1.4_vuoto.docx")
OUTPUT_DIR = "Bluedraft"
OUTPUT = _resolve_path(2, "BLUEPRINT_OUTPUT",
                       os.path.join(OUTPUT_DIR, "Blueprint_PyroguardIAn_SolarEyes_ES.docx"))


# ── helpers ──────────────────────────────────────────────────────────────────

def find_para(doc, fragment, style=None, start=0):
    if start is None:
        start = 0
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if fragment in p.text:
            if style is None or p.style.name == style:
                return i, p
    return None, None


def _get_rpr(para):
    for run in para.runs:
        rpr = run._r.find(qn("w:rPr"))
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None


def set_para_text(para, new_text):
    """Replace all runs in para with a single run containing new_text, preserving formatting."""
    rpr = _get_rpr(para)
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    r_elem = OxmlElement("w:r")
    if rpr is not None:
        r_elem.append(rpr)
    t_elem = OxmlElement("w:t")
    t_elem.text = new_text
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_elem.append(t_elem)
    para._p.append(r_elem)


def set_para_lines(para, lines):
    """Replace paragraph content with multiple lines joined by w:br (soft return)."""
    rpr = _get_rpr(para)
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    for idx, line in enumerate(lines):
        r_elem = OxmlElement("w:r")
        if rpr is not None:
            r_elem.append(copy.deepcopy(rpr))
        t_elem = OxmlElement("w:t")
        t_elem.text = line
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r_elem.append(t_elem)
        para._p.append(r_elem)
        if idx < len(lines) - 1:
            r_br = OxmlElement("w:r")
            if rpr is not None:
                r_br.append(copy.deepcopy(rpr))
            br = OxmlElement("w:br")
            r_br.append(br)
            para._p.append(r_br)


def fill_cell(cell, text):
    """Replace cell text preserving formatting."""
    para = cell.paragraphs[0]
    set_para_text(para, text)


def fill_row(table, row_idx, values):
    """Fill a table row with a list of string values."""
    row = table.rows[row_idx]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            fill_cell(row.cells[ci], val)


# ── carga plantilla ──────────────────────────────────────────────────────────
doc = Document(TEMPLATE)

# ── TÍTULO PRINCIPAL ──────────────────────────────────────────────────────────
_, p = find_para(doc, "Blueprint – xxx", "Heading 1")
if p:
    set_para_text(p, "Blueprint – PyroguardIAn & SolarEyes: Monitoreo IA de Instalaciones Solares EGP")

# ── RESUMEN EJECUTIVO ─────────────────────────────────────────────────────────
_, p = find_para(doc, "1. Sommario Esecutivo", "Heading 2")
if p:
    set_para_text(p, "1. Resumen Ejecutivo")

_, p = find_para(doc, "##TITOLO##")
if p:
    set_para_text(p,
        "Procesos Identificados: Monitoreo de Instalaciones Fotovoltaicas - "
        "Deteccion de Anomalias Termicas e Inspeccion Visual con IA")

_, p = find_para(doc, "##contenuto##")
if p:
    set_para_lines(p, [
        "Contexto General:",
        "El proceso actual de monitoreo de las instalaciones fotovoltaicas EGP/TGX se basa en "
        "inspecciones fisicas periodicas realizadas por tecnicos en campo, consulta manual de "
        "dashboards SCADA/DCS y analisis reactivo de datos de produccion. La deteccion de hotspots "
        "termicos, celulas degradadas y riesgos de incendio se realiza principalmente mediante "
        "campanas de termografia infrarroja con cadencia semestral o anual, sin cobertura en "
        "tiempo real. El sistema actual no puede detectar anticipadamente anomalias generalizadas "
        "sobre grandes superficies de paneles.",
        "",
        "Propuesta PyroguardIAn & SolarEyes: introduccion de un sistema de IA compuesto por dos "
        "modulos integrados — (1) PyroguardIAn: agente IA para la deteccion en tiempo real de "
        "anomalias termicas y riesgos de incendio mediante procesamiento de imagenes termograficas "
        "(drones/sensores fijos); (2) SolarEyes: agente IA para la inspeccion visual automatizada "
        "de paneles mediante vision artificial (imagenes RGB drone/satelite) para detectar danos "
        "fisicos, suciedad y degradacion. Ambos modulos comparten un Agente Orquestador y un "
        "Dashboard HITL para la gestion de anomalias con supervision humana.",
        "",
        "Propietario del proceso: EGP/TGX - O&M Solar | Poblacion impactada: Tecnicos de oficina, "
        "Responsables de instalacion, Equipo HSE, Tecnicos de inspeccion, Proveedores de mantenimiento"
    ])

# ── PROCESO 1 – TÍTULO ────────────────────────────────────────────────────────
_, p = find_para(doc, "Processo 1: xxx", "Heading 1")
if p:
    set_para_text(p, "Proceso 1: Monitoreo IA de Instalaciones Fotovoltaicas - PyroguardIAn & SolarEyes")

# ── 1. CONTEXTO Y FINALIDAD ───────────────────────────────────────────────────
_, p = find_para(doc, "1. Contesto e finalita", "Heading 2")
if p:
    set_para_text(p, "1. Contexto y Finalidad")

# ── 1.1 ALCANCE ───────────────────────────────────────────────────────────────
_, p = find_para(doc, "1.1 Scopo", "Heading 3")
if p:
    set_para_text(p, "1.1 Alcance")
scopo_start = _ if _ is not None else 0
count = 0
for i, para in enumerate(doc.paragraphs):
    if i <= scopo_start:
        continue
    if para.style.name == "Heading 3":
        break
    if para.text.strip() in ("xxxx", "xxx"):
        if count == 0:
            set_para_text(para,
                "El proceso gestiona el monitoreo continuo de las instalaciones fotovoltaicas EGP, "
                "adquiriendo imagenes termograficas y RGB de los drones de inspeccion y de los "
                "sensores fijos instalados en las instalaciones, procesandolas con modelos de IA "
                "para la deteccion de anomalias termicas (hotspots, celulas bypass, riesgos de "
                "incendio) y defectos visuales (micro-grietas, suciedad, roturas fisicas, sombreado parcial).")
        elif count == 1:
            set_para_text(para,
                "Objetivo principal: prevenir incendios y maximizar la produccion de las "
                "instalaciones solares mediante la deteccion temprana de anomalias (PyroguardIAn) "
                "y la planificacion optimizada de intervenciones de mantenimiento basada en las "
                "inspecciones visuales (SolarEyes), reduciendo los costes de inspeccion manual "
                "y los tiempos de respuesta ante eventos criticos.")
        elif count == 2:
            set_para_text(para,
                "Sistemas en alcance: drones de inspeccion con payload termografico y RGB, "
                "sensores termicos fijos en string/inversor, sistema SCADA de la instalacion, "
                "plataforma de gestion de imagenes (Image Repository), CMMS para ordenes de "
                "trabajo, Base de Conocimiento tecnica fotovoltaico.")
        count += 1
        if count >= 3:
            break

# ── 1.2 FINALIDAD ─────────────────────────────────────────────────────────────
_, p = find_para(doc, "1.2 Finalit", "Heading 3")
if p:
    set_para_text(p, "1.2 Finalidad")
h12_idx = _ if _ is not None else 0
count = 0
for i, para in enumerate(doc.paragraphs):
    if i <= h12_idx:
        continue
    if para.style.name == "Heading 3":
        break
    if para.text.strip() in ("• xxx", "xxx"):
        if count == 0:
            set_para_text(para,
                "Reducir el riesgo de incendio en las instalaciones solares mediante la deteccion "
                "en tiempo real de hotspots termicos y condiciones anomalas (modulo PyroguardIAn), "
                "con alerta inmediata al equipo HSE y al responsable de la instalacion.")
        elif count == 1:
            set_para_text(para,
                "Automatizar la inspeccion visual de los paneles fotovoltaicos mediante analisis "
                "IA de imagenes drone (modulo SolarEyes), reduciendo la dependencia de campanas "
                "manuales semestrales/anuales y bajando el coste de inspeccion por MW.")
        elif count == 2:
            set_para_text(para,
                "Optimizar la planificacion de las intervenciones de mantenimiento basandose en "
                "anomalias detectadas y clasificadas por la IA, priorizadas por impacto en la "
                "produccion y riesgo de seguridad, con flujo de aprobacion HITL.")
        count += 1
        if count >= 3:
            break

# ── 1.3 PERIMETRO ─────────────────────────────────────────────────────────────
_, p = find_para(doc, "1.3 Perimetro", "Heading 3")
if p:
    set_para_text(p, "1.3 Perimetro")
h13_idx = _ if _ is not None else 0
in_scope_count = 0
out_scope_count = 0
in_out_mode = None
for i, para in enumerate(doc.paragraphs):
    if i <= h13_idx:
        continue
    if para.style.name in ("Heading 3", "Heading 2"):
        break
    txt = para.text.strip()
    if txt == "IN SCOPE:":
        set_para_text(para, "EN ALCANCE:")
        in_out_mode = "in"
    elif txt == "OUT OF SCOPE:":
        set_para_text(para, "FUERA DE ALCANCE:")
        in_out_mode = "out"
    elif txt == "- xxx":
        if in_out_mode == "in":
            if in_scope_count == 0:
                set_para_text(para,
                    "- Monitoreo termografico en tiempo real/casi tiempo real de instalaciones "
                    "fotovoltaicas EGP en alcance mediante PyroguardIAn")
            elif in_scope_count == 1:
                set_para_text(para,
                    "- Inspeccion visual automatizada de paneles mediante analisis IA de imagenes "
                    "RGB de drone (SolarEyes)")
            elif in_scope_count == 2:
                set_para_text(para,
                    "- Clasificacion IA de anomalias (hotspots, micro-grietas, suciedad, roturas) "
                    "con priorizacion por riesgo e impacto en produccion")
            in_scope_count += 1
        elif in_out_mode == "out":
            if out_scope_count == 0:
                set_para_text(para,
                    "- Intervencion fisica de sustitucion/limpieza de paneles (permanece a cargo "
                    "de tecnicos y proveedores externos)")
            elif out_scope_count == 1:
                set_para_text(para,
                    "- Gestion operativa de flotas de drones (planificacion de vuelos y "
                    "mantenimiento de drones no en alcance en el piloto)")
            out_scope_count += 1

# ── 1.4 RESTRICCIONES CLAVE ───────────────────────────────────────────────────
_, p = find_para(doc, "1.4 Vincoli chiave", "Heading 3")
if p:
    set_para_text(p, "1.4 Restricciones Clave")
h14_idx = _ if _ is not None else 0
_, p = find_para(doc, "Normativi:", "Normal", h14_idx)
if p:
    set_para_text(p, "Normativos:")
_, p = find_para(doc, "xxx", "List Paragraph", h14_idx)
if p:
    set_para_text(p,
        "Conformidad con normativas aeronauticas para uso de drones en instalaciones energeticas "
        "(ENAC D-LG-2020-001, reglamento UE 2019/947); cumplimiento de requisitos de ciberseguridad "
        "para sistemas OT/SCADA (IEC 62443); normativa contra incendios en instalaciones "
        "fotovoltaicas (CEI 82-25).")
_, p = find_para(doc, "Tecnici:", "Normal", h14_idx)
if p:
    set_para_text(p, "Tecnicos:")
_, p = find_para(doc, "xxx", "Normal", h14_idx)
if p:
    set_para_text(p,
        "Disponibilidad de conectividad de datos en las instalaciones para transmision de imagenes "
        "en tiempo real [A DEFINIR con IT Owner]; integracion con SCADA/DCS existente para "
        "correlacion datos energia-anomalias [A DEFINIR]; API CMMS para creacion automatica de "
        "ordenes de trabajo [A DEFINIR]; resolucion minima de imagenes drone para clasificacion "
        "IA (>=5 MP termico, >=20 MP RGB). "
        "\u26a0\ufe0f NOTA: Los nombres exactos de sistemas y protocolos de integracion deben "
        "ser definidos con IT Owner OT antes del kick-off del proyecto.")
_, p = find_para(doc, "Organizzativi:", "Normal", h14_idx)
if p:
    set_para_text(p, "Organizativos:")
_, p = find_para(doc, "xxx", "Normal", h14_idx)
if p:
    set_para_text(p,
        "Definicion del flujo de escalada para eventos criticos (incendio/humo) y "
        "responsabilidades HSE correspondientes; formacion de tecnicos de oficina y responsables "
        "de instalacion en el uso del Dashboard HITL; gobierno para aprobacion de intervenciones "
        "generadas por IA; gestion del cambio para transicion de inspecciones manuales periodicas "
        "a monitoreo continuo.")

# ── STAKEHOLDERS Y PARTICIPANTES ──────────────────────────────────────────────
_, p = find_para(doc, "Stakeholder & Partecipanti", "Heading 2")
if p:
    set_para_text(p, "Stakeholders y Participantes")

# TABLE 0 – Sistemas AS-IS (reusando la tabla de la plantilla, igual que el script italiano)
# NOTE: PREGUNTA ABIERTA — los nombres de roles/stakeholders se documentan como punto abierto
# Roles esperados: Business Owner, BSN Global, Factory AI, Initiative Leader,
#                  AI Hero, IT Owner/SME, AISA — [A COMPLETAR con nombres y unidades org. EGP TGX]
t0 = doc.tables[0]
fill_row(t0, 0, ["Sistema AS-IS", "Descripcion / Rol", "Tipologia"])
# NOTE: PREGUNTA ABIERTA — nombres exactos de sistemas a confirmar con IT Owner
sistemas_asis = [
    ("SCADA / DCS Instalacion",
     "Sistema de supervision y control de la instalacion fotovoltaica; proporciona datos de "
     "produccion, alarmas de string y estado del inversor. "
     "\u26a0\ufe0f NOTA: Nombre exacto del sistema [A DEFINIR con IT Owner OT/SCADA]",
     "SCADA / Control System"),
    ("Sistema de Gestion de Imagenes (Image Repository)",
     "Plataforma de almacenamiento y acceso a imagenes termograficas y RGB adquiridas por los "
     "drones; base de datos para analisis IA",
     "Image Management Platform"),
    ("Drone Termografico",
     "UAV con payload de camara termografica infrarroja para adquisicion de imagenes termicas "
     "de las instalaciones solares",
     "Inspection Hardware / UAV"),
    ("Drone RGB / Multiespectral",
     "UAV con camara RGB de alta resolucion para inspeccion visual de los paneles fotovoltaicos",
     "Inspection Hardware / UAV"),
    ("CMMS (Gestion de Mantenimiento)",
     "Sistema para la gestion de ordenes de trabajo e intervenciones de mantenimiento en las "
     "instalaciones. \u26a0\ufe0f NOTA: Nombre exacto del sistema [A DEFINIR con IT Owner]",
     "CMMS / EAM Platform"),
    ("Portal Meteorologico / Previsiones",
     "Servicio de prediccion meteorologica para planificacion de ventanas de vuelo y correlacion "
     "de produccion con irradiancia",
     "External Weather API"),
    ("Archivo Documental Tecnico",
     "Repositorio de documentacion tecnica de paneles (fichas tecnicas, manuales OEM, historico "
     "de intervenciones anteriores)",
     "Document Management (SharePoint/PLM)"),
]
for i, (s, r, t) in enumerate(sistemas_asis):
    fill_row(t0, i + 1, [s, r, t])

# ── 2. PROCESO AS-IS – DESCRIPCION ESTRUCTURADA ───────────────────────────────
_, p = find_para(doc, "2. Processo AS-IS", "Heading 2")
if p:
    set_para_text(p, "2. Proceso AS-IS - Descripcion Estructurada")

# ── 2.1 SISTEMAS INVOLUCRADOS (AS-IS) ─────────────────────────────────────────
_, p = find_para(doc, "2.1 Sistemi coinvolti", "Heading 3")
if p:
    set_para_text(p, "2.1 Sistemas Involucrados (AS-IS)")

# ── 2.2 AS-IS – SECUENCIA OPERATIVA ──────────────────────────────────────────
_, p = find_para(doc, "2.2 AS-IS", "Heading 3")
if p:
    set_para_text(p, "2.2 AS-IS - Secuencia Operativa")

_, pa = find_para(doc, "Sotto-processo A: xxx", "Normal")
if pa:
    set_para_text(pa, "Subproceso A: Planificacion y Ejecucion de Campanas de Inspeccion Termografica")
_, pb = find_para(doc, "Sotto-processo B: xxx", "Normal")
if pb:
    set_para_text(pb, "Subproceso B: Analisis Manual de Imagenes e Identificacion de Anomalias")
_, pc = find_para(doc, "Sotto-processo C: xxx", "Normal")
if pc:
    set_para_text(pc, "Subproceso C: Planificacion y Coordinacion de Intervenciones de Mantenimiento")

# TABLE 2 – AS-IS Secuencia A
t2 = doc.tables[1]
fill_row(t2, 0, ["Paso", "Actividad", "Actor", "Entrada", "Salida", "Sistemas"])
rows_a = [
    ("A1",
     "Equipo O&M planifica la campana de inspeccion termografica semestral/anual para la instalacion",
     "Responsable O&M",
     "Calendario de mantenimiento, plan anual de inspecciones",
     "Campana de inspeccion planificada con fechas e instalaciones objetivo",
     "Calendario ERP / Email"),
    ("A2",
     "Coordinacion con operador de drone certificado (interno o proveedor externo) para disponibilidad y autorizaciones de vuelo",
     "Responsable O&M / Proveedor",
     "Disponibilidad del operador, autorizaciones ENAC",
     "Operador y drone reservados; autorizaciones obtenidas",
     "Email / Telefono"),
    ("A3",
     "Operador de drone realiza vuelo de inspeccion sobre la instalacion con payload de camara termografica IR",
     "Operador drone / Tecnico",
     "Plan de vuelo, condiciones meteorologicas favorables",
     "Imagenes termograficas adquiridas (formato TIFF/RJPEG)",
     "Drone + Payload Termocamara"),
    ("A4",
     "Imagenes termograficas transferidas manualmente a estacion de trabajo/servidor empresarial",
     "Tecnico de oficina",
     "Archivos de imagenes del drone (tarjeta SD / enlace)",
     "Dataset de imagenes archivado localmente o en servidor",
     "Transferencia de archivos manual (USB/FTP)"),
    ("A5",
     "Operador drone realiza vuelo RGB adicional para documentacion visual y georreferenciacion de anomalias",
     "Operador drone",
     "Plan de vuelo ejecutado tras termografia",
     "Imagenes RGB georreferenciadas adquiridas",
     "Drone + Camara RGB"),
]
for i, row_data in enumerate(rows_a):
    if i + 1 < len(t2.rows):
        fill_row(t2, i + 1, list(row_data))

# TABLE 3 – AS-IS Secuencia B
t3 = doc.tables[2]
fill_row(t3, 0, ["Paso", "Actividad", "Actor", "Entrada", "Salida", "Sistemas"])
rows_b = [
    ("B1",
     "Tecnico especializado analiza manualmente las imagenes termograficas con software de procesamiento (ej. FLIR Tools, IrfanView)",
     "Tecnico termografista",
     "Dataset de imagenes termograficas",
     "Lista de anomalias termicas identificadas (hotspots, fallo bypass diode, etc.)",
     "Software termografia (FLIR Tools / ResearchIR)"),
    ("B2",
     "Clasificacion manual de anomalias segun estandar IEC 62446-3 (niveles de severidad 1-3)",
     "Tecnico termografista",
     "Lista de anomalias identificadas",
     "Anomalias clasificadas por severidad y tipologia",
     "Software termografia + hoja Excel"),
    ("B3",
     "Tecnico elabora informe de inspeccion en formato Word/Excel con fotos y geolocalizacion manual de anomalias",
     "Tecnico termografista",
     "Lista de anomalias clasificadas + imagenes de detalle",
     "Informe de inspeccion termografica (Word/Excel/PDF)",
     "Microsoft Word / Excel"),
    ("B4",
     "Informe enviado por email al Responsable de instalacion y al equipo O&M para planificacion de intervenciones",
     "Tecnico termografista",
     "Informe completado",
     "Email con informe adjunto",
     "Email (Outlook)"),
    ("B5",
     "Responsable de instalacion revisa el informe y evalua prioridades de intervencion",
     "Responsable de instalacion",
     "Informe de inspeccion",
     "Lista de intervenciones priorizadas (subjetiva)",
     "Email / Reunion"),
]
for i, row_data in enumerate(rows_b):
    if i + 1 < len(t3.rows):
        fill_row(t3, i + 1, list(row_data))

# TABLE 4 – AS-IS Secuencia C (doc.tables[4] in template)
t4_asis = doc.tables[3]
fill_row(t4_asis, 0, ["Paso", "Actividad", "Actor", "Entrada", "Salida", "Sistemas"])
rows_c = [
    ("C1",
     "Responsable de instalacion verifica la disponibilidad de empresas y tecnicos de mantenimiento",
     "Responsable de instalacion",
     "Lista de intervenciones priorizadas",
     "Disponibilidad de tecnicos/proveedores verificada",
     "Telefono / Email"),
    ("C2",
     "Verificacion de disponibilidad de repuestos (celulas, paneles de sustitucion) y plazos de aprovisionamiento",
     "Responsable de instalacion",
     "Lista de componentes a sustituir",
     "Disponibilidad de repuestos y lead time conocidos",
     "ERP / Telefono proveedores"),
    ("C3",
     "Planificacion manual de intervenciones con definicion de ventanas operativas favorables (meteorologia, produccion)",
     "Responsable de instalacion",
     "Restricciones de disponibilidad, meteorologia, repuestos",
     "Plan de intervenciones (hoja Excel o CMMS manual)",
     "Excel / CMMS"),
    ("C4",
     "Emision de ordenes de trabajo en el CMMS y comunicacion a proveedores/tecnicos",
     "Responsable de instalacion / Back-office",
     "Plan de intervenciones aprobado",
     "Ordenes de trabajo emitidas en el CMMS",
     "CMMS [A DEFINIR]"),
    ("C5",
     "Ejecucion de intervenciones (sustitucion de paneles, limpieza, reparacion de conexiones) por tecnicos/proveedores",
     "Tecnico de mantenimiento / Proveedor",
     "Ordenes de trabajo recibidas",
     "Intervencion ejecutada; acta en papel cumplimentada",
     "En campo / Papel"),
    ("C6",
     "Verificacion post-intervencion: nuevo vuelo termografico programado (en la siguiente campana) para confirmar eficacia",
     "Responsable O&M",
     "Acta de intervencion",
     "Planificacion de verificacion post-intervencion (ciclo siguiente)",
     "Email / Calendario"),
]
for i, row_data in enumerate(rows_c):
    if i + 1 < len(t4_asis.rows):
        fill_row(t4_asis, i + 1, list(row_data))

# ── 2.3 FICHAS DE PROCESO AS-IS ───────────────────────────────────────────────
_, p = find_para(doc, "2.3 Process Cards AS-IS", "Heading 3")
if p:
    set_para_text(p, "2.3 Fichas de Proceso AS-IS")

_, p = find_para(doc, "Card 1: xxx")
if p:
    set_para_text(p, "Ficha 1: Planificacion y Ejecucion de Campana de Inspeccion Termografica")
idx_c1 = _ if _ is not None else 0

c1_items = []
for i, para in enumerate(doc.paragraphs):
    if i <= idx_c1:
        continue
    if para.text.strip() == "Card 2: xxx":
        break
    if para.text.strip() in ("Input: xxx ", "Input: xxx", "Attivita: xxx", "Attività: xxx",
                              "Output: xxx", "Pain point: xxx"):
        c1_items.append(para)

if len(c1_items) >= 1:
    set_para_text(c1_items[0],
        "Entrada: Calendario de mantenimiento anual; disponibilidad de operadores de drone y "
        "autorizaciones ENAC; condiciones meteorologicas")
if len(c1_items) >= 2:
    set_para_lines(c1_items[1], [
        "Actividades:",
        "1. Equipo O&M planifica la campana de inspeccion (fechas, instalaciones, operadores)",
        "2. Coordinacion con operador de drone para disponibilidad y autorizaciones de vuelo",
        "3. Vuelo termografico IR sobre la instalacion con adquisicion de imagenes",
        "4. Vuelo RGB para georreferenciacion y documentacion visual",
        "5. Transferencia manual de imagenes al servidor empresarial"
    ])
if len(c1_items) >= 3:
    set_para_text(c1_items[2],
        "Salida: Dataset de imagenes termograficas y RGB archivado; campana ejecutada con "
        "cadencia semestral/anual")
if len(c1_items) >= 4:
    set_para_text(c1_items[3],
        "Punto Critico: Cadencia de inspeccion demasiado baja para detectar eventos criticos "
        "en tiempo real; coste elevado por campana (operador + vuelo)")

_, p = find_para(doc, "Card 2: xxx")
if p:
    set_para_text(p, "Ficha 2: Analisis Manual de Imagenes e Identificacion de Anomalias")
idx_c2 = _ if _ is not None else 0

c2_items = []
for i, para in enumerate(doc.paragraphs):
    if i <= idx_c2:
        continue
    if para.text.strip() == "Card 3: xxx":
        break
    if para.text.strip() in ("Input: xxx ", "Input: xxx", "Attivita: xxx", "Attività: xxx",
                              "Output: xxx", "Pain point: xxx"):
        c2_items.append(para)

if len(c2_items) >= 1:
    set_para_text(c2_items[0],
        "Entrada: Dataset de imagenes termograficas y RGB adquiridas durante la campana de inspeccion")
if len(c2_items) >= 2:
    set_para_lines(c2_items[1], [
        "Actividades:",
        "1. Tecnico termografista analiza imagenes con software dedicado (FLIR Tools / ResearchIR)",
        "2. Identificacion manual de hotspots, fallo bypass, sombreado, suciedad",
        "3. Clasificacion de anomalias por severidad (IEC 62446-3)",
        "4. Elaboracion de informe Word/Excel con coordenadas de anomalias",
        "5. Envio de informe por email al responsable de instalacion"
    ])
if len(c2_items) >= 3:
    set_para_text(c2_items[2],
        "Salida: Informe de inspeccion con lista de anomalias clasificadas; enviado al "
        "responsable de instalacion")
if len(c2_items) >= 4:
    set_para_text(c2_items[3],
        "Punto Critico: Analisis manual lento (~2-5 dias por instalacion media); subjetivo; "
        "sin correlacion automatica con datos SCADA")

_, p = find_para(doc, "Card 3: xxx")
if p:
    set_para_text(p, "Ficha 3: Planificacion y Coordinacion de Intervenciones de Mantenimiento")
idx_c3 = _ if _ is not None else 0

c3_items = []
for i, para in enumerate(doc.paragraphs):
    if i <= idx_c3:
        continue
    if para.style.name in ("Heading 3", "Heading 2"):
        break
    if para.text.strip() in ("Input: xxx", "Input: xxx ", "Attivita: xxx", "Attività: xxx",
                              "Output: xxx", "Pain point: xxx"):
        c3_items.append(para)

if len(c3_items) >= 1:
    set_para_text(c3_items[0],
        "Entrada: Informe con lista de anomalias priorizadas recibido del tecnico termografista; "
        "disponibilidad de tecnicos, proveedores y repuestos (verificacion manual)")
if len(c3_items) >= 2:
    set_para_lines(c3_items[1], [
        "Actividades:",
        "1. Responsable de instalacion revisa el informe y define prioridades (subjetivas)",
        "2. Verificacion manual de disponibilidad de tecnicos, proveedores y repuestos",
        "3. Planificacion de intervenciones en Excel o CMMS manual",
        "4. Emision de ordenes de trabajo y comunicacion a tecnicos/proveedores",
        "5. Ejecucion de intervenciones y cumplimentacion de acta en papel"
    ])
if len(c3_items) >= 3:
    set_para_text(c3_items[2],
        "Salida: Intervenciones ejecutadas; actas en papel; sin verificacion post-intervencion "
        "a corto plazo")
if len(c3_items) >= 4:
    set_para_text(c3_items[3],
        "Punto Critico: Planificacion no optimizada; gap temporal elevado entre deteccion de "
        "anomalia e intervencion (semanas/meses)")

# ── 2.4 PUNTOS CRITICOS PRINCIPALES (AS-IS) ───────────────────────────────────
_, p = find_para(doc, "2.4 Pain point principali", "Heading 3")
if p:
    set_para_text(p, "2.4 Puntos Criticos Principales (AS-IS)")

_, p = find_para(doc, "Pain Point Generali", "Normal")
if p:
    set_para_text(p, "Puntos Criticos Generales")

pain_points_es = [
    "1. Esfuerzo FTE elevado: analisis manual de imagenes termograficas requiere termografistas "
    "certificados; estimado 2-5 dias FTE por instalacion media por campana",
    "2. Tiempos de ciclo largos: desde la planificacion de la campana hasta la intervencion de "
    "mantenimiento transcurren de media 4-12 semanas; riesgo de incendio no gestionado en tiempo real",
    "3. Cuellos de botella: disponibilidad limitada de termografistas certificados y drones; "
    "imposibilidad de cubrir continuamente el parque de instalaciones; crecimiento del parque "
    "solar no escalable con el enfoque actual",
    "4. Calidad no uniforme: la calidad del analisis termografico depende de la experiencia del "
    "termografista individual; sin estandar automatizado de clasificacion; variabilidad en el reporting",
    "5. Escasa escalabilidad: con el crecimiento del parque solar EGP, el modelo de inspecciones "
    "periodicas manuales no aguanta el ritmo; imposibilidad de monitorear continuamente superficies "
    "de centenares de MW"
]
for num_txt in ["1. xxx", "2. xxx", "3. xxx", "4. xxx", "5. xxx"]:
    _, p = find_para(doc, num_txt, "Normal")
    if p:
        idx = int(num_txt.split(".")[0]) - 1
        if idx < len(pain_points_es):
            set_para_text(p, pain_points_es[idx])

# ── 2.6 MAPEO DE DATOS-SISTEMAS ───────────────────────────────────────────────
_, p = find_para(doc, "2.6 Mapping Dati-SistemI", "Heading 3")
if p:
    set_para_text(p, "2.6 Mapeo de Datos-Sistemas")

# TABLE 5 – Data Mapping
# NOTE: PREGUNTA ABIERTA — formatos API a confirmar con IT Owner
t5_dm = doc.tables[4]
fill_row(t5_dm, 0, ["Dato", "Sistema Origen", "Sistema Destino", "Formato", "Nota"])
data_mapping = [
    ("Imagen termografica RAW (TIFF radiometrico)",
     "Drone Termografico",
     "Image Repository / Motor IA PyroguardIAn",
     "TIFF / RJPEG con metadata GPS",
     "Adquirida por cada sesion de vuelo; trigger analisis IA en tiempo real"),
    ("Imagen RGB georreferenciada (inspeccion visual)",
     "Drone RGB",
     "Image Repository / Motor IA SolarEyes",
     "GeoTIFF / JPEG con metadata GPS",
     "Adquirida en paralelo al vuelo termografico o en campana dedicada"),
    ("Resultado clasificacion anomalia termica",
     "Motor IA PyroguardIAn",
     "Dashboard HITL / CMMS",
     "JSON estructurado (ID modulo, tipo anomalia, severidad, coordenadas)",
     "Generado en tiempo real tras procesamiento de imagen"),
    ("Resultado clasificacion defecto visual panel",
     "Motor IA SolarEyes",
     "Dashboard HITL / CMMS",
     "JSON estructurado (ID modulo, tipo defecto, score confianza, bounding box)",
     "Generado tras analisis de imagen RGB; batch o casi tiempo real"),
    ("Dato de produccion string / inversor",
     "SCADA / DCS Instalacion",
     "PyroguardIAn / SolarEyes Motor IA",
     "JSON / Modbus / OPC-UA [A DEFINIR] - \u26a0\ufe0f NOTA: protocolo a definir con IT Owner OT",
     "Correlacion automatica anomalia termica - caida de produccion"),
    ("Alarma SCADA (string off, fallo inversor)",
     "SCADA / DCS Instalacion",
     "Agente Orquestador / Dashboard HITL",
     "JSON REST / Webhook [A DEFINIR] - \u26a0\ufe0f NOTA: protocolo a definir con IT Owner",
     "Trigger contextual para analisis de anomalia asociada"),
    ("Orden de trabajo generada por IA",
     "Agente Orquestador (TO-BE)",
     "CMMS",
     "JSON / API REST CMMS [A DEFINIR] - \u26a0\ufe0f NOTA: API CMMS a definir con IT Owner",
     "Flujo HITL: aprobacion del Responsable de instalacion antes de crear la OT"),
    ("Acta post-intervencion digital",
     "Tecnico mantenimiento (App Movil TO-BE)",
     "Base de Conocimiento / CMMS / Dashboard HITL",
     "JSON estructurado / PDF",
     "Feedback loop: alimenta reentrenamiento de modelos IA"),
    ("Datos meteorologicos e irradiancia solar",
     "API Meteorologica Externa (ej. Solargis / ECMWF)",
     "SolarEyes / PyroguardIAn Motores IA",
     "JSON / CSV REST API",
     "Normalizacion de anomalias por condiciones de irradiancia; planificacion de vuelos"),
    ("Ficha tecnica / datasheet panel fotovoltaico",
     "Archivo Documental Tecnico (SharePoint)",
     "Base de Conocimiento (RAG)",
     "PDF / Word",
     "Referencia para umbrales termicos normales por modelo de panel especifico"),
    ("Historico de intervenciones de mantenimiento",
     "CMMS / Archivo en papel",
     "Base de Conocimiento (RAG)",
     "PDF / Excel / JSON",
     "Digitalizacion necesaria; base para modelos predictivos de degradacion"),
    ("Informe de inspeccion termografica historico",
     "Archivo documental",
     "Base de Conocimiento (RAG) / Dataset Training IA",
     "PDF / Word / Excel",
     "Etiquetado retroactivo para dataset de entrenamiento de modelos IA"),
]
for i, row_data in enumerate(data_mapping):
    if i + 1 < len(t5_dm.rows):
        fill_row(t5_dm, i + 1, list(row_data))

# ── 3. PROCESO TO-BE ──────────────────────────────────────────────────────────
_, p = find_para(doc, "3. Processo TO-BE", "Heading 2")
if p:
    set_para_text(p, "3. Proceso TO-BE")

# ── 3.1 ARQUITECTURA FUNCIONAL IA ─────────────────────────────────────────────
_, p = find_para(doc, "3.1 Architettura Funzionale AI", "Heading 3")
if p:
    set_para_text(p, "3.1 Arquitectura Funcional IA")

_, p = find_para(doc, "Pattern tassonomico Catalogo Soluzioni AI:", "Normal")
if p:
    set_para_text(p,
        "Patron taxonomico Catalogo Soluciones IA: "
        "Arquitectura multi-agente IA compuesta por 3 modulos especializados coordinados por "
        "un Orquestador. (1) Orchestrator Agent - coordina el flujo extremo a extremo entre "
        "PyroguardIAn y SolarEyes, gestiona priorizacion de anomalias, flujo HITL e integracion "
        "con CMMS/Dashboard. Patron: Orchestrator/Coordinator Agent. "
        "(2) PyroguardIAn AI Engine - analiza imagenes termograficas IR para detectar hotspots, "
        "fallo bypass diode y riesgos de incendio; clasifica por severidad IEC 62446-3; integra "
        "datos SCADA para correlacion anomalia termica - caida de produccion. Patron: "
        "Analyzer/Classifier Agent (Computer Vision + LLM reasoning). "
        "(3) SolarEyes AI Engine - analiza imagenes RGB drone para detectar defectos visuales "
        "(micro-grietas, suciedad, roturas fisicas); genera mapa geoespacial de defectos. "
        "Patron: Analyzer/Detector Agent (Computer Vision Object Detection). "
        "(4) Knowledge Agent (transversal) - recupera y sintetiza informacion tecnica de la "
        "Base de Conocimiento para contextualizar los analisis IA. Patron: Retriever/RAG Agent.")

# TABLE 6 – Arquitectura Funcional TO-BE
t6 = doc.tables[5]
fill_row(t6, 0, ["Componente", "Funcion", "Tecnologia/Metodo", "Agente Comun"])
arch_rows = [
    ("Agente Orquestador",
     "Coordina el flujo extremo a extremo entre los modulos PyroguardIAn y SolarEyes; gestiona "
     "la priorizacion de anomalias, el flujo HITL y la interfaz con CMMS y Dashboard",
     "LLM multi-agent orchestration (LangGraph / Azure AI Agent Service)",
     ""),
    ("Motor IA PyroguardIAn",
     "Analiza imagenes termograficas IR en tiempo real / casi tiempo real para detectar hotspots, "
     "fallo bypass diode, suciedad termica y potenciales riesgos de incendio; clasifica por "
     "severidad IEC 62446-3",
     "Computer Vision CNN/ViT + Thermal Image Processing (Python/ONNX/TensorRT)",
     ""),
    ("Motor IA SolarEyes",
     "Analiza imagenes RGB drone de alta resolucion para detectar micro-grietas, roturas fisicas, "
     "suciedad, excrementos de aves y sombreado parcial de los paneles",
     "Computer Vision CNN/YOLOv8 + RGB Image Analysis (Python/PyTorch)",
     ""),
    ("ETL & Pipeline de Integracion de Datos",
     "Adquiere y normaliza imagenes de drones/sensores, datos SCADA, meteorologia y CMMS; "
     "alimenta los motores IA con datos contextualizados y georreferenciados",
     "Azure Data Factory / Apache Kafka / Delta Lake (streaming y batch)",
     ""),
    ("Base de Conocimiento Fotovoltaico (RAG)",
     "Repositorio indexado de documentacion tecnica (fichas de paneles, manuales OEM, historico "
     "de intervenciones, informes de inspecciones pasadas) para contextualizacion IA y soporte "
     "a la toma de decisiones",
     "RAG + Vector DB (Azure AI Search / Qdrant) + LLM (GPT-4o)",
     ""),
    ("Dashboard HITL (TO-BE)",
     "Interfaz del operador para visualizacion del mapa de anomalias en la instalacion, informes "
     "IA, flujo de aprobacion de intervenciones y KPI en tiempo real",
     "Power BI Embedded / React Web App / MS Teams Bot",
     ""),
    ("Modulo de Validacion HITL y Alertas",
     "Gestiona el flujo de notificacion y aprobacion humana para eventos criticos (incendio/humo: "
     "escalada inmediata HSE) e intervenciones de mantenimiento planificadas",
     "Power Automate / Motor de flujo personalizado / Push notification movil",
     ""),
    ("Pipeline de Feedback y Reentrenamiento",
     "Recoge las actas post-intervencion y las correcciones HITL para alimentar el reentrenamiento "
     "periodico de los modelos IA y mejorar la precision con el tiempo",
     "MLflow / Azure ML / CI-CD pipeline reentrenamiento",
     ""),
]
for i, row_data in enumerate(arch_rows):
    if i + 1 < len(t6.rows):
        fill_row(t6, i + 1, list(row_data))

# ── 3.1b PUNTOS ABIERTOS – INTEGRACIONES Y ACCESOS ───────────────────────────
_, p = find_para(doc, "3.1b Punti Aperti", "Heading 3")
if p:
    set_para_text(p, "3.1b Puntos Abiertos - Integraciones y Accesos")

_, p = find_para(doc, "- xxx: esposizione API interna/esterna [DA DEFINIRE]")
if p:
    set_para_text(p,
        "- SCADA/DCS: modalidad de exposicion de datos de produccion y alarmas de string hacia "
        "el ETL Pipeline y Motor IA [A DEFINIR con IT Owner OT/SCADA]")
_, p = find_para(doc, "- xxx: API disponibile? [DA DEFINIRE]")
if p:
    set_para_text(p,
        "- CMMS: API disponible para creacion automatica de OT desde IA y lectura de historico "
        "de intervenciones [A DEFINIR]; autenticacion y permisos de escritura a coordinar")
_, p = find_para(doc, "- xxx: modalita lettura/scrittura [DA DEFINIRE]")
if p:
    set_para_text(p,
        "- Image Repository: modalidad de ingestion de imagenes drone (push automatico post-vuelo "
        "vs. carga manual); metadata estructurado minimo requerido por Motor IA [A DEFINIR]")
_, p = find_para(doc, "- xxx: accesso API [DA DEFINIRE]")
if p:
    set_para_text(p,
        "- Drones / Fleet Management: disponibilidad de API para programacion de misiones y "
        "recuperacion automatica de metadata de vuelo; compatibilidad del payload de termocamara "
        "con estandares de datos del Motor IA [A DEFINIR]")
_, p = find_para(doc, "I seguenti accessi/API sono [DA DEFINIRE]")
if p:
    set_para_text(p,
        "Los siguientes accesos/API deben ser [A DEFINIR] con IT Owner y OT Owner antes del "
        "kick-off del proyecto: SCADA/DCS, CMMS, Image Repository, Fleet Management Drones.")

# ── 3.2 SECUENCIA OPERATIVA DETALLADA (TO-BE) ─────────────────────────────────
_, p = find_para(doc, "3.2 Sequenza Operativa di Dettaglio", "Heading 3")
if p:
    set_para_text(p, "3.2 Secuencia Operativa Detallada (TO-BE)")

_, p = find_para(doc, "Sotto-processo A: xxx (TO-BE)")
if p:
    set_para_text(p, "Subproceso A: Adquisicion de Imagenes, Analisis IA y Clasificacion de Anomalias (TO-BE)")
_, p = find_para(doc, "Sotto-processo B: xxx (TO-BE)")
if p:
    set_para_text(p, "Subproceso B: Gestion de Alertas, Flujo HITL y Planificacion de Intervenciones (TO-BE)")
_, p = find_para(doc, "Sotto-processo C: XXXX (TO-BE)")
if p:
    set_para_text(p, "Subproceso C: Ejecucion de Intervenciones, Feedback Loop y Mejora Continua de Modelos IA (TO-BE)")

# TABLE 7 – TO-BE Secuencia A
t7 = doc.tables[6]
fill_row(t7, 0, ["Paso", "Actividad", "Actor", "Entrada", "Salida", "Sistemas", "IA+Intervencion Humana"])
rows_tobe_a = [
    ("A1",
     "ETL Pipeline adquiere automaticamente imagenes del drone post-vuelo (o en streaming desde "
     "sensores fijos) y datos SCADA de produccion",
     "ETL Pipeline (automatico)",
     "Imagenes IR/RGB drone, datos SCADA string/inversor, datos meteorologicos",
     "Dataset normalizado y enriquecido con metadata geoespacial y de produccion",
     "Image Repository, SCADA, API Meteorologica",
     "Ingestion automatica; calidad de datos verificada antes del analisis"),
    ("A2",
     "Motor IA PyroguardIAn analiza las imagenes termograficas y detecta hotspots, fallo bypass "
     "y anomalias termicas por cada modulo",
     "Motor IA PyroguardIAn (automatico)",
     "Dataset imagenes IR con metadata GPS y produccion",
     "Mapa termico de anomalias con: ID modulo, tipo anomalia, DeltaT, severidad IEC",
     "PyroguardIAn AI Engine, Image Repository",
     "CV inference + LLM reasoning para clasificacion contextual"),
    ("A3",
     "Motor IA SolarEyes analiza las imagenes RGB y detecta defectos visuales (micro-grietas, "
     "suciedad, roturas fisicas) en cada panel",
     "Motor IA SolarEyes (automatico)",
     "Dataset imagenes RGB georreferenciadas",
     "Mapa de defectos visuales con: ID modulo, tipo defecto, score confianza, bounding box",
     "SolarEyes AI Engine, Image Repository",
     "CV Object Detection (YOLOv8) + clasificacion de severidad"),
    ("A4",
     "Agente Orquestador correlaciona los resultados de PyroguardIAn y SolarEyes con datos SCADA "
     "y base de conocimiento para contexto y estimacion de impacto en produccion",
     "Agente Orquestador",
     "Output PyroguardIAn + SolarEyes + datos SCADA + Base Conocimiento",
     "Informe de anomalias enriquecido con: impacto estimado en produccion, notas tecnicas de "
     "la BC, recomendacion de intervencion",
     "Todos los sistemas IA, BC, SCADA",
     "LLM synthesis + RAG retrieval desde Base de Conocimiento"),
    ("A5",
     "Para anomalias criticas (riesgo incendio/humo): HITL inmediato - notificacion push a HSE "
     "y Responsable de instalacion con datos e imagen de anomalia",
     "Modulo Validacion HITL (automatico a humano)",
     "Anomalia Critical de PyroguardIAn",
     "Notificacion urgente enviada; flujo de escalada HSE activado",
     "Dashboard HITL, Push Movil, Teams",
     "Alerta inmediata - Humano decide accion de emergencia"),
    ("A6",
     "Para anomalias no criticas: Tecnico de oficina revisa el informe IA en Dashboard HITL, "
     "valida las clasificaciones y aprueba la lista de intervenciones",
     "Tecnico de oficina / Responsable O&M",
     "Informe IA consolidado en Dashboard HITL",
     "Lista de intervenciones validada con prioridades confirmadas",
     "Dashboard HITL",
     "Asistencia - Humano revisa y aprueba clasificacion IA"),
]
for i, row_data in enumerate(rows_tobe_a):
    if i + 1 < len(t7.rows):
        fill_row(t7, i + 1, list(row_data))

# TABLE 8 – TO-BE Secuencia B
t8 = doc.tables[7]
fill_row(t8, 0, ["Paso", "Actividad", "Actor", "Entrada", "Salida", "Sistemas", "IA+Intervencion Humana"])
rows_tobe_b = [
    ("B1",
     "Agente Orquestador recibe la lista de intervenciones validadas e inicia la recopilacion de "
     "restricciones de planificacion",
     "Agente Orquestador",
     "Lista de intervenciones aprobadas (del Tecnico HITL)",
     "Recopilacion de restricciones iniciada",
     "Orquestador, API sistemas externos",
     "Recopilacion automatica de restricciones via function calling"),
    ("B2",
     "Motor IA consulta API meteorologica para previsiones a 10 dias y verifica ventanas de "
     "irradiancia favorables para maximizar la eficacia de las intervenciones",
     "Agente Orquestador",
     "Lista de instalaciones con intervenciones",
     "Previsiones meteorologicas + irradiancia para ventanas de mantenimiento optimas",
     "API Meteorologica (Solargis / OpenWeather)",
     "Llamada API automatica + reasoning LLM"),
    ("B3",
     "Orquestador verifica disponibilidad de repuestos (paneles de sustitucion, bypass diode, "
     "conectores) consultando CMMS/ERP",
     "Agente Orquestador",
     "Lista de componentes a sustituir",
     "Disponibilidad de repuestos, lead time, proveedores alternativos",
     "CMMS / ERP Repuestos [A DEFINIR API]",
     "Consulta BD + LLM para analisis de disponibilidad"),
    ("B4",
     "Orquestador genera plan de intervenciones optimizado considerando criticidad de anomalia, "
     "impacto en produccion, meteorologia y disponibilidad de recursos",
     "Agente Orquestador / Modulo de Planificacion",
     "Intervenciones, meteorologia, repuestos, produccion perdida estimada",
     "Plan de intervenciones optimizado con programacion y prioridades IA",
     "Orquestador (optimizador de restricciones)",
     "Optimizacion multi-restriccion + LLM"),
    ("B5",
     "HITL - Responsable de instalacion revisa el plan en el Dashboard, puede modificarlo y lo aprueba",
     "Responsable de instalacion",
     "Plan IA con explicaciones y motivaciones",
     "Plan aprobado (o modificado y aprobado)",
     "Dashboard HITL",
     "Asistencia - Humano decide y aprueba el plan"),
    ("B6",
     "Post-aprobacion: Orquestador crea automaticamente las OT en el CMMS y notifica a los "
     "tecnicos/proveedores asignados",
     "Agente Orquestador (post-HITL)",
     "Plan aprobado",
     "OT creadas en el CMMS; notificaciones enviadas a tecnicos/proveedores",
     "CMMS, Email/Teams",
     "Automatizacion post-aprobacion HITL"),
]
for i, row_data in enumerate(rows_tobe_b):
    if i + 1 < len(t8.rows):
        fill_row(t8, i + 1, list(row_data))

# TABLE 9 – TO-BE Secuencia C
t9 = doc.tables[8]
fill_row(t9, 0, ["Paso", "Actividad", "Actor", "Entrada", "Salida", "Sistemas", "IA+Intervencion Humana"])
rows_tobe_c = [
    ("C1",
     "Tecnicos/proveedores reciben OT estructuradas con geolocalizacion de anomalia, imagen de "
     "detalle e instrucciones IA de la Base de Conocimiento",
     "Tecnico de mantenimiento / Proveedor",
     "OT del CMMS con datos IA adjuntos",
     "Intervencion iniciada con informacion completa",
     "CMMS, App Movil TO-BE, Dashboard HITL",
     "OT enriquecida por IA con instrucciones contextuales"),
    ("C2",
     "Ejecucion de intervencion fisica (sustitucion de paneles, limpieza, reparacion de conexiones)",
     "Tecnico de mantenimiento / Proveedor",
     "OT estructurada con geolocalizacion de anomalia",
     "Intervencion ejecutada en campo",
     "En campo",
     "N/A - actividad humana fisica"),
    ("C3",
     "Tecnico cumplimenta acta post-intervencion digital via App Movil (tipo anomalia confirmada, "
     "accion ejecutada, resultado)",
     "Tecnico de mantenimiento",
     "Intervencion completada",
     "Acta digital estructurada (tipo anomalia, accion, componentes sustituidos, resultado)",
     "App Movil TO-BE",
     "Formulario guiado por IA con sugerencias de cumplimentacion basadas en tipo de anomalia"),
    ("C4",
     "SolarEyes/PyroguardIAn verifican automaticamente la eficacia de la intervencion mediante "
     "comparacion de imagenes pre/post (siguiente sesion drone)",
     "Motor IA (post-intervencion)",
     "Imagenes post-intervencion + acta digital",
     "Verificacion de eficacia: anomalia resuelta / parcialmente resuelta / persistente",
     "Image Repository, Motores IA",
     "Comparacion automatica pre/post + evaluacion LLM"),
    ("C5",
     "Feedback loop: actas y resultados post-intervencion alimentan el reentrenamiento periodico "
     "de los modelos IA (mensual/trimestral)",
     "Pipeline de Feedback y Reentrenamiento",
     "Actas digitales + etiquetas de correccion HITL + resultados verificaciones",
     "Modelos IA actualizados con nuevos ejemplos; precision monitoreada",
     "MLflow / Azure ML",
     "Pipeline de aprendizaje continuo automatizado"),
    ("C6",
     "Orquestador calcula y actualiza KPI en el Dashboard: tasa de deteccion, falsos positivos, "
     "MTTR, produccion recuperada",
     "Agente Orquestador",
     "Datos de intervenciones completadas, KPI de produccion SCADA",
     "Dashboard KPI actualizado en tiempo real",
     "Dashboard HITL, SCADA, CMMS",
     "Calculo automatico de KPI + insights LLM mensuales"),
]
for i, row_data in enumerate(rows_tobe_c):
    if i + 1 < len(t9.rows):
        fill_row(t9, i + 1, list(row_data))

# ── 3.3 FICHAS DE PROCESO TO-BE ───────────────────────────────────────────────
_, p = find_para(doc, "3.3 Process Cards TO-BE", "Heading 3")
if p:
    set_para_text(p, "3.3 Fichas de Proceso TO-BE")

_, p = find_para(doc, "Card 1: xxx (TO-BE)")
if p:
    set_para_text(p, "Ficha 1: Adquisicion de Imagenes y Clasificacion IA de Anomalias (TO-BE)")
idx_tc1 = _ if _ is not None else 0

tc1_items = []
for i, para in enumerate(doc.paragraphs):
    if i <= idx_tc1:
        continue
    if para.text.strip() == "Card 2: xxx (TO-BE)":
        break
    if para.text.strip() in ("Input: xxx", "- xxx", "1. xxx", "Sistemi:xxx", "Sistemi: xxx"):
        tc1_items.append(para)

if len(tc1_items) >= 1:
    set_para_text(tc1_items[0],
        "Entrada: Imagenes IR drone (PyroguardIAn); Imagenes RGB drone (SolarEyes); Datos SCADA "
        "produccion; Metadata GPS; Base de Conocimiento tecnica fotovoltaico")
if len(tc1_items) >= 2:
    set_para_lines(tc1_items[1], [
        "Actividades:",
        "1. ETL Pipeline ingiere automaticamente imagenes y datos contextuales post-vuelo",
        "2. PyroguardIAn analiza imagenes IR y clasifica anomalias termicas por severidad",
        "3. SolarEyes analiza imagenes RGB y detecta defectos visuales en cada panel",
        "4. Orquestador correlaciona los resultados con SCADA y Base de Conocimiento",
        "5. HITL: Tecnico de oficina revisa y aprueba el informe de anomalias en el Dashboard"
    ])
if len(tc1_items) >= 3:
    set_para_text(tc1_items[2],
        "Salida: Informe de anomalias validado con mapa geoespacial, clasificacion IA, impacto "
        "en produccion y recomendaciones de intervencion")
if len(tc1_items) >= 4:
    set_para_text(tc1_items[3],
        "Sistemas: PyroguardIAn, SolarEyes, ETL Pipeline, Image Repository, SCADA, Base de "
        "Conocimiento, Dashboard HITL")

_, p = find_para(doc, "Card 2: xxx (TO-BE)")
if p:
    set_para_text(p, "Ficha 2: Gestion de Alertas Criticas y Planificacion Optimizada de Intervenciones (TO-BE)")
idx_tc2 = _ if _ is not None else 0

tc2_items = []
for i, para in enumerate(doc.paragraphs):
    if i <= idx_tc2:
        continue
    if para.text.strip() == "Card 3: xxx (TO-BE)":
        break
    if para.text.strip() in ("Input:xxx", "Input: xxx", "- xxx", "1. xxx", "Sistemi:xxx", "Sistemi: xxx"):
        tc2_items.append(para)

if len(tc2_items) >= 1:
    set_para_text(tc2_items[0],
        "Entrada: Lista de intervenciones validadas; Previsiones meteorologicas/irradiancia; "
        "Disponibilidad de tecnicos, proveedores y repuestos; Impacto estimado en produccion "
        "por anomalia")
if len(tc2_items) >= 2:
    set_para_lines(tc2_items[1], [
        "Actividades:",
        "1. Para anomalias criticas (incendio): notificacion inmediata push a HSE y Responsable",
        "2. Orquestador recopila automaticamente restricciones (meteorologia, repuestos, disponibilidad)",
        "3. IA genera plan de intervenciones optimizado multi-restriccion",
        "4. HITL: Responsable de instalacion revisa y aprueba el plan",
        "5. Post-aprobacion: OT creadas automaticamente en CMMS y notificaciones a tecnicos"
    ])
if len(tc2_items) >= 3:
    set_para_text(tc2_items[2],
        "Salida: Plan de intervenciones aprobado; OT creadas en el CMMS; tecnicos/proveedores notificados")
if len(tc2_items) >= 4:
    set_para_text(tc2_items[3],
        "Sistemas: Agente Orquestador, Modulo HITL, Dashboard HITL, CMMS, API Meteorologica, ERP Repuestos")

_, p = find_para(doc, "Card 3: xxx (TO-BE)")
if p:
    set_para_text(p, "Ficha 3: Ejecucion de Intervenciones y Feedback para Mejora Continua IA (TO-BE)")
idx_tc3 = _ if _ is not None else 0

tc3_items = []
for i, para in enumerate(doc.paragraphs):
    if i <= idx_tc3:
        continue
    if para.style.name in ("Heading 3", "Heading 2"):
        break
    if para.text.strip() in ("Input: xxx", "- xxx", "1. xxx", "Sistemi:xxx", "Sistemi: xxx"):
        tc3_items.append(para)

if len(tc3_items) >= 1:
    set_para_text(tc3_items[0],
        "Entrada: OT estructuradas con datos IA (geolocalizacion, imagen anomalia, instrucciones); "
        "actas digitales post-intervencion")
if len(tc3_items) >= 2:
    set_para_lines(tc3_items[1], [
        "Actividades:",
        "1. Tecnicos reciben OT enriquecidas con datos IA e instrucciones de la BC en App Movil",
        "2. Ejecucion de intervencion fisica en campo",
        "3. Tecnico cumplimenta acta digital estructurada via App Movil (resultado, componentes, notas)",
        "4. IA verifica eficacia de intervencion con imagenes post-intervencion (siguiente sesion drone)",
        "5. Feedback loop: actas alimentan reentrenamiento mensual de modelos IA"
    ])
if len(tc3_items) >= 3:
    set_para_text(tc3_items[2],
        "Salida: Acta digital estructurada; verificacion de eficacia IA; KPI actualizados en "
        "Dashboard; modelos IA mejorados")
if len(tc3_items) >= 4:
    set_para_text(tc3_items[3],
        "Sistemas: CMMS, App Movil, Motores IA (post-verificacion), Pipeline de Feedback, "
        "Base de Conocimiento, Dashboard HITL, SCADA")

# ── 3.4 QUE NO HACE LA IA ─────────────────────────────────────────────────────
_, p = find_para(doc, "3.4 Cosa NON fa l", "Heading 3")
if p:
    set_para_text(p, "3.4 Que NO hace la IA")

_, p_cosa = find_para(doc, "xxx", "Normal", _)
if p_cosa:
    set_para_lines(p_cosa, [
        "No ejecuta fisicamente las intervenciones de mantenimiento en los paneles "
        "(permanece responsabilidad de tecnicos/proveedores)",
        "No activa autonomamente protocolos de emergencia contra incendios sin aprobacion "
        "explicita del Responsable HSE (HITL obligatorio para eventos criticos)",
        "No pilota autonomamente los drones de inspeccion ni gestiona las autorizaciones de vuelo ENAC",
        "No accede a los sistemas de control SCADA/DCS para modificar parametros operativos "
        "ni detener inversores",
        "No emite ordenes de trabajo en el CMMS sin validacion humana del Responsable de instalacion",
        "No reemplaza el diagnostico tecnico especializado del termografista certificado para "
        "casos complejos o ambiguos",
        "No se comunica autonomamente con proveedores o empresas contratistas sin aprobacion HITL",
        "No certifica la conformidad de las instalaciones con normativas HSE o autorizaciones "
        "(CEI 82-25, IEC 62446-3)"
    ])

# ── 4. DELTA AS-IS VS TO-BE ───────────────────────────────────────────────────
_, p = find_para(doc, "4. Delta AS-IS vs TO-BE", "Heading 2")
if p:
    set_para_text(p, "4. Delta AS-IS vs TO-BE")

_, p = find_para(doc, "Cosa Cambia", "Heading 3")
if p:
    set_para_text(p, "Que Cambia")

_, p = find_para(doc, "AS-IS:", "Normal")
if p:
    set_para_text(p, "AS-IS:")

_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "- Inspecciones termograficas periodicas (semestrales/anuales) con operadores de drone "
        "y termografistas manuales",
        "- Analisis manual de imagenes con software dedicado (2-5 dias por instalacion); "
        "informes Word/Excel",
        "- Sin correlacion automatica de anomalias termicas con datos de produccion SCADA",
        "- Planificacion manual de intervenciones sin optimizacion formal; gap 4-12 semanas "
        "anomalia a intervencion",
        "- Sin monitoreo continuo; imposible detectar eventos criticos (incendio) en tiempo real"
    ])

_, p = find_para(doc, " TO-BE:", "Normal")
if p:
    set_para_text(p, "TO-BE:")

_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "- PyroguardIAn monitorea continuamente las imagenes termograficas y detecta hotspots y "
        "riesgos de incendio en tiempo real con notificacion inmediata",
        "- SolarEyes analiza automaticamente cada sesion drone y produce mapa geoespacial de "
        "defectos visuales para la instalacion completa",
        "- Orquestador correlaciona anomalias IA con datos SCADA y genera plan de intervenciones "
        "optimizado listo para aprobacion HITL",
        "- Responsable de instalacion aprueba intervenciones y gestiona emergencias via dashboard "
        "dedicado: decision informada en < 2 horas",
        "- Feedback loop continuo: cada intervencion mejora la precision de los modelos IA con el tiempo"
    ])

# ── 4.1 IMPACTOS OPERATIVOS ───────────────────────────────────────────────────
_, p = find_para(doc, "4.1 Impatti Operativi", "Heading 3")
if p:
    set_para_text(p, "4.1 Impactos Operativos")

_, p = find_para(doc, "xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "- Reduccion estimada ~80% del tiempo de analisis de imagenes: de 2-5 dias manuales a "
        "2-4 horas (IA + HITL)",
        "- Deteccion temprana de riesgos de incendio en tiempo real vs. actual imposibilidad de "
        "monitoreo continuo",
        "- Reduccion de coste por MW inspeccionado: eliminacion del coste del termografista "
        "certificado para el analisis basico",
        "- Gap anomalia a intervencion reducido de 4-12 semanas a 3-7 dias "
        "(priorizacion IA + OT automatica)",
        "- Aumento de la frecuencia efectiva de inspeccion: de semestral/anual a continua "
        "(sensores fijos) + post-vuelo (drone)",
        "- Recuperacion de produccion estimada: intervencion oportuna sobre hotspots con DeltaT "
        "> 20 C recupera de media 0.5-2% PR por modulo"
    ])

# ── 4.2 INVARIANTES (NO CAMBIA) ───────────────────────────────────────────────
_, p = find_para(doc, "4.2 Invarianti", "Heading 3")
if p:
    set_para_text(p, "4.2 Invariantes (No Cambia)")

h42_idx = _ if _ is not None else 0
inv_count = 0
for i, para in enumerate(doc.paragraphs):
    if i <= h42_idx:
        continue
    if para.style.name.startswith("Heading"):
        break
    if para.text.strip() == "- xxx":
        if inv_count == 0:
            set_para_text(para,
                "- La decision final de aprobacion de intervenciones y activacion de protocolos "
                "de emergencia permanece a cargo del responsable humano (HITL obligatorio)")
        elif inv_count == 1:
            set_para_text(para,
                "- Los drones fisicos y los sensores de campo siguen siendo los sistemas de "
                "adquisicion primarios; la IA analiza las imagenes pero no las adquiere")
        elif inv_count == 2:
            set_para_text(para,
                "- Los tecnicos ejecutan fisicamente todas las intervenciones de mantenimiento "
                "en los paneles en campo")
        elif inv_count == 3:
            set_para_text(para,
                "- Las normativas aeronauticas (ENAC), contra incendios (CEI 82-25) y tecnicas "
                "(IEC 62446-3) para instalaciones fotovoltaicas permanecen invariadas y vinculantes")
        inv_count += 1

# ── 5. HOJA DE RUTA (DRAFT) ───────────────────────────────────────────────────
_, p = find_para(doc, "5. Roadmap Draft", "Heading 2")
if p:
    set_para_text(p, "5. Hoja de Ruta (Draft)")

_, p = find_para(doc, "Durata Totale Stimata:")
if p:
    set_para_text(p,
        "Duracion Total Estimada: 12 meses (M1-M12; Fase 3 y Fase 4 parcialmente paralelizables)")

# TABLE 10 – Roadmap
t10 = doc.tables[9]
fill_row(t10, 0, ["Fase", "Objetivo", "Entregable", "Duracion Estimada"])
roadmap = [
    ("Fase 1 - Assessment y Dataset Foundation\n(M1-M2)",
     "Recopilacion y catalogacion de imagenes termograficas y RGB historicas; definicion de "
     "estandar metadata drone; mapeo de API SCADA/CMMS; configuracion del Image Repository estructurado",
     "Dataset historico catalogado (>=500 sesiones de vuelo); acuerdos IT para accesos API OT/CMMS; "
     "especificaciones metadata drone definidas",
     "2 meses"),
    ("Fase 2 - Etiquetado de Dataset y Entrenamiento de Modelos CV\n(M2-M5)",
     "Etiquetado del dataset con termografistas certificados (anomalias termicas) y tecnicos O&M "
     "(defectos visuales); entrenamiento de modelos PyroguardIAn (IR) y SolarEyes (RGB); "
     "validacion en test set",
     "Modelos CV con precision >= 85% en test set interno; falsos positivos < 15%; aprobacion "
     "cualitativa de termografistas",
     "3 meses"),
    ("Fase 3 - Desarrollo de Motor y ETL Pipeline\n(M4-M7)",
     "Implementacion del ETL Pipeline; integracion API SCADA/Image Repository; desarrollo del "
     "Agente Orquestador; Base de Conocimiento fotovoltaico; primera version del Dashboard HITL",
     "ETL funcional en entorno dev; Motor IA integrado con Image Repository y SCADA; "
     "Dashboard HITL v0.1 navegable",
     "3 meses"),
    ("Fase 4 - Integracion CMMS y Flujo HITL\n(M6-M9)",
     "Integracion con CMMS para creacion automatica de OT; desarrollo de App Movil para tecnicos; "
     "prueba del flujo HITL completo (anomalia > OT); pruebas con supervisores y tecnicos piloto",
     "Flujo HITL extremo a extremo funcional en entorno pre-produccion; App Movil probada por "
     "3-5 tecnicos piloto; OT creadas automaticamente validadas",
     "3 meses"),
    ("Fase 5 - Piloto Operativo en Instalacion Muestra\n(M8-M10)",
     "Piloto operativo en 1-2 instalaciones EGP seleccionadas; seguimiento de KPI (tasa de "
     "deteccion, falsos positivos, MTTR, produccion recuperada); go/no-go para escalado",
     "Tasa de deteccion de anomalias >= 90%; falsos positivos < 10%; MTTR reducido >= 30%; "
     "adopcion HITL > 80%; informe go/no-go de escalado",
     "3 meses"),
    ("Fase 6 - Escalado al Parque Solar EGP\n(M10-M12+)",
     "Extension gradual a todas las instalaciones solares EGP en alcance; formacion del personal "
     "O&M y HSE; puesta en regimen del reentrenamiento continuo y seguimiento de KPI; plan de "
     "evolucion futura (sensores fijos)",
     "Sistema en produccion en >= 80% de instalaciones objetivo; plan de mejora continua definido; "
     "hoja de ruta de sensores fijos aprobada",
     "3 meses+"),
]
for i, row_data in enumerate(roadmap):
    if i + 1 < len(t10.rows):
        fill_row(t10, i + 1, list(row_data))

# ── 6. KPI ─────────────────────────────────────────────────────────────────────
_, p = find_para(doc, "6. KPI", "Heading 2")
if p:
    set_para_text(p, "6. KPI")

_, p = find_para(doc, "KPI Quantitativi", "Normal")
if p:
    set_para_text(p, "KPI Cuantitativos")

_, p = find_para(doc, "KPI Qualitativi", "Normal")
if p:
    set_para_text(p, "KPI Cualitativos")

# TABLE 11 – KPI Cuantitativos
# NOTE: PREGUNTA ABIERTA — kpi_cuantitativos: baselines AS-IS a confirmar con datos EGP
t11 = doc.tables[10]
fill_row(t11, 0, ["KPI", "AS-IS (Situacion Actual)", "Objetivo TO-BE", "Metodo de Medicion"])
kpi_quant = [
    ("Tasa de deteccion de anomalias termicas criticas (hotspot DeltaT > 20 C)",
     "~30-50% (inspecciones periodicas manuales; muchos eventos no detectados entre campanas) "
     "[A CONFIRMAR CON DATOS EGP] - \u26a0\ufe0f NOTA: confirmar baseline con IT Owner EGP",
     "> 90% (monitoreo continuo IA)",
     "Comparacion anomalias detectadas por IA vs. anomalias confirmadas post-intervencion (recall)"),
    ("Tasa de falsos positivos de clasificacion IA",
     "N/A (baseline: analisis 100% manual)",
     "< 10% de tasa de falsos positivos",
     "% anomalias IA que no fueron confirmadas por tecnico HITL o post-intervencion"),
    ("Tiempo de analisis de imagenes por sesion drone (horas FTE)",
     "~2-5 dias FTE/instalacion (analisis manual de termografista certificado)",
     "< 4 horas (motor IA + revision HITL)",
     "Timestamp ingestion imagenes > aprobacion informe HITL en Dashboard"),
    ("MTTR Solar - gap anomalia > intervencion completada (dias)",
     "~28-84 dias (4-12 semanas: campana > analisis > planificacion > intervencion) "
     "[A CONFIRMAR] - \u26a0\ufe0f NOTA: confirmar MTTR real con datos historicos O&M",
     "< 7 dias para anomalias High/Critical",
     "Tracking timestamp: deteccion anomalia IA > cierre OT CMMS"),
    ("Coste por MW inspeccionado (EUR/MW/anno)",
     "~500-2.000 EUR/MW/anno (campana drone + analisis termografista) "
     "[A CUANTIFICAR CON DATOS EGP] - \u26a0\ufe0f NOTA: confirmar coste base con EGP",
     "Reduccion >= 40% (automatizacion analisis, reduccion campanas manuales)",
     "Coste total inspecciones anual / MW en produccion"),
    ("Produccion recuperada (MWh/anno por anomalias resueltas a tiempo)",
     "Baseline a calcular en historico EGP (produccion perdida por hotspots no detectados) "
     "[A DEFINIR] - \u26a0\ufe0f NOTA: calcular baseline con datos SCADA historicos",
     "+0.5-1.5% PR medio instalaciones monitoreadas",
     "Comparacion PR (Performance Ratio) antes/despues de intervencion en anomalias resueltas; "
     "dato SCADA"),
    ("Numero de eventos incendio/humo detectados en tiempo real vs. detectados a posteriori",
     "~0% deteccion en tiempo real (sin monitoreo continuo actual)",
     "> 95% deteccion en menos de 15 minutos desde aparicion del evento",
     "Log eventos criticos: timestamp aparicion (de sensores/SCADA) vs. timestamp alerta IA"),
    ("Cobertura de inspeccion termografica instalaciones/anno (%)",
     "~50-70% instalaciones cubiertas anualmente (restricciones de disponibilidad operadores/drones)",
     "> 95% instalaciones cubiertas por sesion anual (automatizacion analisis desbloquea escalabilidad)",
     "Ratio instalaciones analizadas / total instalaciones en alcance"),
]
for i, row_data in enumerate(kpi_quant):
    if i + 1 < len(t11.rows):
        fill_row(t11, i + 1, list(row_data))

# TABLE 12 – KPI Cualitativos
t12 = doc.tables[11]
fill_row(t12, 0, ["KPI", "Objetivo TO-BE", "Metodo de Medicion"])
kpi_qual = [
    ("Calidad de clasificacion IA (juicio de termografistas certificados)",
     "Target: >= 85% de las clasificaciones IA consideradas correctas o aceptables por "
     "termografistas en revision periodica; requiere benchmark semestral",
     "Revision semestral a muestra de clasificaciones IA por parte de 2-3 termografistas "
     "certificados EGP"),
    ("Adopcion del Dashboard HITL por los Responsables de instalacion",
     "Target: > 90% de los Responsables de instalacion utiliza regularmente el Dashboard HITL "
     "para revision de anomalias en los 6 meses tras el go-live",
     "Log de accesos al Dashboard HITL; encuesta semestral de adopcion"),
    ("Calidad de los planes de intervencion IA (% aprobados sin modificaciones sustanciales)",
     "Target: > 70% de los planes de intervencion generados por IA aprobados por el Responsable "
     "de instalacion sin modificaciones sustanciales en 12 meses",
     "Tracking HITL: planes aprobados tal cual / total planes; analisis de motivos de modificacion"),
    ("Completitud de actas post-intervencion digitales",
     "Target: > 95% de las actas cumplimentadas en formato digital estructurado via App Movil "
     "en las 4 horas siguientes al fin de la intervencion",
     "Tracking envio actas digitales / total OT cerradas en el CMMS"),
    ("Precision del retrieval de la Base de Conocimiento (% consultas resueltas sin escalada humana)",
     "Target: > 80% de las consultas tecnicas de tecnicos (instrucciones de intervencion, "
     "fichas tecnicas) resueltas autonomamente por la BC sin escalada",
     "Tracking consultas BC: resueltas autonomamente / total consultas tecnicos en campo"),
    ("Satisfaccion de usuarios del sistema IA (NPS interno tecnicos y responsables)",
     "Target: NPS > 35 en los 12 meses tras el go-live; enfoque en reduccion de carga manual "
     "percibida por tecnicos de oficina y termografistas",
     "Encuesta NPS trimestral a todos los usuarios del sistema (tecnicos, responsables, HSE)"),
    ("Actualizacion continua de modelos IA (estabilidad de precision en el tiempo)",
     "Target: sin degradacion de precision > 5% en test set rolling en los 12 meses post "
     "go-live; reentrenamiento mensual con nuevos datos validados",
     "Monitoreo automatico de precision con alertas en caso de degradacion; MLflow tracking"),
    ("Trazabilidad de decisiones IA para auditoria HSE (% eventos con explicacion)",
     "Target: 100% de las alertas criticas (incendio/humo) documentadas con: imagen original, "
     "output IA, timestamp, accion HITL tomada; conformidad para auditoria HSE",
     "Log de auditoria automatico Dashboard HITL + CMMS para cada evento severity Critical"),
]
for i, row_data in enumerate(kpi_qual):
    if i + 1 < len(t12.rows):
        fill_row(t12, i + 1, list(row_data))

# ── GUARDAR DOCUMENTO ─────────────────────────────────────────────────────────
os.makedirs(OUTPUT_DIR, exist_ok=True)
doc.save(OUTPUT)
print(f"Blueprint en espanol guardado en:\n{OUTPUT}")
print()
print("=" * 70)
print("PREGUNTAS ABIERTAS — Secciones que requieren datos del equipo EGP TGX:")
print("=" * 70)
print("  1. roles            — Nombres y unidades organizativas especificas de")
print("                        los roles del proyecto (Business Owner, IT Owner,")
print("                        Data Owner, Product Owner EGP TGX)")
print("  2. sistemas_asis    — Nombres exactos de los sistemas SCADA/DCS, CMMS,")
print("                        ERP en uso en las instalaciones solares EGP")
print("                        (senalados como [A DEFINIR])")
print("  3. kpi_cuantitativos — Baselines AS-IS reales del historico operativo EGP:")
print("                        n. eventos incendio/anno, MTTR real, coste/MW")
print("                        inspeccion, % cobertura actual")
print("                        (senalados como [A CONFIRMAR])")
print("  4. data_mapping     — Detalle formatos API de los sistemas OT/SCADA y")
print("                        CMMS (protocolo, autenticacion, rate limit)")
print("                        (senalados como [A DEFINIR])")
print("  5. arquitectura     — Eleccion de stack tecnologico IA final a validar")
print("                        con IT Owner (cloud provider, MLOps platform,")
print("                        proveedor de drones)")
