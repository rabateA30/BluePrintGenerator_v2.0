"""
Blueprint Generator – IT O&M Gestione Penali (EGP TGX) — versione IT concisa
Progetto:  IT - O&M - Gestione Penali — EGP TGX Italia
Sorgente:  BP al BSN per chiarimenti (SharePoint Grids e EGP TGX)
Template:  Blueprint_1.4_vuoto.docx
Output:    Blueprint_Penali_OM_IT.docx  (root repository)
Lingua:    Italiano (non verbosa)

Struttura template Blueprint_1.4_vuoto.docx (verificata):
  P0:  Heading1 'Blueprint - XXX - XXX'
  P4:  Normal   '##TITOLO##'
  P5:  Normal   '##contenuto##'
  P6:  Heading1 'Processo 1: XXX'
  P9:  Normal   'xxxx'  (Scopo)
  P11-13: Normal '• xxx' (Finalita')
  P16-18: Normal '- xxx' (IN SCOPE)
  P20-21: Normal '- xxx' (OUT SCOPE)
  P24: ListPara 'xxx'   (Normativi)
  P26: Normal   'xxx'   (Tecnici)
  P28: Normal   '\nxxx' (Organizzativi)
  P40,42,44: Normal 'Macroattivita': xxx'
  P47-61: AS-IS Process Cards
  P65-69: Pain points '1. xxx' ... '5. xxx'
  P77: 'Pattern tassonomico...: xxxx'
  P80: 'Sotto-processo A: xxxx'
  P81: 'Sotto-processo B: xxx'
  P83: 'Sotto-processo C: XXXX (TO-BE)'
  P87: Heading3 '3.3 Process Cards TO-BE'
  P88-102: TO-BE Process Cards
  P104: Heading3 "3.4 Cosa NON fa l'AI"
  P105: Normal  '❌ xxx\n❌ xxx...'
  P109: 'AS-IS:'
  P110-113: ListPara 'xxx' (AS-IS bullets x4)
  P114: Normal ' TO-BE:' (spazio iniziale)
  P115-119: ListPara 'xxx' (TO-BE bullets x5)
  P122: '• xxx\n• xxx\n• xxx' (Impatti Operativi)
  P125-126: Normal '- xxx' (Invarianti)
  Tables 0-12

Esecuzione:
    python generate_blueprint_penali_IT.py [TEMPLATE_PATH] [OUTPUT_PATH]
"""

import copy
import os
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── path resolution ────────────────────────────────────────────────────────────

def _resolve_path(cli_index, env_name, default_value):
    if len(sys.argv) > cli_index and sys.argv[cli_index]:
        return sys.argv[cli_index]
    return os.environ.get(env_name, default_value)


TEMPLATE = _resolve_path(1, "BLUEPRINT_TEMPLATE", "Blueprint_1.4_vuoto.docx")
OUTPUT   = _resolve_path(2, "BLUEPRINT_OUTPUT",   "Blueprint_Penali_OM_IT.docx")


# ── helpers ───────────────────────────────────────────────────────────────────

def find_para(doc, fragment, style=None, start=0):
    """Restituisce (indice, paragrafo) del primo paragrafo che contiene fragment."""
    if start is None:
        start = 0
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if fragment in p.text:
            if style is None or p.style.name == style:
                return i, p
    return None, None


def _get_rpr(para):
    for run in para.runs:
        rpr = run._r.find(qn("w:rPr"))
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None


def set_para_text(para, new_text):
    rpr = _get_rpr(para)
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    r_elem = OxmlElement("w:r")
    if rpr is not None:
        r_elem.append(rpr)
    t_elem = OxmlElement("w:t")
    t_elem.text = new_text
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_elem.append(t_elem)
    para._p.append(r_elem)


def set_para_lines(para, lines):
    rpr = _get_rpr(para)
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    for idx, line in enumerate(lines):
        r_elem = OxmlElement("w:r")
        if rpr is not None:
            r_elem.append(copy.deepcopy(rpr))
        t_elem = OxmlElement("w:t")
        t_elem.text = line
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r_elem.append(t_elem)
        para._p.append(r_elem)
        if idx < len(lines) - 1:
            r_br = OxmlElement("w:r")
            if rpr is not None:
                r_br.append(copy.deepcopy(rpr))
            br = OxmlElement("w:br")
            r_br.append(br)
            para._p.append(r_br)


def fill_cell(cell, text):
    para = cell.paragraphs[0]
    set_para_text(para, text)


def fill_row(table, row_idx, values):
    row = table.rows[row_idx]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            fill_cell(row.cells[ci], val)


# ── document ──────────────────────────────────────────────────────────────────

doc = Document(TEMPLATE)


# ── TITOLO PRINCIPALE ─────────────────────────────────────────────────────────
_, p = find_para(doc, "Blueprint \u2013 XXX", "Heading 1")
if p:
    set_para_text(p,
        "Blueprint \u2013 IT O&M Gestione Penali \u2013 "
        "Automazione AI del Ciclo Penali Contrattuali EGP TGX Italia")


# ── SOMMARIO ESECUTIVO ────────────────────────────────────────────────────────
_, p = find_para(doc, "##TITOLO##")
if p:
    set_para_text(p,
        "IT O&M Gestione Penali: Automazione AI del Ciclo Penali Contrattuali EGP TGX Italia")

_, p = find_para(doc, "##contenuto##")
if p:
    set_para_lines(p, [
        "Il processo gestisce il ciclo di vita delle penali contrattuali degli impianti O&M "
        "EGP TGX in Italia: ricezione da committenti/TSO/DSO, verifica del calcolo, contestazione "
        "documentata e monitoraggio proattivo dei KPI.",
        "",
        "Proposta: architettura AI multi-agent (Orchestrator, Penalty Classifier, Calculation "
        "Verifier, Dispute Generator, Penalty Monitor, Knowledge Agent) per automatizzare il ciclo "
        "ricezione-verifica-contestazione, con HITL obbligatorio per ogni decisione.",
        "",
        "Proprietario processo: EGP/TGX \u2013 O&M Italia",
    ])


# ── PROCESSO 1 \u2013 TITOLO ───────────────────────────────────────────────────────
_, p = find_para(doc, "Processo 1: XXX", "Heading 1")
if p:
    set_para_text(p,
        "Processo 1: Gestione Automatizzata delle Penali Contrattuali EGP TGX Italia")


# ── 1.1 SCOPO ─────────────────────────────────────────────────────────────────
h11_idx, _ = find_para(doc, "1.1 Scopo", "Heading 3")
_, p_scopo = find_para(doc, "xxxx", "Normal", h11_idx)
if p_scopo:
    set_para_lines(p_scopo, [
        "Il processo gestisce il ciclo di vita delle penali contrattuali degli impianti O&M "
        "EGP TGX in Italia: ricezione delle notifiche da committenti, TSO/DSO ed enti regolatori, "
        "verifica della correttezza del calcolo, generazione delle contestazioni documentate, "
        "chiusura del contenzioso e monitoraggio dei KPI.",
        "",
        "Obiettivo: ridurre l'esposizione economica tramite verifica sistematica e tempestiva, "
        "contestazione automatizzata delle penali errate e monitoraggio proattivo per prevenire "
        "future penali.",
        "",
        "Sistemi in scope: SAP/ERP, SCADA/PI, piattaforma documentale contratti, email, "
        "Excel/SharePoint.",
    ])


# ── 1.2 FINALITA' ─────────────────────────────────────────────────────────────
h12_idx, _ = find_para(doc, "1.2 Finalit", "Heading 3")
finalita_texts = [
    "- Ridurre il valore complessivo delle penali accettate tramite verifica automatizzata "
    "della correttezza dei calcoli applicati da committenti ed enti.",
    "- Accelerare i tempi di risposta alle notifiche generando automaticamente le bozze di "
    "contestazione documentate, riducendo il rischio di decadenza dei termini contrattuali.",
    "- Monitorare proattivamente i KPI contrattuali (availability, performance ratio, "
    "curtailment) per anticipare e prevenire nuove penali tramite alert predittivi.",
]
count = 0
for i, para in enumerate(doc.paragraphs):
    if i <= h12_idx:
        continue
    if para.style.name in ("Heading 3", "Heading 2"):
        break
    if "\u2022 xxx" in para.text and count < len(finalita_texts):
        set_para_text(para, finalita_texts[count])
        count += 1


# ── 1.3 PERIMETRO ─────────────────────────────────────────────────────────────
h13_idx, _ = find_para(doc, "1.3 Perimetro", "Heading 3")
if h13_idx is None:
    raise ValueError("Sezione '1.3 Perimetro' non trovata nel template.")
in_scope = [
    "- Ricezione, classificazione e registrazione delle notifiche di penale "
    "da committenti, TSO/DSO ed enti regolatori per gli impianti EGP TGX Italia",
    "- Verifica automatizzata della correttezza del calcolo tramite confronto "
    "con dati SCADA/PI e clausole contrattuali",
    "- Generazione automatica delle bozze di contestazione e invio previa approvazione HITL",
]
out_scope = [
    "- Esecuzione degli interventi tecnici correttivi sugli impianti (team O&M di campo)",
    "- Gestione penali fiscali/tributarie e controversie legali in fase giudiziale",
]
in_c, out_c, mode = 0, 0, None
for i, para in enumerate(doc.paragraphs):
    if i <= h13_idx:
        continue
    if para.style.name in ("Heading 3", "Heading 2"):
        break
    t = para.text.strip()
    if t == "IN SCOPE:":
        mode = "in"
    elif t == "OUT OF SCOPE:":
        mode = "out"
    elif t == "- xxx":
        if mode == "in" and in_c < len(in_scope):
            set_para_text(para, in_scope[in_c])
            in_c += 1
        elif mode == "out" and out_c < len(out_scope):
            set_para_text(para, out_scope[out_c])
            out_c += 1


# ── 1.4 VINCOLI CHIAVE ────────────────────────────────────────────────────────
h14_idx, _ = find_para(doc, "1.4 Vincoli chiave", "Heading 3")
_, p_norm = find_para(doc, "xxx", "List Paragraph", h14_idx)
if p_norm:
    set_para_text(p_norm,
        "Conformita' ai termini contrattuali O&M; rispetto delle scadenze per la contestazione "
        "(15-30 giorni dalla notifica); conformita' a normativa vigente (ARERA, GME, GSE/PPA).")
_, p_tec = find_para(doc, "xxx", "Normal", h14_idx)
if p_tec:
    set_para_text(p_tec,
        "Disponibilita' di API verso SAP/ERP, SCADA/PI e piattaforma documentale; "
        "vincoli di accesso ai dati operativi degli impianti (granularita', latenza, NDA).")
_, p_org = find_para(doc, "\nxxx", "Normal", h14_idx)
if p_org:
    set_para_text(p_org,
        "Modello di governance per l'approvazione delle contestazioni (HITL obbligatorio); "
        "change management per l'adozione del sistema da parte del team penali EGP TGX.")


# ── TABLE 0 \u2013 STAKEHOLDER / ROLES ────────────────────────────────────────────
t0 = doc.tables[0]
stakeholders = [
    ("Business Owner",
     "EGP TGX \u2013 O&M Italia",
     "Approva il Blueprint e le decisioni di business; responsabile del valore economico "
     "(riduzione penali pagate)"),
    ("Data Owner",
     "Responsabile dati O&M EGP TGX Italia",
     "Garantisce qualita' e governance dei dati operativi (SCADA/PI), contrattuali e "
     "del registro storico penali"),
    ("IT Owner",
     "GICT / Digital Infrastructure EGP",
     "Supervisiona architettura tecnica, integrazioni SAP/SCADA/PI e conformita' "
     "agli standard IT Enel"),
    ("Product Owner",
     "AISA Factory EGP TGX",
     "Gestisce il backlog e coordina le iterazioni di sviluppo; interfaccia tra "
     "team AI e Business Owner"),
    ("Referente O&M Penali",
     "Team Penali EGP TGX Italia",
     "Utente principale: revisiona le analisi AI, approva le contestazioni tramite "
     "Dashboard HITL, gestisce il contenzioso"),
    ("Responsabile Legale / Contratti",
     "Legal & Compliance EGP TGX",
     "Revisiona le contestazioni di importo rilevante; garantisce conformita' alle "
     "clausole contrattuali e alla normativa"),
    ("BSN (Business Solution Network)",
     "AI Scale Up Accelerator \u2013 Grids & EGP TGX",
     "Supporto metodologico per la definizione, validazione e governance del Blueprint"),
]
for i, row_data in enumerate(stakeholders):
    if i + 1 < len(t0.rows):
        fill_row(t0, i + 1, list(row_data))


# ── TABLE 1 \u2013 SISTEMI COINVOLTI AS-IS ────────────────────────────────────────
t1 = doc.tables[1]
systems = [
    ("SAP / ERP",
     "Gestione contratti O&M e contabilizzazione delle penali (modulo PM/FI)",
     "ERP / Contract Management"),
    ("SCADA / Sistema PI",
     "Monitoraggio dati di produzione e disponibilita' impianti per il periodo di penale",
     "SCADA / Data Historian"),
    ("Piattaforma Documentale Contratti",
     "Archiviazione contratti O&M e clausole penali (SharePoint / Documentum)",
     "Document Management"),
    ("Sistema di Ticketing",
     "Tracciamento contestazioni e stato contenzioso (ServiceNow / sistema interno)",
     "Ticketing / Workflow"),
    ("Excel / Fogli di Calcolo",
     "Ricalcolo manuale delle penali e confronto con i dati di produzione",
     "Spreadsheet (manuale)"),
    ("Email",
     "Ricezione notifiche penale da committenti/TSO e invio contestazioni",
     "Comunicazione"),
]
for i, (s, r, tp) in enumerate(systems):
    if i + 1 < len(t1.rows):
        fill_row(t1, i + 1, [s, r, tp])


# ── AS-IS SOTTO-PROCESSI LABEL ────────────────────────────────────────────────
_, pa = find_para(doc, "Sotto-processo A: xxxx", "Normal")
if pa:
    set_para_text(pa,
        "Sotto-processo A: Ricezione, Registrazione e Classificazione delle Notifiche di Penale")
_, pb = find_para(doc, "Sotto-processo B: xxx", "Normal")
if pb:
    set_para_text(pb,
        "Sotto-processo B: Verifica Correttezza Calcolo e Preparazione Contestazione")


# ── TABLE 2 \u2013 AS-IS Sotto-processo A ─────────────────────────────────────────
t2 = doc.tables[2]
rows_a = [
    ("A1",
     "Ricezione della notifica di penale via email dal committente/TSO/DSO; "
     "il referente O&M prende in carico e valuta l'entita'",
     "Referente O&M",
     "Email notifica penale",
     "Notifica ricevuta e presa in carico",
     "Email"),
    ("A2",
     "Registrazione manuale della notifica nel registro penali (Excel o ticketing)",
     "Referente O&M",
     "Email notifica penale",
     "Penale registrata nel registro",
     "Excel / Ticketing"),
    ("A3",
     "Classificazione manuale per tipologia: availability, performance ratio, "
     "curtailment, SLA fornitore",
     "Referente O&M",
     "Notifica penale registrata",
     "Penale classificata per tipologia e impianto",
     "Email / Excel"),
    ("A4",
     "Identificazione del contratto di riferimento e delle clausole penali applicabili",
     "Referente O&M / Legale",
     "Notifica classificata",
     "Contratto e clausole penali identificati",
     "Piattaforma documentale / SharePoint"),
]
for i, row_data in enumerate(rows_a):
    if i + 1 < len(t2.rows):
        fill_row(t2, i + 1, list(row_data))


# ── TABLE 3 \u2013 AS-IS Sotto-processo B ─────────────────────────────────────────
t3 = doc.tables[3]
rows_b = [
    ("B1",
     "Estrazione manuale dei dati di produzione/disponibilita' dal SCADA/PI "
     "per il periodo di riferimento della penale",
     "Referente O&M / IT",
     "Periodo di riferimento penale",
     "Dataset produzione/disponibilita'",
     "SCADA / Sistema PI"),
    ("B2",
     "Ricalcolo della penale su foglio Excel applicando la formula contrattuale",
     "Referente O&M",
     "Dati produzione + clausole contrattuali",
     "Calcolo penale rielaborato",
     "Excel"),
    ("B3",
     "Confronto tra penale notificata dal committente e valore ricalcolato; "
     "identificazione del delta",
     "Referente O&M",
     "Penale committente + ricalcolo O&M",
     "Delta identificato (congruente / discrepanza)",
     "Excel"),
    ("B4",
     "Redazione della lettera di contestazione con allegati tecnici; "
     "approvazione interna prima dell'invio",
     "Referente O&M / Legale",
     "Ricalcolo + documentazione tecnica",
     "Bozza contestazione con allegati",
     "Word / Email"),
]
for i, row_data in enumerate(rows_b):
    if i + 1 < len(t3.rows):
        fill_row(t3, i + 1, list(row_data))


# ── TABLE 4 \u2013 AS-IS Sotto-processo C ─────────────────────────────────────────
t4 = doc.tables[4]
rows_c = [
    ("C1",
     "Invio contestazione via email al committente/TSO entro il termine contrattuale; "
     "protocollazione e tracciamento stato contenzioso",
     "Referente O&M",
     "Contestazione approvata",
     "Contestazione inviata; stato contenzioso tracciato",
     "Email / Excel / Ticketing"),
    ("C2",
     "Chiusura penale dopo risposta committente (storno SAP/ERP o pagamento); "
     "monitoraggio KPI per prevenire nuove penali",
     "Referente O&M / Finance",
     "Risposta committente + KPI impianto",
     "Penale chiusa in SAP/ERP; alert impianti a rischio",
     "SAP / ERP / SCADA"),
]
for i, row_data in enumerate(rows_c):
    if i + 1 < len(t4.rows):
        fill_row(t4, i + 1, list(row_data))


# ── PROCESS CARDS AS-IS ───────────────────────────────────────────────────────
# P47: 'Card 1: xxx' (Normal)  P48-51: List Paragraph o Normal
# P52: 'Card 2: xxx' (Normal)  P53-56: Normal
# P57: 'Card 3: xxx' (Normal)  P58-61: Normal
asis_cards = [
    {
        "title":     "Card 1: Ricezione, Registrazione e Prima Valutazione",
        "input":     "Email notifica penale, bollettino TSO, comunicazione GME/ARERA, notifica PPA",
        "attivita":  ("1. Referente O&M riceve notifica via email\n"
                      "2. Registrazione nel registro penali (Excel / ticketing)\n"
                      "3. Classificazione per tipologia (availability, PR, curtailment, SLA)\n"
                      "4. Prima valutazione di congruita'"),
        "output":    "Penale registrata e classificata; decisione iter (verificare / contestare / accettare)",
        "painpoint": ("Nessun sistema centralizzato; classificazione soggettiva; "
                      "rischio di perdere notifiche o superare i termini contrattuali"),
    },
    {
        "title":     "Card 2: Verifica Calcolo e Preparazione Contestazione",
        "input":     "Notifica penale registrata, dati SCADA/PI, clausole contrattuali",
        "attivita":  ("1. Estrazione manuale dati produzione/disponibilita' dal SCADA/PI\n"
                      "2. Ricalcolo su Excel applicando la formula contrattuale\n"
                      "3. Confronto penale notificata vs. ricalcolo; identificazione delta\n"
                      "4. Redazione contestazione e approvazione interna"),
        "output":    "Contestazione redatta con allegati tecnici oppure accettazione motivata",
        "painpoint": ("Estrazione SCADA laboriosa; ricalcolo non standardizzato; "
                      "elevato rischio di errore umano; processo accentrato su pochi esperti"),
    },
    {
        "title":     "Card 3: Gestione Contenzioso e Monitoraggio KPI",
        "input":     "Contestazione inviata, risposta committente, registro penali, KPI impianto",
        "attivita":  ("1. Invio contestazione entro i termini contrattuali\n"
                      "2. Tracking manuale dello stato (Excel / email)\n"
                      "3. Gestione risposta committente (storno / rifiuto / negoziazione)\n"
                      "4. Chiusura in SAP/ERP (storno o pagamento)"),
        "output":    "Penale chiusa (stornata o pagata); report periodico management",
        "painpoint": ("Nessun tracking automatico; visibilita' limitata su penali aggregate; "
                      "monitoraggio KPI preventivo assente o manuale"),
    },
]
card_search_start = 0
for ci, cd in enumerate(asis_cards):
    c_idx, p_card = find_para(doc, f"Card {ci + 1}: xxx", "Normal", card_search_start)
    if p_card:
        set_para_text(p_card, cd["title"])
        cur = c_idx
        # Cerca Input, Attivita, Output, Pain point nell'ordine
        for frag_key, search_frags in [
            ("input",     ["Input: xxx", "Input:xxx"]),
            ("attivita",  ["Attivit\u00e0: xxx", "Attivit\u00e0:xxx",
                           "Attivit\u00e0 xxx"]),
            ("output",    ["Output: xxx", "Output:xxx"]),
            ("painpoint", ["Pain point: xxx", "Pain point:xxx"]),
        ]:
            for frag in search_frags:
                fi, fp = find_para(doc, frag, None, cur)
                if fp and fi is not None and fi <= c_idx + 10:
                    set_para_text(fp, cd[frag_key])
                    cur = fi
                    break
        card_search_start = c_idx + 1


# ── MACROATTIVITA' LABELS AS-IS ───────────────────────────────────────────────
macro_labels = [
    "Macroattivita' A: Ricezione, Registrazione e Classificazione Penali",
    "Macroattivita' B: Verifica Calcolo e Preparazione Contestazione",
    "Macroattivita' C: Gestione Contenzioso, Chiusura e Monitoraggio KPI",
]
macro_start = 0
for label in macro_labels:
    mi, mp = find_para(doc, "Macroattivit\u00e0: xxx", "Normal", macro_start)
    if mp:
        set_para_text(mp, label)
        macro_start = mi + 1


# ── PAIN POINT GENERALI ───────────────────────────────────────────────────────
pp_idx, _ = find_para(doc, "Pain Point Generali", "Normal")
pain_points = [
    "1. Effort FTE elevato: estrazione dati SCADA, ricalcolo Excel e redazione contestazioni "
    "richiedono molte ore/penale; knowledge contrattuale distribuita informalmente",
    "2. Tempi di attraversamento lunghi: dal ricevimento alla trasmissione della contestazione; "
    "rischio concreto di superare i termini contrattuali",
    "3. Processo accentrato su pochi referenti esperti; nessuna visibilita' aggregata sul portfolio",
    "4. Nessun monitoraggio preventivo dei KPI contrattuali; penali rilevate solo a posteriori",
    "5. Nessun sistema centralizzato di acquisizione e registrazione delle notifiche di penale",
]
if pp_idx is not None:
    for pp_n, pp_text in enumerate(pain_points, 1):
        pi, pp = find_para(doc, f"{pp_n}. xxx", "Normal", pp_idx)
        if pp:
            set_para_text(pp, pp_text)


# ── PATTERN TASSONOMICO ───────────────────────────────────────────────────────
_, p = find_para(doc, "Pattern tassonomico Catalogo Soluzioni AI", "Normal")
if p:
    set_para_text(p,
        "Pattern tassonomico: Multi-Agent Orchestration + RAG + Document Intelligence "
        "+ Predictive Monitor")


# ── TABLE 5 \u2013 DATA MAPPING ────────────────────────────────────────────────────
t5 = doc.tables[5]
data_mapping = [
    ("Notifiche di penale (tipo, importo, periodo, impianto)",
     "Email committente / TSO / ARERA",
     "Penalty Classifier Agent / Registro Penali",
     "Email / PDF",
     "Parsing strutturato o OCR/LLM in base al formato committente"),
    ("Dati produzione e disponibilita' impianto (periodo penale)",
     "SCADA / Sistema PI",
     "Calculation Verifier Agent",
     "CSV / REST API / PI Web API",
     "Granularita' 15 min / oraria; storico minimo 2 anni"),
    ("Clausole contrattuali e formule calcolo penali",
     "Piattaforma documentale / SharePoint",
     "Knowledge Agent / Calculation Verifier Agent",
     "PDF / Word",
     "Indicizzazione RAG; parsing clausole per tipologia committente"),
    ("Storico penali gestite (tipologia, importo, esito)",
     "Excel / Ticketing / SAP",
     "Knowledge Agent / Dashboard HITL",
     "Excel / CSV / API SAP",
     "Training set per accuratezza classificazione e contestazione"),
    ("KPI contrattuali (availability, PR, curtailment per impianto)",
     "SCADA / Reporting operativo",
     "Penalty Monitor Agent / Dashboard HITL",
     "CSV / REST API",
     "Soglie contrattuali configurate per impianto"),
    ("Risposte committente e dati meteo/curtailment TSO (exclusion clause)",
     "Email committente / Fonte meteo / Log TSO",
     "Registro Penali / Calculation Verifier Agent",
     "Email / PDF / CSV",
     "Dato critico per applicazione delle clausole di esclusione"),
]
for i, row_data in enumerate(data_mapping):
    if i + 1 < len(t5.rows):
        fill_row(t5, i + 1, list(row_data))


# ── TABLE 6 \u2013 ARCHITETTURA FUNZIONALE TO-BE ───────────────────────────────────
t6 = doc.tables[6]
arch_rows = [
    ("Orchestrator Agent",
     "Coordina il flusso end-to-end tra agenti; gestisce l'interfaccia HITL con il "
     "referente O&M; consolida analisi e raccomandazioni",
     "LLM multi-agent orchestration (LangGraph / AutoGen)",
     "No \u2013 specifico processo penali"),
    ("Penalty Classifier Agent",
     "Acquisisce e classifica le notifiche per tipologia, impianto, committente e periodo; "
     "registra e allerta il referente O&M",
     "LLM NLP classification + document parsing (GPT-4o)",
     "No \u2013 specifico processo penali"),
    ("Calculation Verifier Agent",
     "Estrae dati SCADA/PI, applica formula contrattuale, quantifica la discrepanza "
     "e applica exclusion clause",
     "Formula engine + LLM reasoning + API SCADA/PI",
     "No \u2013 specifico processo penali"),
    ("Dispute Generator Agent",
     "Genera la bozza di contestazione con allegati tecnici per revisione HITL",
     "LLM text generation + document assembly (GPT-4o)",
     "No \u2013 specifico processo penali"),
    ("Penalty Monitor Agent",
     "Monitora proattivamente i KPI contrattuali di tutti gli impianti; "
     "genera alert predittivi prima dell'insorgenza di penali",
     "Rule-based + ML trend forecasting + LLM alerts",
     "Si \u2013 riusabile cross-progetto O&M"),
    ("Knowledge Agent",
     "Recupera clausole, exclusion clause, storico contestazioni e best practice "
     "tramite RAG; supporta Calculation Verifier e Dispute Generator",
     "RAG + Vector DB (Azure AI Search) + LLM",
     "Si \u2013 pattern RAG comune"),
    ("ETL Pipeline + Dashboard HITL",
     "ETL: acquisisce e normalizza dati da SCADA/PI, SAP, email, piattaforma documentale. "
     "Dashboard: interfaccia HITL per validazione, approvazione, tracking e KPI portfolio",
     "Azure Data Factory + Power BI Embedded + React + Teams Bot",
     "ETL: Si \u2013 infrastruttura comune; Dashboard: No \u2013 specifica penali"),
]
for i, row_data in enumerate(arch_rows):
    if i + 1 < len(t6.rows):
        fill_row(t6, i + 1, list(row_data))


# ── TO-BE SOTTO-PROCESSI LABEL ────────────────────────────────────────────────
# P83: 'Sotto-processo C: XXXX (TO-BE)'
_, p_c_tobe = find_para(doc, "Sotto-processo C: XXXX (TO-BE)", "Normal")
if p_c_tobe:
    set_para_text(p_c_tobe,
        "Sotto-processo C: Gestione Contenzioso, Monitoraggio Preventivo e Feedback Loop (TO-BE)")


# ── TABLE 7 \u2013 TO-BE Sotto-processo A ──────────────────────────────────────────
t7 = doc.tables[7]
rows_tobe_a = [
    ("A1",
     "ETL Pipeline acquisisce automaticamente le notifiche; Penalty Classifier Agent "
     "le classifica e allerta il referente O&M via Dashboard/Teams",
     "ETL Pipeline + Penalty Classifier Agent",
     "Email/notifica penale (PDF, email, portale committente)",
     "Penale classificata, registrata; alert referente O&M",
     "Email gateway, portale committente/TSO, Registro centralizzato",
     "Document parsing + OCR; LLM NLP classification; alert Teams"),
    ("A2",
     "Calculation Verifier estrae dati SCADA/PI; Knowledge Agent recupera clausole "
     "ed exclusion clause; Calculation Verifier applica formula e quantifica la discrepanza",
     "Calculation Verifier Agent + Knowledge Agent",
     "Periodo di riferimento + tipologia penale + ID contratto",
     "Ricalcolo penale con delta vs. notifica e breakdown esclusioni",
     "SCADA/PI Web API, Knowledge Base contratti",
     "API SCADA automatica; RAG su clausole; LLM per exclusion clause"),
    ("A3",
     "Orchestrator Agent consolida risultati e genera raccomandazione motivata; "
     "HITL \u2013 Referente O&M valida e approva la decisione finale",
     "Orchestrator Agent + Referente O&M (HITL)",
     "Output Calculation Verifier + Knowledge Agent",
     "Decisione validata: contestare (totalmente/parzialmente) o accettare",
     "Dashboard HITL, Teams Bot",
     "LLM synthesis; Human decide e approva; audit log obbligatorio"),
]
for i, row_data in enumerate(rows_tobe_a):
    if i + 1 < len(t7.rows):
        fill_row(t7, i + 1, list(row_data))


# ── TABLE 8 \u2013 TO-BE Sotto-processo B ──────────────────────────────────────────
t8 = doc.tables[8]
rows_tobe_b = [
    ("B1",
     "Dispute Generator riceve la decisione HITL; Knowledge Agent recupera template "
     "e storico contestazioni analoghe per la strategia",
     "Dispute Generator Agent + Knowledge Agent",
     "Decisione HITL: contestare + motivazioni",
     "Template contestazione + best practice da casi analoghi",
     "Knowledge Base (RAG)",
     "Trigger automatico post-HITL; RAG su storico contestazioni"),
    ("B2",
     "Dispute Generator genera bozza con dati tecnici, ricalcolo, exclusion clause, "
     "riferimenti normativi/contrattuali e richiesta di storno",
     "Dispute Generator Agent",
     "Template + dati tecnici + ricalcolo + exclusion clause",
     "Bozza lettera contestazione completa con allegati tecnici",
     "LLM text generation + document assembly",
     "LLM generation con template strutturato; allegati auto-generati"),
    ("B3",
     "HITL \u2013 Referente O&M revisiona e approva la bozza; sistema invia la "
     "contestazione approvata con timestamp e imposta reminder scadenza",
     "Orchestrator Agent + Referente O&M (HITL)",
     "Bozza contestazione AI",
     "Contestazione approvata inviata; reminder scadenza impostato",
     "Dashboard HITL, Email gateway / portale committente",
     "HITL obbligatorio \u2013 nessun invio automatico; audit log completo"),
]
for i, row_data in enumerate(rows_tobe_b):
    if i + 1 < len(t8.rows):
        fill_row(t8, i + 1, list(row_data))


# ── TABLE 9 \u2013 TO-BE Sotto-processo C ──────────────────────────────────────────
t9 = doc.tables[9]
rows_tobe_c = [
    ("C1",
     "Sistema monitora i termini di risposta e notifica il referente O&M; "
     "Penalty Classifier acquisisce la risposta del committente e aggiorna lo stato; "
     "HITL \u2013 Referente O&M gestisce l'esito; SAP/ERP aggiornato solo post-HITL",
     "Orchestrator + Penalty Classifier Agent + Referente O&M (HITL)",
     "Contestazione inviata + risposta committente",
     "Stato contenzioso aggiornato; SAP/ERP allineato post-HITL",
     "Dashboard HITL, Teams Bot, SAP/ERP",
     "Reminder automatici 7/3/1 gg; NLP classification risposta; SAP update post-HITL"),
    ("C2",
     "Knowledge Agent aggiorna la KB con l'esito; Penalty Monitor monitora i KPI "
     "di tutti gli impianti e genera alert predittivi; Orchestrator genera report management",
     "Knowledge Agent + Penalty Monitor Agent + Orchestrator Agent",
     "Esito contenzioso + KPI operativi real-time (SCADA/PI)",
     "KB aggiornata; alert preventivi per impianti a rischio; report management",
     "Knowledge Base, SCADA/PI, Dashboard HITL, Power BI",
     "Document ingestion + embedding update; ML forecasting; LLM insights"),
]
for i, row_data in enumerate(rows_tobe_c):
    if i + 1 < len(t9.rows):
        fill_row(t9, i + 1, list(row_data))


# ── PROCESS CARDS TO-BE ───────────────────────────────────────────────────────
# P87: '3.3 Process Cards TO-BE'
# P88-92: Card 1 | P93-97: Card 2 | P98-102: Card 3
h33_idx, _ = find_para(doc, "3.3 Process Cards TO-BE", "Heading 3")

tobe_cards = [
    {
        "title":    "Card 1: Ricezione Automatizzata, Classificazione e Verifica Calcolo (TO-BE)",
        "input":    "Notifiche penale (email/portale); Dati SCADA/PI; Clausole (KB); Dati meteo/TSO",
        "attivita": ("1. ETL Pipeline acquisisce automaticamente le notifiche\n"
                     "2. Penalty Classifier classifica e allerta il referente O&M\n"
                     "3. Calculation Verifier estrae dati SCADA e applica formula\n"
                     "4. Knowledge Agent recupera exclusion clause e confronta con storico\n"
                     "5. HITL: Referente O&M valida il ricalcolo e decide se contestare"),
        "output":   "Report analitico con raccomandazione motivata (contestare / accettare)",
        "sistemi":  "SCADA/PI, Knowledge Base, Penalty Classifier, Calculation Verifier, Dashboard HITL",
    },
    {
        "title":    "Card 2: Generazione Contestazione e Approvazione HITL (TO-BE)",
        "input":    "Decisione HITL; Ricalcolo; Exclusion clause; Template da KB; Storico analoghi",
        "attivita": ("1. Dispute Generator genera bozza contestazione con allegati tecnici\n"
                     "2. Knowledge Agent recupera best practice da casi analoghi vinti\n"
                     "3. HITL: Referente O&M revisiona e approva la bozza\n"
                     "4. Sistema invia contestazione approvata con timestamp"),
        "output":   "Contestazione inviata; Reminder scadenza; Audit log decisione HITL",
        "sistemi":  "Knowledge Base (RAG), Dispute Generator Agent, Dashboard HITL, Email/Portale",
    },
    {
        "title":    "Card 3: Gestione Contenzioso, Chiusura e Monitoraggio Preventivo (TO-BE)",
        "input":    "Risposta committente; Registro penali; KPI operativi real-time impianti",
        "attivita": ("1. Penalty Classifier acquisisce risposta e aggiorna stato contenzioso\n"
                     "2. HITL: Referente O&M gestisce esito (storno / rifiuto / accordo)\n"
                     "3. SAP/ERP aggiornato automaticamente post-HITL\n"
                     "4. Knowledge Agent aggiorna KB con caso documentato\n"
                     "5. Penalty Monitor monitora KPI e genera alert preventivi"),
        "output":   "KB aggiornata; SAP/ERP allineato; Report management; Alert preventivi",
        "sistemi":  "Knowledge Base, Orchestrator Agent, SAP/ERP, SCADA/PI, Dashboard HITL, Power BI",
    },
]

# Usa offset diretto: i 4 campi sono sempre in posizione c_idx+1, +2, +3, +4
# (Input, Attività, Output, Sistemi) per tutte e tre le card TO-BE.
tobe_field_keys_ordered = ["input", "attivita", "output", "sistemi"]

card_s = h33_idx if h33_idx is not None else 0
for ci, cd in enumerate(tobe_cards):
    c_idx, p_card = find_para(doc, f"Card {ci + 1}: xxx", "Normal", card_s)
    if p_card:
        set_para_text(p_card, cd["title"])
        for offset, fk in enumerate(tobe_field_keys_ordered, 1):
            idx = c_idx + offset
            if idx < len(doc.paragraphs):
                set_para_text(doc.paragraphs[idx], cd[fk])
        card_s = c_idx + 1


# ── 3.4 COSA NON FA L'AI ──────────────────────────────────────────────────────
p34_idx, _ = find_para(doc, "3.4 Cosa NON fa l'AI", "Heading 3")
_, p_cosa = find_para(doc, "\u274c xxx", "Normal", p34_idx)
if p_cosa:
    set_para_lines(p_cosa, [
        "Non invia contestazioni al committente senza approvazione del referente O&M (HITL obbligatorio)",
        "Non aggiorna SAP/ERP senza validazione umana",
        "Non prende decisioni su escalation legale o contenzioso formale in autonomia",
        "Non interpreta clausole ambigue senza revisione del referente O&M o del legale",
        "Non negozia direttamente con committenti o TSO",
        "Non accede a sistemi di pagamento o gestisce flussi finanziari",
        "Non modifica parametri operativi degli impianti (team O&M di campo)",
        "Non divulga dati contrattuali o produttivi a soggetti non autorizzati",
    ])


# ── 4. DELTA AS-IS vs TO-BE ───────────────────────────────────────────────────
delta_idx, _ = find_para(doc, "Cosa Cambia", "Heading 3")

# Bullet AS-IS (List Paragraph 'xxx' x4)
asis_bullets = [
    "Ricezione manuale delle notifiche via email; nessun sistema centralizzato",
    "Verifica calcolo su Excel: non standardizzata, elevato rischio di errore umano",
    "Knowledge contrattuale distribuita informalmente; penali rilevate a posteriori",
    "Tempi di risposta lunghi; rischio di decadenza dei termini contrattuali",
]
asis_b_idx = find_para(doc, "AS-IS:", "Normal", delta_idx)[0]
if asis_b_idx is not None:
    bc = 0
    for i, para in enumerate(doc.paragraphs):
        if i <= asis_b_idx:
            continue
        if " TO-BE:" in para.text or "TO-BE:" in para.text:
            break
        if para.style.name == "List Paragraph" and para.text.strip() == "xxx":
            if bc < len(asis_bullets):
                set_para_text(para, asis_bullets[bc])
                bc += 1

# Aggiorna ' TO-BE:' (spazio iniziale nel template)
_, p_tobe = find_para(doc, " TO-BE:", "Normal", delta_idx)
if p_tobe:
    set_para_text(p_tobe, "TO-BE:")

# Bullet TO-BE (List Paragraph 'xxx' x5)
tobe_bullets = [
    "ETL Pipeline acquisisce automaticamente le notifiche; Penalty Classifier allerta in real time",
    "Calculation Verifier verifica il calcolo tramite dati SCADA e clausole contrattuali",
    "Dispute Generator genera bozze documentate; il referente approva tramite Dashboard HITL",
    "Penalty Monitor monitora proattivamente i KPI e genera alert preventivi per impianti a rischio",
    "Knowledge Base strutturata rende clausole e best practice accessibili a tutti i referenti O&M",
]
tobe_b_idx_val = find_para(doc, " TO-BE:", "Normal", delta_idx)[0]
if tobe_b_idx_val is None:
    tobe_b_idx_val = find_para(doc, "TO-BE:", "Normal", delta_idx)[0]
if tobe_b_idx_val is not None:
    bc = 0
    for i, para in enumerate(doc.paragraphs):
        if i <= tobe_b_idx_val:
            continue
        if para.style.name in ("Heading 2", "Heading 3"):
            break
        if para.style.name == "List Paragraph" and para.text.strip() == "xxx":
            if bc < len(tobe_bullets):
                set_para_text(para, tobe_bullets[bc])
                bc += 1

# 4.1 Impatti Operativi: '• xxx\n• xxx\n• xxx'
_, p_imp = find_para(doc, "\u2022 xxx", "Normal", delta_idx)
if p_imp:
    set_para_lines(p_imp, [
        "Riduzione del valore delle penali pagate tramite contestazione sistematica e documentata",
        "Riduzione del tempo di risposta alle notifiche (<3 giorni lavorativi vs. settimane AS-IS)",
        "Eliminazione del rischio di decadenza dei termini contrattuali di contestazione",
    ])

# 4.2 Invarianti: '- xxx' Normal x2
_, p_inv1 = find_para(doc, "4.2 Invarianti", "Heading 3")
invariants_start = find_para(doc, "4.2 Invarianti", "Heading 3")[0]
invariants = [
    "- Responsabilita' umana (HITL) per tutte le decisioni di contestazione e di storno",
    "- Processi legali e contrattuali di gestione del contenzioso con i committenti",
]
inv_count = 0
if invariants_start is not None:
    for i, para in enumerate(doc.paragraphs):
        if i <= invariants_start:
            continue
        if para.style.name in ("Heading 2", "Heading 3"):
            break
        if para.text.strip() == "- xxx" and inv_count < len(invariants):
            set_para_text(para, invariants[inv_count])
            inv_count += 1


# ── TABLE 10 \u2013 ROADMAP ─────────────────────────────────────────────────────────
t10 = doc.tables[10]
roadmap = [
    ("Fase 1 \u2013 Requisiti e Accessi\n(M1-M2)",
     "Allineamento con BSN/EGP TGX; mappatura API SAP/ERP, SCADA/PI, documentale; "
     "inventario contratti e committenti in scope",
     "Architettura tecnica validata; accordi IT per accessi API; inventario contratti",
     "2 mesi"),
    ("Fase 2 \u2013 Data Pipeline e Knowledge Base\n(M2-M4)",
     "ETL per acquisizione notifiche (email gateway), SCADA/PI e SAP; "
     "indicizzazione KB contratti (clausole, formule, exclusion clause, storico)",
     "ETL funzionante in dev; KB popolata con contratti e storico penali",
     "2 mesi"),
    ("Fase 3 \u2013 AI PoC\n(M3-M6)",
     "Sviluppo e test Penalty Classifier, Calculation Verifier e Dispute Generator "
     "su dati storici; validazione accuracy",
     "Accuracy classificazione >90%; match calcolo >95%; bozze validate da O&M",
     "3 mesi"),
    ("Fase 4 \u2013 Integrazione e Dashboard HITL\n(M5-M8)",
     "Integrazione con sistemi reali (SAP, SCADA/PI, email, documentale); "
     "sviluppo Dashboard HITL e Teams Bot; test operativi con referenti pilota",
     "Sistema integrato in pre-produzione; Dashboard validata dai referenti pilota",
     "3 mesi"),
    ("Fase 5 \u2013 Pilota e Validazione\n(M7-M9)",
     "Pilota operativo su subset committenti/impianti; monitoraggio KPI (savings, "
     "tempi, accuracy); go/no-go per scale-up",
     "Savings verificati; adozione >80% referenti coinvolti; report go/no-go",
     "3 mesi"),
    ("Fase 6 \u2013 Scale-up\n(M9-M12)",
     "Estensione a tutti i contratti/committenti EGP TGX Italia; attivazione "
     "Penalty Monitor; formazione referenti O&M; messa a regime",
     "Sistema in produzione su portfolio completo; piano ottimizzazione continua",
     "3 mesi"),
]
for i, row_data in enumerate(roadmap):
    if i + 1 < len(t10.rows):
        fill_row(t10, i + 1, list(row_data))


# ── TABLE 11 \u2013 KPI QUANTITATIVI ───────────────────────────────────────────────
t11 = doc.tables[11]
kpi_quant = [
    ("Valore penali contestate e vinte / anno (\u20ac)",
     "Baseline AS-IS da rilevare con BSN",
     "Target: +30% savings da contestazione sistematica",
     "Report SAP: penali stornate / totale contestate (\u20ac e %)"),
    ("Tempo medio risposta alle notifiche "
     "(giorni: ricezione \u2192 invio contestazione)",
     "Baseline AS-IS da rilevare con BSN",
     "<3 giorni lavorativi (AI + HITL)",
     "Timestamp ricezione \u2192 data invio contestazione nel sistema"),
    ("Tempo medio verifica e ricalcolo penale (ore/penale)",
     "Baseline AS-IS da rilevare con BSN",
     "<30 minuti (AI-assisted + HITL)",
     "Timestamp avvio verifica \u2192 approvazione HITL su Dashboard"),
    ("Tasso rispetto scadenze contrattuali contestazione (%)",
     "Baseline AS-IS da rilevare con BSN",
     "100% contestazioni entro i termini contrattuali",
     "Contestazioni inviate entro scadenza / totale avviate"),
]
for i, row_data in enumerate(kpi_quant):
    if i + 1 < len(t11.rows):
        fill_row(t11, i + 1, list(row_data))


# ── TABLE 12 \u2013 KPI QUALITATIVI ────────────────────────────────────────────────
t12 = doc.tables[12]
kpi_qual = [
    ("Soddisfazione referenti O&M",
     "Riduzione percepita del carico manuale; maggioranza lo valuta ridotto (scala Likert)",
     "Survey semestrale ai referenti O&M del team penali EGP TGX"),
    ("Qualita' contestazioni AI "
     "(% bozze approvate senza modifiche sostanziali)",
     ">70% bozze approvate senza modifiche entro 12 mesi dal go-live",
     "Tracking HITL: bozze approvate as-is / totale generate"),
    ("Tracciabilita' e auditabilita' del processo penali",
     "100% decisioni (accettare/contestare) documentate con motivazione nel sistema",
     "Audit log Dashboard HITL: % decisioni con motivazione registrata"),
    ("Completezza e aggiornamento Knowledge Base",
     "KB aggiornata entro 5 gg da ogni esito; copertura 100% contratti in scope",
     "Monitoraggio KB: % contratti indicizzati; lag medio aggiornamento post-esito"),
]
for i, row_data in enumerate(kpi_qual):
    if i + 1 < len(t12.rows):
        fill_row(t12, i + 1, list(row_data))


# ── SAVE ──────────────────────────────────────────────────────────────────────
output_dir = os.path.dirname(OUTPUT)
if output_dir:
    os.makedirs(output_dir, exist_ok=True)
doc.save(OUTPUT)
print(f"\nBlueprint IT O&M Gestione Penali (IT concisa) salvato in: {OUTPUT}")
print("Documento in italiano, non verboso, tutte le sezioni compilate.")
print("Valori baseline KPI AS-IS (in tabella KPI) da confermare con BSN EGP TGX.")
