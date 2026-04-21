"""
Blueprint Generator – PyroguardIAn & SolarEyes
Progetto: IT - BTP - PyroguardIAn - SolarEyes
Business Unit: EGP TGX (Enel Green Power – Technology and GrideXpansion)
Fonte materiale: EGP TGX\\IT - BTP -PyroguardIAn - SolarEyes
Template: Blueprint_1.4_vuoto.docx
Output language: Italian

OPEN QUESTIONS (sezioni con dati da integrare dal materiale sorgente):
  - roles: nomi e unità organizzative specifiche dei ruoli di progetto
  - sistemi_asis: nomi esatti dei sistemi SCADA/CMMS/ERP in uso negli impianti solari EGP
  - kpi_quantitativi: metriche baseline AS-IS da storico operativo EGP (es. n. fire event/anno, MTTR reale)
  - data_mapping: dettaglio formati dati e API disponibili per i sistemi esistenti
"""

import copy
import os
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _resolve_path(cli_index, env_name, default_value):
    """Resolve a path from CLI argument, environment variable, or a relative default."""
    if len(sys.argv) > cli_index and sys.argv[cli_index]:
        return sys.argv[cli_index]
    return os.environ.get(env_name, default_value)


TEMPLATE = _resolve_path(1, "BLUEPRINT_TEMPLATE", "Blueprint_1.4_vuoto.docx")
OUTPUT = _resolve_path(2, "BLUEPRINT_OUTPUT", "Blueprint_PyroguardIAn_SolarEyes.docx")

# ── helpers ──────────────────────────────────────────────────────────────────

def find_para(doc, fragment, style=None, start=0):
    if start is None:
        start = 0
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if fragment in p.text:
            if style is None or p.style.name == style:
                return i, p
    return None, None


def _get_rpr(para):
    """Return a deep copy of the first run's rPr, or None."""
    for run in para.runs:
        rpr = run._r.find(qn("w:rPr"))
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None


def set_para_text(para, new_text):
    """Replace all runs in para with a single run containing new_text, preserving formatting."""
    rpr = _get_rpr(para)
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    r_elem = OxmlElement("w:r")
    if rpr is not None:
        r_elem.append(rpr)
    t_elem = OxmlElement("w:t")
    t_elem.text = new_text
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_elem.append(t_elem)
    para._p.append(r_elem)


def set_para_lines(para, lines):
    """Replace paragraph content with multiple lines joined by w:br (soft return)."""
    rpr = _get_rpr(para)
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    for idx, line in enumerate(lines):
        r_elem = OxmlElement("w:r")
        if rpr is not None:
            r_elem.append(copy.deepcopy(rpr))
        t_elem = OxmlElement("w:t")
        t_elem.text = line
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r_elem.append(t_elem)
        para._p.append(r_elem)
        if idx < len(lines) - 1:
            r_br = OxmlElement("w:r")
            if rpr is not None:
                r_br.append(copy.deepcopy(rpr))
            br = OxmlElement("w:br")
            r_br.append(br)
            para._p.append(r_br)


def fill_cell(cell, text):
    """Replace cell text preserving formatting."""
    para = cell.paragraphs[0]
    set_para_text(para, text)


def fill_row(table, row_idx, values):
    """Fill a table row with a list of string values."""
    row = table.rows[row_idx]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            fill_cell(row.cells[ci], val)


# ── main ─────────────────────────────────────────────────────────────────────

doc = Document(TEMPLATE)

# ── TITOLO ────────────────────────────────────────────────────────────────────
_, p = find_para(doc, "Blueprint – xxx", "Heading 1")
if p:
    set_para_text(p, "Blueprint – PyroguardIAn & SolarEyes: Monitoraggio AI Impianti Solari EGP")

# ── SOMMARIO ESECUTIVO ────────────────────────────────────────────────────────
_, p = find_para(doc, "Processi Identificati")
if p:
    set_para_text(p, "Processi Identificati: Monitoraggio Impianti Fotovoltaici – Rilevazione Anomalie Termiche e Ispezione Visiva con AI")

_, p = find_para(doc, "Contesto Generale")
if p:
    set_para_lines(p, [
        "Contesto Generale:",
        "Il processo attuale di monitoraggio degli impianti fotovoltaici EGP/TGX si basa su ispezioni "
        "fisiche periodiche eseguite da tecnici sul campo, consultazione manuale di dashboard SCADA/DCS "
        "e analisi reattiva dei dati di producibilità. La rilevazione di hotspot termici, celle degradate "
        "e rischi di incendio avviene principalmente mediante campagne di termografia infrarosso con "
        "cadenza semestrale o annuale, senza copertura in tempo reale. Il sistema attuale non è in grado "
        "di rilevare precocemente anomalie diffuse su grandi superfici di pannelli.",
        "",
        "Proposta PyroguardIAn & SolarEyes: introduzione di un sistema AI composto da due moduli "
        "integrati — (1) PyroguardIAn: agente AI per la rilevazione real-time di anomalie termiche e "
        "rischi d'incendio tramite elaborazione di immagini termografiche (droni/sensori fissi); "
        "(2) SolarEyes: agente AI per l'ispezione visiva automatizzata dei pannelli tramite computer "
        "vision (immagini RGB drone/satellite) per rilevare danni fisici, sporco e degradazione. "
        "Entrambi i moduli condividono un Orchestrator Agent e una Dashboard HITL per la gestione "
        "delle anomalie con supervisione umana.",
        "",
        "Proprietario processo: EGP/TGX – O&M Solar | Popolazione impattata: Tecnici d'ufficio, "
        "Responsabili impianto, Team HSE, Tecnici di ispezione, Fornitori manutenzione"
    ])

# ── PROCESSO 1 – TITOLO ───────────────────────────────────────────────────────
_, p = find_para(doc, "Processo 1: xxx", "Heading 1")
if p:
    set_para_text(p, "Processo 1: Monitoraggio AI Impianti Fotovoltaici – PyroguardIAn & SolarEyes")

# ── 1.1 SCOPO ─────────────────────────────────────────────────────────────────
_, p = find_para(doc, "1.1 Scopo", "Heading 3")
scopo_start = _ if _ else 0
count = 0
for i, para in enumerate(doc.paragraphs):
    if i <= scopo_start:
        continue
    if para.style.name == "Heading 3":
        break
    if para.text.strip() == "xxx":
        if count == 0:
            set_para_text(para,
                "Il processo gestisce il monitoraggio continuo degli impianti fotovoltaici EGP, "
                "acquisendo immagini termografiche e RGB dai droni di ispezione e dai sensori fissi "
                "installati sugli impianti, elaborandole con modelli AI per la rilevazione di anomalie "
                "termiche (hotspot, celle by-pass, rischi d'incendio) e difetti visivi (micro-crack, "
                "soiling, rotture fisiche, ombreggiamento parziale).")
        elif count == 1:
            set_para_text(para,
                "Obiettivo principale: prevenire incendi e massimizzare la producibilità degli impianti "
                "solari attraverso il rilevamento precoce di anomalie (PyroguardIAn) e la pianificazione "
                "ottimizzata degli interventi manutentivi sulla base delle ispezioni visive (SolarEyes), "
                "riducendo i costi di ispezione manuale e i tempi di risposta agli eventi critici.")
        elif count == 2:
            set_para_text(para,
                "Sistemi in scope: droni di ispezione con payload termografico e RGB, sensori termici "
                "fissi su stringa/inverter, sistema SCADA impianto, piattaforma di gestione immagini "
                "(Image Repository), CMMS per ordini di lavoro, Knowledge Base tecnica fotovoltaico.")
        count += 1
        if count >= 3:
            break

# ── 1.2 FINALITÀ ──────────────────────────────────────────────────────────────
_, h12 = find_para(doc, "1.2 Finalità", "Heading 3")
_ = _ if _ is not None else 0
count = 0
for i, para in enumerate(doc.paragraphs):
    if i <= _:
        continue
    if para.style.name == "Heading 3":
        break
    if para.text.strip() == "xxx":
        if count == 0:
            set_para_text(para,
                "- Ridurre il rischio d'incendio negli impianti solari attraverso il rilevamento "
                "real-time di hotspot termici e condizioni anomale (modulo PyroguardIAn), con allerta "
                "immediata al team HSE e responsabile impianto.")
        elif count == 1:
            set_para_text(para,
                "- Automatizzare l'ispezione visiva dei pannelli fotovoltaici tramite analisi AI di "
                "immagini drone (modulo SolarEyes), riducendo la dipendenza da campagne manuali "
                "semestral/annuali e abbattendo il costo di ispezione per MW.")
        elif count == 2:
            set_para_text(para,
                "- Ottimizzare la pianificazione degli interventi manutentivi basandosi su anomalie "
                "rilevate e classificate dall'AI, prioritizzate per impatto sulla producibilità e "
                "rischio di sicurezza, con workflow di approvazione HITL.")
        count += 1
        if count >= 3:
            break

# ── 1.3 PERIMETRO ─────────────────────────────────────────────────────────────
_, h13 = find_para(doc, "1.3 Perimetro", "Heading 3")
h13_idx = _
in_scope_count = 0
out_scope_count = 0
in_out_mode = None
for i, para in enumerate(doc.paragraphs):
    if i <= h13_idx:
        continue
    if para.style.name == "Heading 3" or para.style.name == "Heading 2":
        break
    txt = para.text.strip()
    if txt == "IN SCOPE:":
        in_out_mode = "in"
    elif txt == "OUT OF SCOPE:":
        in_out_mode = "out"
    elif txt == "- xxx":
        if in_out_mode == "in":
            if in_scope_count == 0:
                set_para_text(para, "- Monitoraggio termografico real-time/near-real-time di impianti fotovoltaici EGP in scope tramite PyroguardIAn")
            elif in_scope_count == 1:
                set_para_text(para, "- Ispezione visiva automatizzata dei pannelli tramite analisi AI immagini RGB drone (SolarEyes)")
            elif in_scope_count == 2:
                set_para_text(para, "- Classificazione AI anomalie (hotspot, micro-crack, soiling, rotture) con prioritizzazione per rischio e impatto producibilità")
            in_scope_count += 1
        elif in_out_mode == "out":
            if out_scope_count == 0:
                set_para_text(para, "- Intervento fisico di sostituzione/pulizia pannelli (rimane in carico a tecnici e fornitori esterni)")
            elif out_scope_count == 1:
                set_para_text(para, "- Gestione operativa delle flotte droni (pianificazione voli e manutenzione droni non in scope nel pilota)")
            out_scope_count += 1

# ── 1.4 VINCOLI CHIAVE ────────────────────────────────────────────────────────
_, h14 = find_para(doc, "1.4 Vincoli chiave", "Heading 3")
_, pnorm = find_para(doc, "xxx", "List Paragraph", _)
if pnorm:
    set_para_text(pnorm,
        "Conformità alle normative aeronautiche per l'utilizzo di droni su impianti energetici "
        "(ENAC D-LG-2020-001, regolamento UE 2019/947); rispetto dei requisiti di cybersecurity "
        "per sistemi OT/SCADA (IEC 62443); normativa antincendio impianti fotovoltaici (CEI 82-25).")
_, ptec = find_para(doc, "Tecnici:", "Normal", _)
_, ptec_val = find_para(doc, "xxx", "Normal", _)
if ptec_val:
    set_para_text(ptec_val,
        "Disponibilità di connettività dati sugli impianti per trasmissione immagini real-time "
        "[DA DEFINIRE con IT Owner]; integrazione con SCADA/DCS esistente per correlazione dati "
        "energia-anomalie [DA DEFINIRE]; API CMMS per creazione automatica ordini di lavoro "
        "[DA DEFINIRE]; risoluzione minima immagini drone per classificazione AI (≥ 5 MP termico, "
        "≥ 20 MP RGB).")
_, porg = find_para(doc, "Organizzativi:", "Normal", _)
_, porg_val = find_para(doc, "xxx", "Normal", _)
if porg_val:
    set_para_text(porg_val,
        "Definizione del workflow di escalation per eventi critici (incendio/fumo) e relative "
        "responsabilità HSE; formazione tecnici d'ufficio e responsabili impianto sull'utilizzo "
        "della Dashboard HITL; governance per approvazione interventi AI-generated; "
        "change management per transizione da ispezioni manuali periodiche a monitoraggio continuo.")

# ── TABLE 0 – Sistemi coinvolti AS-IS ─────────────────────────────────────────
# OPEN QUESTION: i nomi esatti dei sistemi SCADA/ERP/CMMS in uso vanno confermati dal materiale sorgente
t0 = doc.tables[0]
systems = [
    ("SCADA / DCS Impianto", "Sistema di supervisione e controllo impianto fotovoltaico; fornisce dati di producibilità, allarmi stringa e stato inverter [NOME DA DEFINIRE]", "SCADA / Control System"),
    ("Sistema Gestione Immagini (Image Repository)", "Piattaforma di archiviazione e accesso alle immagini termografiche e RGB acquisite dai droni; base dati per analisi AI", "Image Management Platform"),
    ("Drone Termografico", "Velivolo UAV con payload termocamera infrarosso per acquisizione immagini termografiche degli impianti solari", "Inspection Hardware / UAV"),
    ("Drone RGB / Multispettrale", "Velivolo UAV con fotocamera RGB ad alta risoluzione per ispezione visiva dei pannelli fotovoltaici", "Inspection Hardware / UAV"),
    ("CMMS (Gestione Manutenzione)", "Sistema per la gestione degli ordini di lavoro e degli interventi manutentivi sugli impianti [NOME DA DEFINIRE]", "CMMS / EAM Platform"),
    ("Portale Meteo / Previsioni", "Servizio di previsione meteo per pianificazione finestre di volo e correlazione producibilità con irraggiamento", "External Weather API"),
    ("Archivio Documentale Tecnico", "Repository documentazione tecnica pannelli (schede tecniche, manuali OEM, storici interventi precedenti)", "Document Management (SharePoint/PLM)"),
]
for i, (s, r, t) in enumerate(systems):
    fill_row(t0, i + 1, [s, r, t])

# ── AS-IS SOTTO-PROCESSI LABEL ────────────────────────────────────────────────
_, pa = find_para(doc, "Sotto-processo A: xxx", "Normal")
if pa:
    set_para_text(pa, "Sotto-processo A: Pianificazione ed Esecuzione Campagne di Ispezione Termografica")
_, pb = find_para(doc, "Sotto-processo B: xxx", "Normal")
if pb:
    set_para_text(pb, "Sotto-processo B: Analisi Manuale Immagini e Identificazione Anomalie")
_, pc = find_para(doc, "Sotto-processo C: xxx", "Normal")
if pc:
    set_para_text(pc, "Sotto-processo C: Pianificazione e Coordinamento Interventi Manutentivi")

# ── TABLE 1 – AS-IS Sotto-processo A ─────────────────────────────────────────
t1 = doc.tables[1]
rows_a = [
    ("A1", "Team O&M pianifica campagna di ispezione termografica semestrale/annuale per impianto", "Responsabile O&M", "Calendario manutenzione, piano ispezioni annuali", "Campagna di ispezione pianificata con date e impianti target", "Calendario ERP / Email"),
    ("A2", "Coordinamento con operatore droni certificato (interno o fornitore esterno) per disponibilità e autorizzazioni volo", "Responsabile O&M / Fornitore", "Disponibilità operatore, autorizzazioni ENAC", "Operatore e drone prenotati; autorizzazioni acquisite", "Email / Telefono"),
    ("A3", "Operatore drone effettua volo di ispezione sull'impianto con payload termocamera IR", "Operatore drone / Tecnico", "Piano di volo, condizioni meteo favorevoli", "Immagini termografiche acquisite (formato TIFF/RJPEG)", "Drone + Payload Termocamera"),
    ("A4", "Immagini termografiche trasferite manualmente su workstation/server aziendale", "Tecnico d'ufficio", "File immagini da drone (SD card / link)", "Dataset immagini archiviato localmente o su server", "File transfer manuale (USB/FTP)"),
    ("A5", "Operatore drone effettua ulteriore volo RGB per documentazione visiva e georeferenziazione anomalie", "Operatore drone", "Piano di volo eseguito termografia", "Immagini RGB georeferenziate acquisite", "Drone + Fotocamera RGB"),
]
for i, row_data in enumerate(rows_a):
    if i + 1 < len(t1.rows):
        fill_row(t1, i + 1, list(row_data))

# ── TABLE 2 – AS-IS Sotto-processo B ─────────────────────────────────────────
t2 = doc.tables[2]
rows_b = [
    ("B1", "Tecnico specializzato analizza manualmente le immagini termografiche con software di elaborazione (es. FLIR Tools, IrfanView)", "Tecnico termografista", "Dataset immagini termografiche", "Lista anomalie termiche identificate (hotspot, bypass diode failure, ecc.)", "Software termografia (FLIR Tools / ResearchIR)"),
    ("B2", "Classificazione manuale delle anomalie secondo standard IEC 62446-3 (livelli severità 1-3)", "Tecnico termografista", "Lista anomalie identificate", "Anomalie classificate per severità e tipologia", "Software termografia + foglio Excel"),
    ("B3", "Tecnico compila report di ispezione in formato Word/Excel con foto e geolocalizzazione manuale delle anomalie", "Tecnico termografista", "Lista anomalie classificate + immagini di dettaglio", "Report ispezione termografica (Word/Excel/PDF)", "Microsoft Word / Excel"),
    ("B4", "Report inviato via email al Responsabile impianto e al team O&M per pianificazione interventi", "Tecnico termografista", "Report completato", "Email con report in allegato", "Email (Outlook)"),
    ("B5", "Responsabile impianto revisiona il report e valuta priorità di intervento", "Responsabile impianto", "Report ispezione", "Lista interventi prioritizzati (soggettiva)", "Email / Riunione"),
]
for i, row_data in enumerate(rows_b):
    if i + 1 < len(t2.rows):
        fill_row(t2, i + 1, list(row_data))

# ── TABLE 3 – AS-IS Sotto-processo C ─────────────────────────────────────────
t3 = doc.tables[3]
rows_c = [
    ("C1", "Responsabile impianto verifica la disponibilità delle imprese e dei tecnici manutentori", "Responsabile impianto", "Lista interventi prioritizzati", "Disponibilità tecnici/fornitori verificata", "Telefono / Email"),
    ("C2", "Verifica disponibilità ricambi (celle, pannelli sostituzione) e tempi di approvvigionamento", "Responsabile impianto", "Lista componenti da sostituire", "Disponibilità ricambi e lead time noti", "ERP / Telefono fornitori"),
    ("C3", "Pianificazione manuale degli interventi con definizione delle finestre operative favorevoli (meteo, produzione)", "Responsabile impianto", "Vincoli disponibilità, meteo, ricambi", "Piano interventi (foglio Excel o CMMS manuale)", "Excel / CMMS"),
    ("C4", "Emissione ordini di lavoro nel CMMS e comunicazione ai fornitori/tecnici", "Responsabile impianto / Back-office", "Piano interventi approvato", "Ordini di lavoro emessi nel CMMS", "CMMS [DA DEFINIRE]"),
    ("C5", "Esecuzione degli interventi (sostituzione pannelli, pulizia, riparazione connessioni) da parte di tecnici/fornitori", "Tecnico manutentore / Fornitore", "Ordini di lavoro ricevuti", "Intervento eseguito; verbale cartaceo compilato", "On-field / Carta"),
    ("C6", "Verifica post-intervento: nuovo volo termografico programmato (nella successiva campagna) per conferma efficacia", "Responsabile O&M", "Verbale intervento", "Pianificazione verifica post-intervento (ciclo successivo)", "Email / Calendario"),
]
for i, row_data in enumerate(rows_c):
    if i + 1 < len(t3.rows):
        fill_row(t3, i + 1, list(row_data))

# ── PROCESS CARDS AS-IS ───────────────────────────────────────────────────────
_, p = find_para(doc, "Card 1: xxx")
if p:
    set_para_text(p, "Card 1: Pianificazione ed Esecuzione Campagna Ispezione Termografica")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- Calendario manutenzione annuale; disponibilità operatori drone e autorizzazioni ENAC; condizioni meteo")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "1. Team O&M pianifica la campagna di ispezione (date, impianti, operatori)",
        "2. Coordinamento con operatore drone per disponibilità e autorizzazioni volo",
        "3. Volo termografico IR sull'impianto con acquisizione immagini",
        "4. Volo RGB per georeferenziazione e documentazione visiva",
        "5. Trasferimento manuale immagini su server aziendale"
    ])
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- Dataset immagini termografiche e RGB archiviato; campagna eseguita con cadenza semestrale/annuale")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- Cadenza di ispezione troppo bassa per rilevare eventi critici in tempo reale; costo elevato per campagna (operatore + volo)")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "N/A — processo manuale AS-IS")

_, p = find_para(doc, "Card 2: xxx")
if p:
    set_para_text(p, "Card 2: Analisi Manuale Immagini e Identificazione Anomalie")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- Dataset immagini termografiche e RGB acquisite durante la campagna di ispezione")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "1. Tecnico termografista analizza immagini con software dedicato (FLIR Tools / ResearchIR)",
        "2. Identificazione manuale hotspot, bypass failure, shading, soiling",
        "3. Classificazione anomalie per severità (IEC 62446-3)",
        "4. Compilazione report Word/Excel con coordinate anomalie",
        "5. Invio report via email al responsabile impianto"
    ])
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- Report di ispezione con lista anomalie classificate; inviato al responsabile impianto")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- Analisi manuale lenta (~2-5 giorni per impianto medio); soggettiva; nessuna correlazione automatica con dati SCADA")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "N/A — processo manuale AS-IS")

_, p = find_para(doc, "Card 3: xxx")
if p:
    set_para_text(p, "Card 3: Pianificazione e Coordinamento Interventi Manutentivi")

idx3 = _
c3_xxx = []
for i, para in enumerate(doc.paragraphs):
    if i <= idx3:
        continue
    if para.style.name == "Heading 3":
        break
    if para.text.strip() == "- xxx":
        c3_xxx.append((i, para))
if len(c3_xxx) >= 1:
    set_para_text(c3_xxx[0][1], "- Report ispezione con lista anomalie prioritizzate ricevuto da tecnico termografista")
if len(c3_xxx) >= 2:
    set_para_text(c3_xxx[1][1], "- Disponibilità fornitori, ricambi, condizioni meteo (verifica manuale)")

_, p = find_para(doc, "1. xxx", "Normal", idx3)
if p:
    set_para_lines(p, [
        "1. Responsabile impianto revisiona il report e definisce le priorità (soggettive)",
        "2. Verifica manuale disponibilità tecnici, fornitori e ricambi",
        "3. Pianificazione interventi su Excel o CMMS manuale",
        "4. Emissione ordini di lavoro e comunicazione ai tecnici/fornitori",
        "5. Esecuzione interventi e compilazione verbale cartaceo"
    ])
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- Interventi eseguiti; verbali su carta; nessuna verifica post-intervento in tempi brevi")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- Pianificazione non ottimizzata; gap temporale elevato tra rilevazione anomalia e intervento (settimane/mesi)")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "N/A — processo manuale AS-IS")

# ── PAIN POINT GENERALI ───────────────────────────────────────────────────────
_, p = find_para(doc, "1. Effort FTE elevato")
if p:
    set_para_text(p, "1. Effort FTE elevato: analisi manuale immagini termografiche richiede termografisti certificati; stimato 2-5 giorni FTE per impianto medio per campagna")
_, p = find_para(doc, "2. Tempi di attraversamento lunghi")
if p:
    set_para_text(p, "2. Tempi di attraversamento lunghi: dalla pianificazione campagna all'intervento manutentivo trascorrono mediamente 4-12 settimane; rischio d'incendio non gestito in tempo reale")
_, p = find_para(doc, "3. Colli di bottiglia")
if p:
    set_para_text(p, "3. Colli di bottiglia: disponibilità di termografisti certificati e droni limitata; impossibilità di coprire continuamente il parco impianti; incremento parco solare non scalabile con approccio attuale")
_, p = find_para(doc, "4. Qualità non uniforme:")
if p:
    set_para_text(p, "4. Qualità non uniforme:")
_, p = find_para(doc, "   - xxx", "Normal")
if p:
    set_para_text(p, "   - La qualità dell'analisi termografica dipende dall'esperienza del singolo termografista; nessuno standard automatizzato di classificazione; variabilità nel reporting")
_, p = find_para(doc, "5. Scarsa scalabilità:")
if p:
    set_para_text(p, "5. Scarsa scalabilità:")
_, p = find_para(doc, "   - xxx", "Normal")
if p:
    set_para_text(p, "   - Con la crescita del parco solare EGP, il modello a ispezioni periodiche manuali non regge il ritmo; impossibilità di monitorare continuativamente superfici di centinaia di MW")

pp_systems = [
    ("SCADA / DCS", "Dati di producibilità e allarmi stringa non correlati automaticamente con anomalie termiche; nessun trigger automatico per attivazione ispezione"),
    ("Image Repository", "Immagini archiviate senza metadati strutturati di anomalia; ricerca e consultazione solo manuale; nessuna analisi automatica pregresso"),
    ("Drone / Payload", "Acquisizione dati solo su richiesta (campagne pianificate); nessuna capacità di triggering automatico da eventi SCADA; operatività limitata da meteo e autorizzazioni"),
    ("CMMS", "Ordini di lavoro creati manualmente senza correlazione con la severità delle anomalie rilevate; nessuna prioritizzazione automatica basata su rischio"),
]
pp_idx = None
_, p = find_para(doc, "Pain Point per Documenti Specifici")
if _:
    pp_idx = _
    count = 0
    for i, para in enumerate(doc.paragraphs):
        if i <= pp_idx:
            continue
        if para.style.name.startswith("Heading"):
            break
        if para.text.strip() in ("xxx:", "xxx") and ":" not in para.text.replace("xxx:", ""):
            if count < len(pp_systems):
                set_para_text(para, pp_systems[count][0] + ":")
                count += 1
        elif para.text.strip() == "- xxx":
            idx_s = count - 1
            if 0 <= idx_s < len(pp_systems):
                set_para_text(para, "- " + pp_systems[idx_s][1])

# ── 2.5 CONTROLLI SPECIFICI ───────────────────────────────────────────────────
verifiche = [
    ("Drone Termografico / Immagini IR",
     ["Risoluzione termocamera ≥ 5 MP e NETD ≤ 50 mK", "File immagini in formato calibrato (TIFF radiometrico o RJPEG)", "Metadata GPS/georeferenziazione associato ad ogni frame"],
     ["Temperatura ambiente tra -10°C e +50°C al momento della ripresa", "Irraggiamento solare > 600 W/m² (per discriminare hotspot reali)", "Volo eseguito nelle ore centrali della giornata (10:00-15:00 ora solare)"]
    ),
    ("Image Repository / Dataset AI",
     ["ID impianto e ID stringa/modulo correttamente valorizzati nei metadati", "Associazione timestamp acquisizione – sessione volo – impianto", "Dataset di training con etichette validate da termografisti certificati"],
     ["Immagini non corrotte e leggibili dal modello AI (qualità > soglia minima)", "Completezza dataset: ≥ 90% superficie impianto coperta per ogni sessione", "Etichette anomalie conformi a tassonomia IEC 62446-3"]
    ),
    ("SCADA / DCS – Dati Producibilità",
     ["ID stringa e ID inverter correttamente associati all'anagrafica impianto", "Valore producibilità (kWh/h) e PR (Performance Ratio) disponibili per correlazione", "Timestamp dati SCADA sincronizzato con timestamp immagini drone"],
     ["Dati SCADA disponibili entro 15 minuti dalla misura per correlazione anomalie-producibilità", "Nessun gap > 1h nei dati di producibilità per il periodo di monitoraggio", "Allarmi stringa attivi correttamente segnalati (non mascherati da reset manuali)"]
    ),
    ("CMMS – Ordini di Lavoro",
     ["Codice anomalia AI presente nel campo note dell'ordine di lavoro generato", "Priorità intervento valorizzata secondo scala AI (Critical/High/Medium/Low)", "ID modulo/stringa anomalo correttamente referenziato nell'OdL"],
     ["Ordine di lavoro creato entro 4h dalla classificazione AI anomalia (per severity High/Critical)", "OdL linkato all'anomalia sorgente per tracciabilità completa", "Status OdL aggiornato al completamento per alimentare il feedback loop AI"]
    ),
]

_, h25 = find_para(doc, "2.5 Controlli Specifici", "Heading 3")
if h25 is None:
    _, h25 = find_para(doc, "2.5 Controlli", "Heading 3")

_, h25 = find_para(doc, "2.5 Controlli")
if _ is not None:
    vi = 0
    fi = 0
    ci = 0
    mode = "sys"

    for i, para in enumerate(doc.paragraphs):
        if i <= _:
            continue
        if para.style.name.startswith("Heading 2") or para.style.name.startswith("Heading 1"):
            break
        txt = para.text.strip()

        if mode == "sys" and txt in ("xxx", "xxx:"):
            if vi < len(verifiche):
                set_para_text(para, verifiche[vi][0])
                mode = "fields_header"
        elif mode == "fields_header" and txt == "Campi da Verificare:":
            mode = "fields"
            fi = 0
        elif mode == "fields" and txt == "- xxx":
            if fi < len(verifiche[vi][1]):
                set_para_text(para, "- " + verifiche[vi][1][fi])
                fi += 1
        elif mode == "fields" and txt == "Condizioni di Validità:":
            mode = "cond"
            ci = 0
        elif mode == "cond" and txt == "- xxx":
            if ci < len(verifiche[vi][2]):
                set_para_text(para, "- " + verifiche[vi][2][ci])
                ci += 1
        elif mode == "cond" and txt in ("xxx", "xxx:"):
            vi += 1
            if vi < len(verifiche):
                set_para_text(para, verifiche[vi][0])
                mode = "fields_header"

# ── TABLE 4 – DATA MAPPING ────────────────────────────────────────────────────
# OPEN QUESTION: formati API/dati da confermare con IT Owner
t4 = doc.tables[4]
data_mapping = [
    ("Immagine termografica RAW (TIFF radiometrico)", "Drone Termografico", "Image Repository / PyroguardIAn AI Engine", "TIFF / RJPEG con metadata GPS", "Acquisita per ogni sessione di volo; trigger analisi AI real-time"),
    ("Immagine RGB georeferenziata (ispezione visiva)", "Drone RGB", "Image Repository / SolarEyes AI Engine", "GeoTIFF / JPEG con metadata GPS", "Acquisita in parallelo al volo termografico o in campagna dedicata"),
    ("Risultato classificazione anomalia termica", "PyroguardIAn AI Engine", "Dashboard HITL / CMMS", "JSON strutturato (ID modulo, tipo anomalia, severità, coordinate)", "Generato in tempo reale post-elaborazione immagine"),
    ("Risultato classificazione difetto visivo pannello", "SolarEyes AI Engine", "Dashboard HITL / CMMS", "JSON strutturato (ID modulo, tipo difetto, confidence score, bounding box)", "Generato post-analisi immagine RGB; batch o near-real-time"),
    ("Dato producibilità stringa / inverter", "SCADA / DCS Impianto", "PyroguardIAn / SolarEyes AI Engine", "JSON / Modbus / OPC-UA [DA DEFINIRE]", "Correlazione automatica anomalia termica – calo producibilità"),
    ("Allarme SCADA (stringa off, inverter fault)", "SCADA / DCS Impianto", "Orchestrator Agent / Dashboard HITL", "JSON REST / Webhook [DA DEFINIRE]", "Trigger contestuale per analisi anomalia associata"),
    ("Ordine di lavoro generato da AI", "Orchestrator Agent (TO-BE)", "CMMS", "JSON / API REST CMMS [DA DEFINIRE]", "Workflow HITL: approvazione Responsabile impianto prima della creazione OdL"),
    ("Verbale post-intervento digitale", "Tecnico manutentore (Mobile App TO-BE)", "Knowledge Base / CMMS / Dashboard HITL", "JSON strutturato / PDF", "Feedback loop: alimenta retraining modelli AI"),
    ("Dati meteo e irraggiamento solare", "API Meteo Esterna (es. Solargis / ECMWF)", "SolarEyes / PyroguardIAn Engines", "JSON / CSV REST API", "Normalizzazione anomalie per condizioni irraggiamento; pianificazione voli"),
    ("Scheda tecnica / datasheet pannello fotovoltaico", "Archivio Documentale Tecnico (SharePoint)", "Knowledge Base (RAG)", "PDF / Word", "Reference per soglie termiche normali per modello pannello specifico"),
    ("Storico interventi manutentivi", "CMMS / Archivio cartaceo", "Knowledge Base (RAG)", "PDF / Excel / JSON", "Digitalizzazione necessaria; base per modelli predittivi degradazione"),
    ("Report ispezione termografica storico", "Archivio documentale", "Knowledge Base (RAG) / Training Dataset AI", "PDF / Word / Excel", "Etichettatura retroattiva per dataset di training modelli AI"),
]
for i, row_data in enumerate(data_mapping):
    if i + 1 < len(t4.rows):
        fill_row(t4, i + 1, list(row_data))

# ── TABLE 5 – ARCHITETTURA FUNZIONALE TO-BE ───────────────────────────────────
t5 = doc.tables[5]
arch_rows = [
    ("Orchestrator Agent", "Coordina il flusso end-to-end tra i moduli PyroguardIAn e SolarEyes; gestisce la prioritizzazione delle anomalie, il workflow HITL e l'interfaccia con CMMS e Dashboard", "LLM multi-agent orchestration (LangGraph / Azure AI Agent Service)"),
    ("PyroguardIAn AI Engine", "Analizza immagini termografiche IR in real-time / near-real-time per rilevare hotspot, bypass diode failure, soiling termico e potenziali rischi d'incendio; classifica per severità IEC 62446-3", "Computer Vision CNN/ViT + Thermal Image Processing (Python/ONNX/TensorRT)"),
    ("SolarEyes AI Engine", "Analizza immagini RGB drone ad alta risoluzione per rilevare micro-crack, rotture fisiche, soiling, bird droppings e ombreggiamento parziale dei pannelli", "Computer Vision CNN/YOLOv8 + RGB Image Analysis (Python/PyTorch)"),
    ("ETL & Data Integration Pipeline", "Acquisisce e normalizza immagini da droni/sensori, dati SCADA, meteo e CMMS; alimenta gli AI engine con dati contestualizzati e geo-referenziati", "Azure Data Factory / Apache Kafka / Delta Lake (streaming e batch)"),
    ("Knowledge Base Fotovoltaico (RAG)", "Repository indicizzato di documentazione tecnica (datasheet pannelli, manuali OEM, storici interventi, report ispezioni passati) per contestualizzazione AI e supporto decisionale", "RAG + Vector DB (Azure AI Search / Qdrant) + LLM (GPT-4o)"),
    ("Dashboard HITL (TO-BE)", "Interfaccia operatore per visualizzazione mappa anomalie su impianto, report AI, workflow approvazione interventi e KPI in tempo reale", "Power BI Embedded / React Web App / MS Teams Bot"),
    ("HITL Validation & Alerting Module", "Gestisce il workflow di notifica e approvazione umana per eventi critici (incendio/fumo: escalation immediata HSE) e interventi manutentivi pianificati", "Power Automate / Custom workflow engine / Push notification mobile"),
    ("Feedback & Retraining Pipeline", "Raccoglie i verbali post-intervento e le correzioni HITL per alimentare il retraining periodico dei modelli AI e migliorare l'accuracy nel tempo", "MLflow / Azure ML / CI-CD pipeline retraining"),
]
for i, row_data in enumerate(arch_rows):
    if i + 1 < len(t5.rows):
        fill_row(t5, i + 1, list(row_data))

# ── PUNTI APERTI 3.1b ─────────────────────────────────────────────────────────
_, p = find_para(doc, "- xxx: esposizione API interna/esterna [DA DEFINIRE]")
if p:
    set_para_text(p, "- SCADA/DCS: modalità esposizione dati producibilità e allarmi stringa verso ETL Pipeline e AI Engine [DA DEFINIRE con IT Owner OT/SCADA]")
_, p = find_para(doc, "- xxx: API disponibile? [DA DEFINIRE]")
if p:
    set_para_text(p, "- CMMS: API disponibile per creazione automatica OdL da AI e lettura storico interventi [DA DEFINIRE]; autenticazione e permessi di scrittura da concordare")
_, p = find_para(doc, "- xxx: modalità lettura/scrittura [DA DEFINIRE]")
if p:
    set_para_text(p, "- Image Repository: modalità di ingestione immagini drone (push automatico post-volo vs. upload manuale); metadati strutturati minimi richiesti da AI Engine [DA DEFINIRE]")
_, p = find_para(doc, "- xxx: accesso API [DA DEFINIRE]")
if p:
    set_para_text(p, "- Droni / Fleet Management: disponibilità API per scheduling missioni e recupero automatico metadati volo; compatibilità payload termocamera con standard dati AI Engine [DA DEFINIRE]")
_, p = find_para(doc, "I seguenti accessi/API sono [DA DEFINIRE]")
if p:
    set_para_text(p, "I seguenti accessi/API sono [DA DEFINIRE] con IT Owner e OT Owner prima del kick-off del progetto: SCADA/DCS, CMMS, Image Repository, Fleet Management Droni.")

# ── TO-BE SOTTO-PROCESSI LABEL ────────────────────────────────────────────────
_, p = find_para(doc, "Sotto-processo A: xxx (TO-BE)")
if p:
    set_para_text(p, "Sotto-processo A: Acquisizione Immagini, Analisi AI e Classificazione Anomalie (TO-BE)")
_, p = find_para(doc, "Sotto-processo B: xxx (TO-BE)")
if p:
    set_para_text(p, "Sotto-processo B: Gestione Alert, Workflow HITL e Pianificazione Interventi (TO-BE)")
_, p = find_para(doc, "Sotto-processo C: xxx (TO-BE)")
if p:
    set_para_text(p, "Sotto-processo C: Esecuzione Interventi, Feedback Loop e Miglioramento Continuo Modelli AI (TO-BE)")

# ── TABLE 6 – TO-BE Sotto-processo A ──────────────────────────────────────────
t6 = doc.tables[6]
rows_tobe_a = [
    ("A1", "ETL Pipeline acquisisce automaticamente immagini dal drone post-volo (o in streaming da sensori fissi) e dati SCADA di producibilità", "ETL Pipeline (automatico)", "Immagini IR/RGB drone, dati SCADA stringa/inverter, dati meteo", "Dataset normalizzato e arricchito con metadati geospaziali e producibilità", "Image Repository, SCADA, API Meteo", "Ingestione automatica; qualità dati verificata prima dell'analisi"),
    ("A2", "PyroguardIAn AI Engine analizza le immagini termografiche e rileva hotspot, bypass failure e anomalie termiche per ogni modulo", "PyroguardIAn AI Engine (automatico)", "Dataset immagini IR con metadati GPS e producibilità", "Mappa termica anomalie con: ID modulo, tipo anomalia, ΔT, severità IEC", "PyroguardIAn AI Engine, Image Repository", "CV inference + LLM reasoning per classificazione contestuale"),
    ("A3", "SolarEyes AI Engine analizza le immagini RGB e rileva difetti visivi (micro-crack, soiling, rotture fisiche) su ogni pannello", "SolarEyes AI Engine (automatico)", "Dataset immagini RGB georeferenziate", "Mappa difetti visivi con: ID modulo, tipo difetto, confidence score, bounding box", "SolarEyes AI Engine, Image Repository", "CV Object Detection (YOLOv8) + classificazione severità"),
    ("A4", "Orchestrator Agent correla i risultati PyroguardIAn e SolarEyes con i dati SCADA e knowledge base per contesto e stima impatto producibilità", "Orchestrator Agent", "Output PyroguardIAn + SolarEyes + dati SCADA + KB", "Report anomalie arricchito con: impatto producibilità stimato, note tecniche da KB, raccomandazione intervento", "Tutti i sistemi AI, KB, SCADA", "LLM synthesis + RAG retrieval da Knowledge Base"),
    ("A5", "Per anomalie critiche (rischio incendio/fumo): HITL immediato — notifica push a HSE e Responsabile impianto con dati e immagine anomalia", "HITL Validation Module (automatico → umano)", "Anomalia Critical da PyroguardIAn", "Notifica urgente inviata; workflow escalation HSE attivato", "Dashboard HITL, Mobile Push, Teams", "Alert immediato — Human decides azione di emergenza"),
    ("A6", "Per anomalie non critiche: Tecnico d'ufficio revisiona il report AI su Dashboard HITL, valida le classificazioni e approva la lista interventi", "Tecnico d'ufficio / Responsabile O&M", "Report AI consolidato su Dashboard HITL", "Lista interventi validata con priorità confermate", "Dashboard HITL", "Assist — Human reviews e approva classificazione AI"),
]
for i, row_data in enumerate(rows_tobe_a):
    if i + 1 < len(t6.rows):
        fill_row(t6, i + 1, list(row_data))

# ── TABLE 7 – TO-BE Sotto-processo B ──────────────────────────────────────────
t7 = doc.tables[7]
rows_tobe_b = [
    ("B1", "Orchestrator Agent riceve la lista interventi validati e avvia raccolta vincoli di pianificazione", "Orchestrator Agent", "Lista interventi approvati (dal Tecnico HITL)", "Raccolta vincoli avviata", "Orchestrator, API sistemi esterni", "Automatic constraint collection via function calling"),
    ("B2", "AI Engine interroga API meteo per previsioni 10 giorni e verifica finestre di irraggiamento favorevoli per massimizzare l'efficacia degli interventi", "Orchestrator Agent", "Lista impianti con interventi", "Previsioni meteo + irraggiamento per finestre manutenzione ottimali", "API Meteo (Solargis / OpenWeather)", "API call automatica + reasoning LLM"),
    ("B3", "Orchestrator verifica disponibilità ricambi (pannelli sostituzione, by-pass diode, connettori) interrogando CMMS/ERP", "Orchestrator Agent", "Lista componenti da sostituire", "Disponibilità ricambi, lead time, fornitori alternativi", "CMMS / ERP Ricambi [DA DEFINIRE API]", "DB query + LLM per analisi disponibilità"),
    ("B4", "Orchestrator genera piano interventi ottimizzato considerando criticità anomalia, impatto producibilità, meteo e disponibilità risorse", "Orchestrator Agent / Planning Module", "Interventi, meteo, ricambi, producibilità persa stimata", "Piano interventi ottimizzato con scheduling e priorità AI", "Orchestrator (constraint optimizer)", "Multi-constraint optimization + LLM"),
    ("B5", "HITL — Responsabile impianto revisiona il piano su Dashboard, può modificarlo e lo approva", "Responsabile impianto", "Piano AI con spiegazioni e motivazioni", "Piano approvato (o modificato e approvato)", "Dashboard HITL", "Assist — Human decides e approva piano"),
    ("B6", "Post-approvazione: Orchestrator crea automaticamente gli OdL nel CMMS e notifica i tecnici/fornitori assegnati", "Orchestrator Agent (post-HITL)", "Piano approvato", "OdL creati nel CMMS; notifiche inviate ai tecnici/fornitori", "CMMS, Email/Teams", "Automazione post-approvazione HITL"),
]
for i, row_data in enumerate(rows_tobe_b):
    if i + 1 < len(t7.rows):
        fill_row(t7, i + 1, list(row_data))

# ── TABLE 8 – TO-BE Sotto-processo C ──────────────────────────────────────────
t8 = doc.tables[8]
rows_tobe_c = [
    ("C1", "Tecnici/fornitori ricevono OdL strutturati con geolocalizzazione anomalia, immagine di dettaglio e istruzioni AI da Knowledge Base", "Tecnico manutentore / Fornitore", "OdL dal CMMS con dati AI allegati", "Intervento avviato con informazioni complete", "CMMS, Mobile App TO-BE, Dashboard HITL", "AI-enriched OdL con istruzioni contestuali"),
    ("C2", "Esecuzione intervento fisico (sostituzione pannelli, pulizia, riparazione connessioni)", "Tecnico manutentore / Fornitore", "OdL strutturato con geolocalizzazione anomalia", "Intervento eseguito sul campo", "On-field", "N/A – attività umana fisica"),
    ("C3", "Tecnico compila verbale post-intervento digitale via Mobile App (tipo anomalia confermata, azione eseguita, esito)", "Tecnico manutentore", "Intervento completato", "Verbale digitale strutturato (tipo anomalia, azione, componenti sostituiti, esito)", "Mobile App TO-BE", "Form guidato AI con suggerimenti compilazione basati su tipo anomalia"),
    ("C4", "SolarEyes/PyroguardIAn verificano automaticamente l'efficacia dell'intervento tramite confronto immagini pre/post (prossima sessione drone)", "AI Engine (post-intervento)", "Immagini post-intervento + verbale digitale", "Verifica efficacia: anomalia risolta / parzialmente risolta / persistente", "Image Repository, AI Engines", "Automated pre/post comparison + LLM assessment"),
    ("C5", "Feedback loop: verbali e esiti post-intervento alimentano retraining periodico dei modelli AI (mensile/trimestrale)", "Feedback & Retraining Pipeline", "Verbali digitali + etichette correzione HITL + esiti verifiche", "Modelli AI aggiornati con nuovi esempi; accuracy monitorata", "MLflow / Azure ML", "Continual learning pipeline automatizzata"),
    ("C6", "Orchestrator calcola e aggiorna KPI su Dashboard: tasso rilevamento, falsi positivi, MTTR, producibilità recuperata", "Orchestrator Agent", "Dati interventi completati, KPI produzione SCADA", "Dashboard KPI aggiornata in tempo reale", "Dashboard HITL, SCADA, CMMS", "KPI calculation automatica + LLM insights mensili"),
]
for i, row_data in enumerate(rows_tobe_c):
    if i + 1 < len(t8.rows):
        fill_row(t8, i + 1, list(row_data))

# ── PROCESS CARDS TO-BE ───────────────────────────────────────────────────────
_, p = find_para(doc, "Card 1: xxx (TO-BE)")
if p:
    set_para_text(p, "Card 1: Acquisizione Immagini e Classificazione AI Anomalie (TO-BE)")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- Immagini IR drone (PyroguardIAn); Immagini RGB drone (SolarEyes); Dati SCADA producibilità; Metadati GPS; Knowledge Base tecnica fotovoltaico")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "1. ETL Pipeline ingerisce automaticamente immagini e dati contestuali post-volo",
        "2. PyroguardIAn analizza immagini IR e classifica anomalie termiche per severità",
        "3. SolarEyes analizza immagini RGB e rileva difetti visivi su ogni pannello",
        "4. Orchestrator correla i risultati con SCADA e Knowledge Base",
        "5. HITL: Tecnico d'ufficio revisiona e approva il report anomalie su Dashboard"
    ])
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- Report anomalie validato con mappa geospaziale, classificazione AI, impatto producibilità e raccomandazioni intervento")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- PyroguardIAn, SolarEyes, ETL Pipeline, Image Repository, SCADA, Knowledge Base, Dashboard HITL")

_, p = find_para(doc, "Card 2: xxx (TO-BE)")
if p:
    set_para_text(p, "Card 2: Gestione Alert Critici e Pianificazione Interventi Ottimizzata (TO-BE)")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- Lista interventi validati; Previsioni meteo/irraggiamento; Disponibilità tecnici, fornitori e ricambi; Impatto producibilità stimato per anomalia")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "1. Per critical (incendio): notifica immediata push a HSE e Responsabile impianto",
        "2. Orchestrator raccoglie automaticamente vincoli (meteo, ricambi, disponibilità)",
        "3. AI genera piano interventi ottimizzato multi-vincolo",
        "4. HITL: Responsabile impianto revisiona e approva il piano",
        "5. Post-approvazione: OdL creati automaticamente nel CMMS e notifiche ai tecnici"
    ])
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- Piano interventi approvato; OdL creati nel CMMS; tecnici/fornitori notificati")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- Orchestrator Agent, HITL Module, Dashboard HITL, CMMS, API Meteo, ERP Ricambi")

_, p = find_para(doc, "Card 3: xxx (TO-BE)")
if p:
    set_para_text(p, "Card 3: Esecuzione Interventi e Feedback per Miglioramento Continuo AI (TO-BE)")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- OdL strutturati con dati AI (geolocalizzazione, immagine anomalia, istruzioni); verbali digitali post-intervento")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "1. Tecnici ricevono OdL arricchiti con dati AI e istruzioni Knowledge Base su Mobile App",
        "2. Esecuzione intervento fisico sul campo",
        "3. Tecnico compila verbale digitale strutturato via Mobile App (esito, componenti, note)",
        "4. AI verifica efficacia intervento con immagini post-intervento (prossima sessione drone)",
        "5. Feedback loop: verbali alimentano retraining mensile dei modelli AI"
    ])
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- Verbale digitale strutturato; verifica efficacia AI; KPI aggiornati su Dashboard; modelli AI migliorati")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- CMMS, Mobile App, AI Engines (post-verifica), Feedback Pipeline, Knowledge Base, Dashboard HITL, SCADA")

# ── 3.4 COSA NON FA L'AI ──────────────────────────────────────────────────────
_, p = find_para(doc, "3.4 Cosa NON fa l'AI", "Heading 3")
_, p_cosa = find_para(doc, "xxx", "Normal", _)
if p_cosa:
    set_para_lines(p_cosa, [
        "❌ Non esegue fisicamente gli interventi di manutenzione sui pannelli (rimane responsabilità di tecnici/fornitori)",
        "❌ Non attiva autonomamente protocolli di emergenza antincendio senza esplicita approvazione del Responsabile HSE (HITL obbligatorio per eventi critici)",
        "❌ Non pilota autonomamente i droni di ispezione né gestisce le autorizzazioni di volo ENAC",
        "❌ Non accede ai sistemi di controllo SCADA/DCS per modificare parametri operativi o arrestare inverter",
        "❌ Non emette ordini di lavoro nel CMMS senza validazione umana del Responsabile impianto",
        "❌ Non sostituisce la diagnosi tecnica specialistica del termografista certificato per casi complessi o ambigui",
        "❌ Non comunica autonomamente con fornitori o imprese appaltatrici senza approvazione HITL",
        "❌ Non certifica la conformità degli impianti a normative HSE o autorizzative (CEI 82-25, IEC 62446-3)"
    ])

# ── 3.5 COMPONENTE AI ─────────────────────────────────────────────────────────
_, p = find_para(doc, "3.5 Componente ai", "Heading 3")
_, p_ai = find_para(doc, "Xxx", "Normal", _)
if p_ai:
    set_para_lines(p_ai, [
        "Architettura multi-agent AI composta da 3 moduli specializzati coordinati da un Orchestrator:",
        "",
        "(1) Orchestrator Agent – coordina il flusso end-to-end tra PyroguardIAn e SolarEyes, "
        "gestisce prioritizzazione anomalie, workflow HITL e integrazione con CMMS/Dashboard.",
        "Pattern tassonomico: Orchestrator / Coordinator Agent",
        "",
        "(2) PyroguardIAn AI Engine – analizza immagini termografiche IR per rilevare hotspot, "
        "bypass diode failure e rischi d'incendio; classifica per severità IEC 62446-3; "
        "integra dati SCADA per correlazione anomalia termica - calo producibilità.",
        "Pattern tassonomico: Analyzer / Classifier Agent (Computer Vision + LLM reasoning)",
        "",
        "(3) SolarEyes AI Engine – analizza immagini RGB drone per rilevare difetti visivi "
        "(micro-crack, soiling, rotture fisiche); genera mappa geospaziale difetti con confidence score.",
        "Pattern tassonomico: Analyzer / Detector Agent (Computer Vision Object Detection)",
        "",
        "(4) Knowledge Agent (trasversale) – recupera e sintetizza informazioni tecniche da KB "
        "(datasheet pannelli, manuali OEM, storici interventi) per contestualizzare le analisi AI "
        "e fornire istruzioni operative ai tecnici via OdL enriched.",
        "Pattern tassonomico: Retriever / RAG Agent"
    ])

_, p_ai2 = find_para(doc, "**in questa sezione va indicata la componente AI", "Normal", _)
if p_ai2:
    set_para_text(p_ai2, "")

# ── 4. DELTA AS-IS vs TO-BE ───────────────────────────────────────────────────
_, p = find_para(doc, "Operatività")
if p:
    set_para_text(p, "Operatività – Confronto AS-IS / TO-BE")
_, p = find_para(doc, "AS-IS:", "Normal")
if p:
    set_para_text(p, "AS-IS:")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "- Ispezioni termografiche periodiche (semestrali/annuali) con operatori drone e termografisti manuali",
        "- Analisi immagini manuale con software dedicato (2-5 giorni per impianto); report Word/Excel",
        "- Nessuna correlazione automatica anomalie termiche - dati producibilità SCADA",
        "- Pianificazione interventi manuale senza ottimizzazione formale; gap 4-12 settimane anomalia→intervento",
        "- Nessun monitoraggio continuo; impossibile rilevare eventi critici (incendio) in real-time"
    ])

# ── 4.1 IMPATTI OPERATIVI ─────────────────────────────────────────────────────
_, h41 = find_para(doc, "4.1 Impatti Operativi", "Heading 3")
_, p = find_para(doc, "xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "- Riduzione stimata ~80% del tempo di analisi immagini: da 2-5 giorni manuali a 2-4 ore (AI + HITL)",
        "- Rilevamento precoce rischi incendio in tempo reale vs. attuale impossibilità di monitoraggio continuo",
        "- Riduzione costo per MW ispezionato: eliminazione costo termografista certificato per analisi di base",
        "- Gap anomalia→intervento ridotto da 4-12 settimane a 3-7 giorni (prioritizzazione AI + OdL automatico)",
        "- Aumento frequenza efficace di ispezione: da semestrale/annuale a continua (sensori fissi) + post-volo (drone)",
        "- Recupero producibilità stimato: intervento tempestivo su hotspot con ΔT > 20°C recupera mediamente 0.5-2% PR per modulo"
    ])

# ── 4.2 INVARIANTI ────────────────────────────────────────────────────────────
_, h42 = find_para(doc, "4.2 Invarianti", "Heading 3")
_ = _ if _ is not None else 0
inv_count = 0
for i, para in enumerate(doc.paragraphs):
    if i <= _:
        continue
    if para.style.name.startswith("Heading"):
        break
    if para.text.strip() == "- xxx":
        if inv_count == 0:
            set_para_text(para, "- La decisione finale di approvazione degli interventi e di attivazione protocolli di emergenza rimane in carico al responsabile umano (HITL obbligatorio)")
        elif inv_count == 1:
            set_para_text(para, "- I droni fisici e i sensori di campo rimangono i sistemi di acquisizione primari; l'AI analizza le immagini ma non le acquisisce")
        elif inv_count == 2:
            set_para_text(para, "- I tecnici eseguono fisicamente tutti gli interventi di manutenzione sui pannelli sul campo")
        elif inv_count == 3:
            set_para_text(para, "- Le normative aeronautiche (ENAC), antincendio (CEI 82-25) e tecniche (IEC 62446-3) per gli impianti fotovoltaici rimangono invariate e vincolanti")
        inv_count += 1

# ── 4.3 NUOVI REQUISITI ABILITANTI ───────────────────────────────────────────
_, h43 = find_para(doc, "4.3 Nuovi Requisiti", "Heading 3")
_ = _ if _ is not None else 0
req_count = 0
for i, para in enumerate(doc.paragraphs):
    if i <= _:
        continue
    if para.style.name == "Heading 3" or para.style.name == "Heading 2":
        break
    if para.text.strip() == "- xxx":
        if req_count == 0:
            set_para_text(para, "- Definizione e implementazione API/interfacce verso SCADA/DCS, CMMS e Image Repository [DA DEFINIRE con IT Owner OT]")
        elif req_count == 1:
            set_para_text(para, "- Creazione dataset di training etichettato per modelli CV (termografico + RGB): raccolta e annotazione immagini storiche esistenti con supporto di termografisti certificati")
        elif req_count == 2:
            set_para_text(para, "- Sviluppo interfaccia HITL (Dashboard TO-BE e Mobile App per tecnici) per visualizzazione anomalie georeferenziate e workflow approvazione")
        elif req_count == 3:
            set_para_text(para, "- Standardizzazione metadata drone (formato file, GPS, calibrazione termocamera) per garantire compatibilità con AI Engine [DA DEFINIRE con fornitori droni]")
        req_count += 1
    if para.text.strip() == "TO-BE:":
        break

# ── TO-BE bullets ─────────────────────────────────────────────────────────────
_, p = find_para(doc, "TO-BE:", "Normal")
if p:
    set_para_text(p, "TO-BE:")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "- PyroguardIAn monitora in continuo le immagini termografiche e rileva hotspot e rischi d'incendio in real-time con notifica immediata",
        "- SolarEyes analizza automaticamente ogni sessione drone e produce mappa difetti visivi geospaziale per l'intero impianto",
        "- Orchestrator correla anomalie AI con dati SCADA e genera piano interventi ottimizzato pronto per approvazione HITL",
        "- Responsabile impianto approva interventi e gestisce emergenze tramite dashboard dedicata: decisione informata in < 2h",
        "- Feedback loop continuo: ogni intervento migliora l'accuracy dei modelli AI nel tempo"
    ])

# ── TABLE 9 – ROADMAP ─────────────────────────────────────────────────────────
t9 = doc.tables[9]
roadmap = [
    ("Fase 1 – Assessment e Dataset Foundation\n(M1–M2)", "Raccolta e catalogazione immagini termografiche e RGB storiche; definizione standard metadata drone; mappatura API SCADA/CMMS; setup Image Repository strutturato", "Dataset storico catalogato (≥500 sessioni volo); accordi IT per accessi API OT/CMMS; specifiche metadata drone definite", "2 mesi"),
    ("Fase 2 – Dataset Labeling e Training Modelli CV\n(M2–M5)", "Etichettatura dataset con termografisti certificati (anomalie termiche) e tecnici O&M (difetti visivi); training modelli PyroguardIAn (IR) e SolarEyes (RGB); validazione su test set", "Modelli CV con accuracy ≥ 85% su test set interno; falsi positivi < 15%; approvazione qualitativa da termografisti", "3 mesi"),
    ("Fase 3 – Sviluppo Engine e ETL Pipeline\n(M4–M7)", "Implementazione ETL Pipeline; integrazione API SCADA/Image Repository; sviluppo Orchestrator Agent; Knowledge Base fotovoltaico; prima versione Dashboard HITL", "ETL funzionante su ambiente dev; AI Engine integrato con Image Repository e SCADA; Dashboard HITL v0.1 navigabile", "3 mesi"),
    ("Fase 4 – Integrazione CMMS e HITL Workflow\n(M6–M9)", "Integrazione con CMMS per creazione OdL automatica; sviluppo Mobile App tecnici; test workflow HITL completo (anomalia → OdL); test sistema con supervisori e tecnici pilota", "Workflow HITL end-to-end funzionante su ambiente pre-prod; Mobile App testata da 3-5 tecnici pilota; OdL creati automaticamente validati", "3 mesi"),
    ("Fase 5 – Pilota Operativo su Impianto Campione\n(M8–M10)", "Pilota operativo su 1-2 impianti EGP selezionati; monitoraggio KPI (tasso rilevamento, falsi positivi, MTTR, producibilità recuperata); go/no-go per scale-up", "Tasso rilevamento anomalie ≥ 90%; falsi positivi < 10%; MTTR ridotto ≥ 30%; adozione HITL > 80%; report go/no-go scale-up", "3 mesi"),
    ("Fase 6 – Scale-up Parco Solare EGP\n(M10–M12+)", "Estensione graduale a tutti gli impianti solari EGP in scope; formazione personale O&M e HSE; messa a regime retraining continuo e monitoring KPI; piano evoluzione futura (sensori fissi)", "Sistema in produzione su ≥ 80% impianti target; piano di miglioramento continuo definito; roadmap sensori fissi approvata", "3 mesi+"),
]
for i, row_data in enumerate(roadmap):
    if i + 1 < len(t9.rows):
        fill_row(t9, i + 1, list(row_data))

# ── TABLE 10 – KPI QUANTITATIVI ───────────────────────────────────────────────
# OPEN QUESTION: baseline AS-IS da storico operativo EGP da integrare dal materiale sorgente
t10 = doc.tables[10]
kpi_quant = [
    ("Tasso di rilevamento anomalie termiche critiche (hotspot ΔT > 20°C)", "~30-50% (ispezioni periodiche manuali; molti eventi non rilevati tra una campagna e l'altra) [DA CONFERMARE CON DATI EGP]", "> 90% (monitoraggio continuo AI)", "Confronto anomalie rilevate AI vs. anomalie confermate post-intervento (recall)"),
    ("Tasso falsi positivi classificazione AI", "N/A (baseline: analisi 100% manuale)", "< 10% false positive rate", "% anomalie AI risultate non confermate da tecnico HITL o post-intervento"),
    ("Tempo analisi immagini per sessione drone (ore FTE)", "~2-5 giorni FTE/impianto (analisi manuale termografista certificato)", "< 4 ore (AI engine + HITL review)", "Timestamp ingestione immagini → approvazione report HITL su Dashboard"),
    ("MTTR Solar – gap anomalia → intervento completato (giorni)", "~28-84 giorni (4-12 settimane: campagna → analisi → pianificazione → intervento) [DA CONFERMARE]", "< 7 giorni per anomalie High/Critical", "Tracking timestamp: rilevamento anomalia AI → chiusura OdL CMMS"),
    ("Costo per MW ispezionato (€/MW/anno)", "~500-2.000 €/MW/anno (campagna drone + analisi termografista) [DA QUANTIFICARE CON DATI EGP]", "Riduzione ≥ 40% (automazione analisi, riduzione campagne manuali)", "Costo totale ispezioni annuale / MW in produzione"),
    ("Producibilità recuperata (MWh/anno per anomalie risolte tempestivamente)", "Baseline da calcolare su storico EGP (producibilità persa da hotspot non rilevati) [DA DEFINIRE]", "+0.5-1.5% PR medio impianti monitorati", "Confronto PR (Performance Ratio) prima/dopo intervento su anomalie risolte; dato SCADA"),
    ("Numero eventi incendio/fumo rilevati in real-time vs. rilevati post-facto", "~0% real-time detection (nessun monitoraggio continuo attuale)", "> 95% detection entro 15 minuti da insorgenza evento", "Log eventi critici: timestamp insorgenza (da sensori/SCADA) vs. timestamp alert AI"),
    ("Coverage ispezione termografica impianti/anno (%)", "~50-70% impianti coperti annualmente (vincoli disponibilità operatori/droni)", "> 95% impianti coperti per sessione annuale (automazione analisi sblocca scalabilità)", "Ratio impianti analizzati / totale impianti in scope"),
]
for i, row_data in enumerate(kpi_quant):
    if i + 1 < len(t10.rows):
        fill_row(t10, i + 1, list(row_data))

# ── TABLE 11 – KPI QUALITATIVI ────────────────────────────────────────────────
t11 = doc.tables[11]
kpi_qual = [
    ("Qualità classificazione AI (giudizio termografisti certificati)", "Target: ≥ 85% delle classificazioni AI giudicate 'corrette o accettabili' da termografisti in review periodica; richiede benchmark semestrale", "Review semestrale a campione di classificazioni AI da parte di 2-3 termografisti certificati EGP"),
    ("Adozione Dashboard HITL da parte dei Responsabili impianto", "Target: > 90% dei Responsabili impianto utilizza regolarmente la Dashboard HITL per revisione anomalie entro 6 mesi dal go-live", "Log accessi Dashboard HITL; survey semestrale adozione"),
    ("Qualità dei piani intervento AI (% approvati senza modifiche sostanziali)", "Target: > 70% dei piani intervento generati dall'AI approvati dal Responsabile impianto senza modifiche sostanziali entro 12 mesi", "Tracking HITL: piani approvati as-is / totale piani; analisi motivazioni modifiche"),
    ("Completezza verbali post-intervento digitali", "Target: > 95% dei verbali compilati in formato digitale strutturato tramite Mobile App entro 4h dalla fine intervento", "Tracking invio verbali digitali / totale OdL chiusi nel CMMS"),
    ("Accuracy Knowledge Base retrieval (% query risolte senza escalation umana)", "Target: > 80% delle query tecniche da tecnici (istruzioni intervento, datasheet) risolte autonomamente dalla KB senza escalation", "Tracking query KB: risolte autonomamente / totale query tecnici in campo"),
    ("Soddisfazione utenti sistema AI (NPS interno tecnici e responsabili)", "Target: NPS > 35 entro 12 mesi dal go-live; focus su riduzione carico manuale percepito da tecnici d'ufficio e termografisti", "Survey NPS trimestrale a tutti gli utenti del sistema (tecnici, responsabili, HSE)"),
    ("Aggiornamento continuo modelli AI (stabilità accuracy nel tempo)", "Target: nessuna degradazione accuracy > 5% su test set rolling nei 12 mesi post go-live; retraining mensile con nuovi dati validati", "Monitoring automatico accuracy con alerting su degradazione; MLflow tracking"),
    ("Tracciabilità decisioni AI per audit HSE (% eventi con spiegazione)", "Target: 100% degli alert critici (incendio/fumo) documentati con: immagine originale, output AI, timestamp, azione HITL intrapresa; conformità audit HSE", "Audit log automatico Dashboard HITL + CMMS per ogni evento severity Critical"),
]
for i, row_data in enumerate(kpi_qual):
    if i + 1 < len(t11.rows):
        fill_row(t11, i + 1, list(row_data))

# ── SAVE ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Blueprint saved to:\n{OUTPUT}")
print()
print("=" * 70)
print("OPEN QUESTIONS — Sezioni che richiedono dati dal materiale sorgente:")
print("=" * 70)
print("  1. roles           — Nomi e unità organizzative specifiche dei ruoli")
print("                       di progetto (Business Owner, IT Owner, Data Owner,")
print("                       Product Owner EGP TGX)")
print("  2. sistemi_asis    — Nomi esatti dei sistemi SCADA/DCS, CMMS, ERP in uso")
print("                       negli impianti solari EGP (segnalati come [DA DEFINIRE])")
print("  3. kpi_quantitativi — Baseline AS-IS reali da storico operativo EGP:")
print("                       n. eventi incendio/anno, MTTR reale, costo/MW ispezione,")
print("                       % coverage attuale (segnalati come [DA CONFERMARE])")
print("  4. data_mapping    — Dettaglio formati API dei sistemi OT/SCADA e CMMS")
print("                       (protocollo, autenticazione, rate limit)")
print("  5. architettura    — Scelta stack tecnologico AI finale da validare con")
print("                       IT Owner (cloud provider, MLOps platform, drone vendor)")
