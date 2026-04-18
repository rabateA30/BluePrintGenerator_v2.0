"""
Blueprint Generator – Manutenzione Predittiva Impianti Eolici
Source: Ipotesi implementazione GenAI su impianto Wind.docx
Template: Blueprint_1.3_enhanced.docx
Output language: Italian
"""

import copy
import os
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree


def _resolve_path(cli_index, env_name, default_value):
    """Resolve a path from CLI argument, environment variable, or a relative default."""
    if len(sys.argv) > cli_index and sys.argv[cli_index]:
        return sys.argv[cli_index]
    return os.environ.get(env_name, default_value)


TEMPLATE = _resolve_path(1, "BLUEPRINT_TEMPLATE", "Blueprint_1.3_enhanced.docx")
OUTPUT = _resolve_path(2, "BLUEPRINT_OUTPUT", "Blueprint_PredictiveMaint_Wind.docx")
# ── helpers ──────────────────────────────────────────────────────────────────

def find_para(doc, fragment, style=None, start=0):
    if start is None:
        start = 0
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if fragment in p.text:
            if style is None or p.style.name == style:
                return i, p
    return None, None


def _get_rpr(para):
    """Return a deep copy of the first run's rPr, or None."""
    for run in para.runs:
        rpr = run._r.find(qn("w:rPr"))
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None


def set_para_text(para, new_text):
    """Replace all runs in para with a single run containing new_text, preserving formatting."""
    rpr = _get_rpr(para)
    # clear existing runs
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    # create new run
    r_elem = OxmlElement("w:r")
    if rpr is not None:
        r_elem.append(rpr)
    t_elem = OxmlElement("w:t")
    t_elem.text = new_text
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_elem.append(t_elem)
    para._p.append(r_elem)


def set_para_lines(para, lines):
    """Replace paragraph content with multiple lines joined by w:br (soft return)."""
    rpr = _get_rpr(para)
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    for idx, line in enumerate(lines):
        r_elem = OxmlElement("w:r")
        if rpr is not None:
            r_elem.append(copy.deepcopy(rpr))
        t_elem = OxmlElement("w:t")
        t_elem.text = line
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r_elem.append(t_elem)
        para._p.append(r_elem)
        if idx < len(lines) - 1:
            r_br = OxmlElement("w:r")
            if rpr is not None:
                r_br.append(copy.deepcopy(rpr))
            br = OxmlElement("w:br")
            r_br.append(br)
            para._p.append(r_br)


def fill_cell(cell, text):
    """Replace cell text preserving formatting."""
    para = cell.paragraphs[0]
    set_para_text(para, text)


def fill_row(table, row_idx, values):
    """Fill a table row with a list of string values."""
    row = table.rows[row_idx]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            fill_cell(row.cells[ci], val)


# ── main ─────────────────────────────────────────────────────────────────────

doc = Document(TEMPLATE)

# ── TITOLO ────────────────────────────────────────────────────────────────────
_, p = find_para(doc, "Blueprint – xxx", "Heading 1")
if p:
    set_para_text(p, "Blueprint – Manutenzione Predittiva Impianti Eolici")

# ── SOMMARIO ESECUTIVO ────────────────────────────────────────────────────────
_, p = find_para(doc, "Processi Identificati")
if p:
    set_para_text(p, "Processi Identificati: Manutenzione Predittiva Impianti Eolici – Automazione con Agenti AI")

_, p = find_para(doc, "Contesto Generale")
if p:
    set_para_lines(p, [
        "Contesto Generale:",
        "Il processo attuale di manutenzione predittiva degli impianti eolici EGP/TGX si basa su analisi manuale "
        "delle dashboard fornite dal sistema SAS, che elabora dati provenienti da sensori di vibrazione (POM) e "
        "dal sistema Iceberg. Il personale d'ufficio analizza gli alert predittivi e comunica manualmente al "
        "supervisore d'impianto i componenti che necessitano di intervento. Il supervisore pianifica gli "
        "interventi considerando disponibilità di imprese/fornitori, condizioni meteo e disponibilità di ricambi.",
        "",
        "Proposta: introduzione di un'architettura multi-agent AI (Orchestrator, Alarm Analysis, Planning, "
        "Knowledge, Spare Parts Agent) per automatizzare l'analisi degli allarmi, ottimizzare la pianificazione "
        "degli interventi e strutturare la knowledge tecnica, mantenendo il supervisore come punto di validazione "
        "HITL prima dell'esecuzione.",
        "",
        "Proprietario processo: EGP/TGX – O&M Wind | Popolazione impattata: Tecnici d'ufficio, Supervisori impianto, Fornitori manutenzione"
    ])

# ── PROCESSO 1 – TITOLO ───────────────────────────────────────────────────────
_, p = find_para(doc, "Processo 1: xxx", "Heading 1")
if p:
    set_para_text(p, "Processo 1: Manutenzione Predittiva Impianti Eolici")

# ── 1.1 SCOPO ─────────────────────────────────────────────────────────────────
_, p = find_para(doc, "1.1 Scopo", "Heading 3")
scopo_start = _ if _ else 0
# fill the three xxx paragraphs after 1.1
count = 0
for i, para in enumerate(doc.paragraphs):
    if i <= scopo_start:
        continue
    if para.style.name == "Heading 3":
        break
    if para.text.strip() == "xxx":
        if count == 0:
            set_para_text(para,
                "Il processo gestisce la manutenzione predittiva degli impianti eolici, "
                "acquisendo dati di vibrazione, temperatura e produzione da sensori di campo "
                "tramite il sistema POM e gli eventi predittivi classificati da SAS e Iceberg.")
        elif count == 1:
            set_para_text(para,
                "Obiettivo principale: identificare tempestivamente i componenti (gearbox, cuscinetti, "
                "generatore) a rischio di guasto e pianificare interventi correttivi prima che si "
                "verifichino fermate impreviste dell'impianto.")
        elif count == 2:
            set_para_text(para,
                "Sistemi in scope: SAS (analytics predittivi), Iceberg (eventi predittivi), "
                "POM (dati vibrazione sensori), Knowledge Base tecnica (documentazione e storici).")
        count += 1
        if count >= 3:
            break

# ── 1.2 FINALITÀ ──────────────────────────────────────────────────────────────
# three xxx paragraphs after heading 1.2
_, h12 = find_para(doc, "1.2 Finalità", "Heading 3")
count = 0
for i, para in enumerate(doc.paragraphs):
    if i <= _:
        continue
    if para.style.name == "Heading 3":
        break
    if para.text.strip() == "xxx":
        if count == 0:
            set_para_text(para,
                "- Ridurre il Mean Time To Repair (MTTR) della manutenzione predittiva "
                "attraverso analisi automatizzata degli allarmi e pianificazione ottimizzata degli interventi.")
        elif count == 1:
            set_para_text(para,
                "- Automatizzare l'analisi degli alert predittivi provenienti da SAS/Iceberg, "
                "classificandoli per priorità e urgenza senza intervento manuale del personale d'ufficio.")
        elif count == 2:
            set_para_text(para,
                "- Ottimizzare la pianificazione degli interventi considerando vincoli economici "
                "(prezzo energia), previsioni meteo, disponibilità fornitori/ricambi e finestre di manutenzione.")
        count += 1
        if count >= 3:
            break

# ── 1.3 PERIMETRO ─────────────────────────────────────────────────────────────
_, h13 = find_para(doc, "1.3 Perimetro", "Heading 3")
# IN SCOPE / OUT OF SCOPE: iterate all paragraphs in 1.3 section
h13_idx = _
in_scope_count = 0
out_scope_count = 0
in_out_mode = None
for i, para in enumerate(doc.paragraphs):
    if i <= h13_idx:
        continue
    if para.style.name == "Heading 3" or para.style.name == "Heading 2":
        break
    txt = para.text.strip()
    if txt == "IN SCOPE:":
        in_out_mode = "in"
    elif txt == "OUT OF SCOPE:":
        in_out_mode = "out"
    elif txt == "- xxx":
        if in_out_mode == "in":
            if in_scope_count == 0:
                set_para_text(para, "- Monitoraggio predittivo di turbine eoliche: analisi componenti critici (gearbox, cuscinetti, generatore)")
            elif in_scope_count == 1:
                set_para_text(para, "- Generazione automatizzata di piani di manutenzione ottimizzati multi-vincolo")
            elif in_scope_count == 2:
                set_para_text(para, "- Integrazione dati da SAS, Iceberg e POM tramite pipeline ETL per alimentare gli agenti AI")
            in_scope_count += 1
        elif in_out_mode == "out":
            if out_scope_count == 0:
                set_para_text(para, "- Esecuzione fisica degli interventi di manutenzione (rimane in carico a tecnici/fornitori)")
            elif out_scope_count == 1:
                set_para_text(para, "- Gestione diretta degli acquisti e approvvigionamento ricambi (fuori scope nel pilota)")
            out_scope_count += 1

# ── 1.4 VINCOLI CHIAVE ────────────────────────────────────────────────────────
_, h14 = find_para(doc, "1.4 Vincoli chiave", "Heading 3")
# Normativi bullet
_, pnorm = find_para(doc, "xxx", "List Paragraph", _)
if pnorm: set_para_text(pnorm,
    "Conformità alle normative di sicurezza per l'esercizio di impianti eolici (CEI, IEC 61400); "
    "rispetto dei requisiti di data governance e privacy per i dati di produzione e manutenzione.")
# Tecnici
_, ptec = find_para(doc, "Tecnici:", "Normal", _)
_, ptec_val = find_para(doc, "xxx", "Normal", _)
if ptec_val: set_para_text(ptec_val,
    "Disponibilità di API/interfacce verso SAS, Iceberg e POM (da definire con IT Owner); "
    "consistenza dei formati dati tra sistemi sorgente; latenza accettabile per analisi near-real-time.")
# Organizzativi
_, porg = find_para(doc, "Organizzativi:", "Normal", _)
_, porg_val = find_para(doc, "xxx", "Normal", _)
if porg_val: set_para_text(porg_val,
    "Formazione supervisori e tecnici d'ufficio sull'utilizzo della dashboard TO-BE e sui flussi HITL; "
    "definizione della governance per l'approvazione dei piani AI-generated; change management organizzativo.")

# ── TABLE 0 – Sistemi coinvolti AS-IS ─────────────────────────────────────────
t0 = doc.tables[0]
systems = [
    ("SAS", "Sistema di analytics predittivi per componenti turbina (gearbox, cuscinetti); genera score di rischio e alert", "Analytics / ML Platform"),
    ("Iceberg", "Piattaforma di gestione eventi predittivi; classifica e cataloga gli eventi di anomalia per impianto", "Event Management Platform"),
    ("POM", "Sistema di acquisizione dati vibrazione e sensori di campo (accelerometri, temperatura, RPM)", "SCADA / Sensor Data Platform"),
    ("Dashboard BI", "Strumento di visualizzazione delle predittive; accesso manuale da parte del personale d'ufficio", "Business Intelligence / Reporting"),
    ("Sistema Gestione Fornitori", "Gestione disponibilità imprese e fornitori per pianificazione interventi", "ERP / Supplier Management"),
    ("Email / Telefono", "Canale di comunicazione informale tra ufficio e supervisore impianto per notifica interventi", "Communication (informal)"),
]
for i, (s, r, t) in enumerate(systems):
    fill_row(t0, i + 1, [s, r, t])

# ── AS-IS SOTTO-PROCESSI LABEL ────────────────────────────────────────────────
_, pa = find_para(doc, "Sotto-processo A: xxx", "Normal")
if pa: set_para_text(pa, "Sotto-processo A: Acquisizione e Analisi Dati Sensori e Predittivi")
_, pb = find_para(doc, "Sotto-processo B: xxx", "Normal")
if pb: set_para_text(pb, "Sotto-processo B: Monitoraggio Dashboard e Comunicazione al Supervisore")
_, pc = find_para(doc, "Sotto-processo C: xxx", "Normal")
if pc: set_para_text(pc, "Sotto-processo C: Pianificazione e Coordinamento Interventi Manutenzione")

# ── TABLE 1 – AS-IS Sotto-processo A ─────────────────────────────────────────
t1 = doc.tables[1]
rows_a = [
    ("A1", "Sensori di campo acquisiscono dati di vibrazione, temperatura e RPM dalle turbine", "Sensori / Impianto", "Dati raw sensori (accelerometria, temperatura)", "Stream dati acquisito", "POM"),
    ("A2", "POM raccoglie e aggrega i dati dai sensori e li trasmette a SAS per l'elaborazione", "POM (automatico)", "Stream dati sensori", "Dataset aggregato per turbina", "POM → SAS"),
    ("A3", "SAS elabora i dati con modelli predittivi e calcola score di rischio per ogni componente", "SAS (automatico)", "Dataset aggregato POM", "Score predittivo per componente (es. gearbox)", "SAS"),
    ("A4", "SAS genera alert predittivi per i componenti con score sopra soglia configurata", "SAS (automatico)", "Score predittivo", "Alert predittivi classificati", "SAS → Iceberg"),
    ("A5", "Iceberg riceve e cataloga gli eventi predittivi generati da SAS", "Iceberg (automatico)", "Alert da SAS", "Evento predittivo catalogato nel sistema Iceberg", "Iceberg"),
    ("A6", "Dashboard BI si aggiorna con i nuovi alert e score predittivi disponibili", "Dashboard BI (automatico)", "Dati SAS e Iceberg", "Dashboard aggiornata con stato impianti", "SAS, Iceberg, Dashboard BI"),
]
for i, row_data in enumerate(rows_a):
    if i + 1 < len(t1.rows):
        fill_row(t1, i + 1, list(row_data))

# ── TABLE 2 – AS-IS Sotto-processo B ─────────────────────────────────────────
t2 = doc.tables[2]
rows_b = [
    ("B1", "Personale d'ufficio accede alla dashboard BI per il monitoraggio giornaliero degli impianti", "Tecnico d'ufficio", "Dashboard BI aggiornata", "Sessione di analisi avviata", "Dashboard BI"),
    ("B2", "Analisi manuale degli alert predittivi: il tecnico esamina componenti a rischio, livello severità e storico", "Tecnico d'ufficio", "Alert predittivi su dashboard", "Lista componenti a rischio identificati", "Dashboard BI, SAS"),
    ("B3", "Il tecnico valuta la priorità degli interventi basandosi su esperienza personale e soglie di allerta", "Tecnico d'ufficio", "Lista componenti a rischio", "Priorità intervento assegnata (urgente/pianificabile)", "Dashboard BI"),
    ("B4", "Il tecnico d'ufficio contatta il supervisore dell'impianto telefonicamente o via email", "Tecnico d'ufficio", "Priorità interventi identificate", "Comunicazione inviata al supervisore", "Email / Telefono"),
    ("B5", "Il supervisore riceve la comunicazione e ne prende nota per pianificazione", "Supervisore impianto", "Comunicazione tecnico d'ufficio", "Informazioni ricevute, avvio pianificazione", "Email / Telefono"),
]
for i, row_data in enumerate(rows_b):
    if i + 1 < len(t2.rows):
        fill_row(t2, i + 1, list(row_data))

# ── TABLE 3 – AS-IS Sotto-processo C ─────────────────────────────────────────
t3 = doc.tables[3]
rows_c = [
    ("C1", "Supervisore verifica la disponibilità delle imprese e dei fornitori per le date ipotizzate", "Supervisore impianto", "Lista componenti da manutenere", "Disponibilità fornitori verificata", "Telefono / Email"),
    ("C2", "Supervisore controlla le previsioni meteo per identificare finestre operative favorevoli", "Supervisore impianto", "Previsioni meteo esterne", "Finestre meteo idonee identificate", "Portale meteo (manuale)"),
    ("C3", "Supervisore verifica la disponibilità di ricambi necessari per gli interventi", "Supervisore impianto", "Lista componenti da sostituire", "Disponibilità ricambi verificata", "ERP / Telefono fornitori"),
    ("C4", "Supervisore pianifica gli interventi integrando tutti i vincoli identificati (disponibilità, meteo, ricambi)", "Supervisore impianto", "Vincoli identificati", "Piano manutenzione bozza", "Manuale / Foglio Excel"),
    ("C5", "Supervisore coordina la logistica: accessi impianto, dotazioni sicurezza, trasporto ricambi", "Supervisore impianto", "Piano manutenzione bozza", "Logistica confermata", "Telefono / Email"),
    ("C6", "Emissione ordini di lavoro per i tecnici e i fornitori", "Supervisore impianto", "Piano manutenzione approvato", "Ordini di lavoro emessi", "ERP / Manuale"),
    ("C7", "Esecuzione degli interventi di manutenzione da parte di tecnici/fornitori", "Tecnico / Fornitore", "Ordini di lavoro", "Verbale intervento compilato", "Mobile App / Carta"),
]
for i, row_data in enumerate(rows_c):
    if i + 1 < len(t3.rows):
        fill_row(t3, i + 1, list(row_data))

# ── PROCESS CARDS AS-IS ───────────────────────────────────────────────────────
# Card 1
_, p = find_para(doc, "Card 1: xxx")
if p: set_para_text(p, "Card 1: Acquisizione Dati e Generazione Alert Predittivi")
_, p = find_para(doc, "- xxx", "Normal", _)  # input
if p: set_para_text(p, "- Dati vibrazione, temperatura, RPM da sensori di campo (turbine eoliche)")
_, p = find_para(doc, "1. xxx", "Normal", _)  # attività
if p: set_para_lines(p, [
    "1. POM acquisisce dati sensori e li trasmette a SAS",
    "2. SAS elabora modelli predittivi e calcola score di rischio per componenti (gearbox, cuscinetti)",
    "3. Alert predittivi generati e catalogati su Iceberg",
    "4. Dashboard BI aggiornata con stato impianti e alert attivi"
])
_, p = find_para(doc, "- xxx", "Normal", _)  # output
if p: set_para_text(p, "- Alert predittivi classificati; Dashboard BI aggiornata con score di rischio per componente")
_, p = find_para(doc, "- xxx", "Normal", _)  # pain point
if p: set_para_text(p, "- Latenza nell'aggiornamento dati; mancata correlazione automatica tra alert SAS e contesto storico manutenzione")
_, p = find_para(doc, "- xxx", "Normal", _)  # HITL
if p: set_para_text(p, "N/A — processo automatico AS-IS (solo raccolta dati e generazione alert)")

# Card 2
_, p = find_para(doc, "Card 2: xxx")
if p: set_para_text(p, "Card 2: Monitoraggio Dashboard e Comunicazione al Supervisore")
_, p = find_para(doc, "- xxx", "Normal", _)  # input
if p: set_para_text(p, "- Dashboard BI con alert predittivi aggiornati da SAS/Iceberg")
_, p = find_para(doc, "1. xxx", "Normal", _)  # attività
if p: set_para_lines(p, [
    "1. Tecnico d'ufficio apre la dashboard BI e visualizza gli alert attivi",
    "2. Analisi manuale dei componenti a rischio per impianto",
    "3. Assegnazione soggettiva delle priorità di intervento",
    "4. Comunicazione informale (telefono/email) al supervisore d'impianto"
])
_, p = find_para(doc, "- xxx", "Normal", _)  # output
if p: set_para_text(p, "- Lista componenti da manutenere comunicata verbalmente/email al supervisore")
_, p = find_para(doc, "- xxx", "Normal", _)  # pain point
if p: set_para_text(p, "- Analisi manuale soggetta a errori; comunicazione informale senza tracciabilità; tempo analisi ~4h/ciclo")
_, p = find_para(doc, "- xxx", "Normal", _)  # HITL
if p: set_para_text(p, "N/A — processo manuale AS-IS")

# Card 3
_, p = find_para(doc, "Card 3: xxx")
if p: set_para_text(p, "Card 3: Pianificazione e Coordinamento Interventi di Manutenzione")

# Remaining xxx for card 3 inputs (two)
idx3 = _
c3_xxx = []
for i, para in enumerate(doc.paragraphs):
    if i <= idx3:
        continue
    if para.style.name == "Heading 3":
        break
    if para.text.strip() == "- xxx":
        c3_xxx.append((i, para))
if len(c3_xxx) >= 1: set_para_text(c3_xxx[0][1], "- Lista componenti da manutenere ricevuta dal tecnico d'ufficio")
if len(c3_xxx) >= 2: set_para_text(c3_xxx[1][1], "- Disponibilità fornitori, previsioni meteo, stato ricambi (verifica manuale)")

_, p = find_para(doc, "1. xxx", "Normal", idx3)  # attività
if p: set_para_lines(p, [
    "1. Supervisore raccoglie tutte le informazioni necessarie (fornitori, meteo, ricambi)",
    "2. Pianificazione manuale integrando vincoli multipli (esperienza tacita)",
    "3. Coordinamento logistica (accessi, trasporti, dotazioni sicurezza)",
    "4. Emissione ordini di lavoro e comunicazione ai tecnici/fornitori"
])
_, p = find_para(doc, "- xxx", "Normal", _)  # output
if p: set_para_text(p, "- Piano manutenzione e ordini di lavoro emessi (su carta o foglio Excel)")
_, p = find_para(doc, "- xxx", "Normal", _)  # pain point
if p: set_para_text(p, "- Alta complessità di pianificazione manuale multi-vincolo; dipendenza da conoscenza tacita del supervisore; nessuna ottimizzazione economica formale")
_, p = find_para(doc, "- xxx", "Normal", _)  # HITL
if p: set_para_text(p, "N/A — processo manuale AS-IS")

# ── PAIN POINT GENERALI ───────────────────────────────────────────────────────
_, p = find_para(doc, "1. Effort FTE elevato")
if p: set_para_text(p, "1. Effort FTE elevato: analisi manuale delle dashboard richiede personale dedicato giornalmente; stimato ~1 FTE/impianto")
_, p = find_para(doc, "2. Tempi di attraversamento lunghi")
if p: set_para_text(p, "2. Tempi di attraversamento lunghi: dalla rilevazione dell'anomalia all'intervento trascorrono mediamente 10-15 giorni (MTTR elevato)")
_, p = find_para(doc, "3. Colli di bottiglia")
if p: set_para_text(p, "3. Colli di bottiglia: la pianificazione è centralizzata sul supervisore che gestisce manualmente tutti i vincoli; singolo punto di failure")
_, p = find_para(doc, "4. Qualità non uniforme:")
if p: set_para_text(p, "4. Qualità non uniforme:")
_, p = find_para(doc, "   - xxx", "Normal")
if p: set_para_text(p, "   - La qualità dell'analisi dipende dall'esperienza del tecnico d'ufficio; variabilità nelle decisioni di prioritizzazione (~70% accuracy stimata)")
_, p = find_para(doc, "5. Scarsa scalabilità:")
if p: set_para_text(p, "5. Scarsa scalabilità:")
_, p = find_para(doc, "   - xxx", "Normal")
if p: set_para_text(p, "   - L'approccio manuale non scala con l'aumento del numero di impianti/turbine monitorati; ogni nuovo impianto richiede FTE aggiuntivi")

# Pain Point per Sistema
pp_systems = [
    ("SAS", "Capacità predittive avanzate sottoutilizzate; output non integrabili automaticamente con sistemi downstream"),
    ("Iceberg", "Dati eventi predittivi non sfruttati automaticamente; accesso solo tramite consultazione manuale"),
    ("POM", "Dati di vibrazione non correlati automaticamente con le decisioni di manutenzione; senza contesto storico"),
    ("Dashboard BI", "Strumento puramente passivo; nessuna automazione notifiche; analisi completamente manuale"),
]
pp_idx = None
_, p = find_para(doc, "Pain Point per Documenti Specifici")
if _:
    pp_idx = _
    count = 0
    for i, para in enumerate(doc.paragraphs):
        if i <= pp_idx:
            continue
        if para.style.name.startswith("Heading"):
            break
        if para.text.strip() in ("xxx:", "xxx") and ":" not in para.text.replace("xxx:", ""):
            if count < len(pp_systems):
                set_para_text(para, pp_systems[count][0] + ":")
                count += 1
        elif para.text.strip() == "- xxx":
            idx_s = count - 1
            if 0 <= idx_s < len(pp_systems):
                set_para_text(para, "- " + pp_systems[idx_s][1])

# ── 2.5 CONTROLLI SPECIFICI ───────────────────────────────────────────────────
verifiche = [
    ("Sistema SAS",
     ["Score predittivo componenti (gearbox, cuscinetti, generatore)", "Completezza serie temporale input (vibrazione, temperatura, RPM) – minimo 30 giorni rolling", "Soglie di allerta configurate per tipologia componente"],
     ["Score calcolato su serie temporali complete senza gap > 24h", "Modelli predittivi aggiornati con dati degli ultimi 90 giorni", "Alert generato solo se score > soglia configurata per almeno 3 giorni consecutivi"]
    ),
    ("Sistema Iceberg",
     ["ID evento e ID turbina/componente correttamente valorizzati", "Timestamp evento e livello severità (1-4) presenti", "Classificazione evento secondo tassonomia standard EGP"],
     ["Evento correttamente associato a turbina e componente in anagrafica impianti", "Severity > 2 per trigger processo manutentivo", "Evento non già gestito (stato ≠ 'chiuso' o 'in lavorazione')"]
    ),
    ("Sistema POM",
     ["ID sensore e ID turbina correttamente associati", "Valore vibrazione entro range fisico plausibile (m/s²)", "Timestamp acquisizione e frequenza di campionamento rispettata"],
     ["Segnale non in stato di errore o flat-line (qualità segnale > 80%)", "Frequenza di campionamento conforme a specifica sensore", "Dati trasmessi a SAS entro la finestra temporale prevista (< 1h)"]
    ),
    ("Dashboard BI",
     ["Data ultimo aggiornamento dati (< 24h rispetto al momento di consultazione)", "Corrispondenza tra alert visualizzati e eventi presenti in SAS/Iceberg"],
     ["Dashboard sincronizzata con sistemi sorgente; nessun dato mancante per impianti in scope", "Tutti gli alert con severity ≥ 2 visibili senza filtri aggiuntivi", "Storico alert degli ultimi 90 giorni disponibile per consultazione"]
    ),
]

_, h25 = find_para(doc, "2.5 Controlli Specifici", "Heading 3")
if h25 is None:
    _, h25 = find_para(doc, "2.5 Controlli", "Heading 3")

# Iterate and fill sequentially
_, h25 = find_para(doc, "2.5 Controlli")
if _ is not None:
    vi = 0  # which system
    fi = 0  # which field
    ci = 0  # which condition
    mode = "sys"  # sys | fields | cond
    sys_done = [False] * len(verifiche)

    for i, para in enumerate(doc.paragraphs):
        if i <= _:
            continue
        if para.style.name.startswith("Heading 2") or para.style.name.startswith("Heading 1"):
            break
        txt = para.text.strip()

        if mode == "sys" and txt in ("xxx", "xxx:"):
            if vi < len(verifiche):
                set_para_text(para, verifiche[vi][0])
                mode = "fields_header"
        elif mode == "fields_header" and txt == "Campi da Verificare:":
            mode = "fields"
            fi = 0
        elif mode == "fields" and txt == "- xxx":
            if fi < len(verifiche[vi][1]):
                set_para_text(para, "- " + verifiche[vi][1][fi])
                fi += 1
        elif mode == "fields" and txt == "Condizioni di Validità:":
            mode = "cond"
            ci = 0
        elif mode == "cond" and txt == "- xxx":
            if ci < len(verifiche[vi][2]):
                set_para_text(para, "- " + verifiche[vi][2][ci])
                ci += 1
        elif mode == "cond" and txt in ("xxx", "xxx:"):
            vi += 1
            if vi < len(verifiche):
                set_para_text(para, verifiche[vi][0])
                mode = "fields_header"

# ── TABLE 4 – DATA MAPPING ────────────────────────────────────────────────────
t4 = doc.tables[4]
data_mapping = [
    ("Score predittivo gearbox (0-1)", "SAS", "Dashboard BI / Alarm Analysis Agent", "Float JSON", "Aggiornamento giornaliero; trigger alert se > 0.7"),
    ("Alert vibrazione anomala", "POM", "SAS", "JSON stream", "Trigger analisi predittiva real-time"),
    ("Evento predittivo classificato", "Iceberg", "Dashboard BI / Alarm Analysis Agent", "JSON REST", "Severity 1-4; include ID turbina e componente"),
    ("Serie temporale vibrazione (FFT)", "POM", "SAS / Knowledge Base", "CSV / Parquet", "Finestra temporale 30 giorni rolling"),
    ("Score predittivo cuscinetti", "SAS", "Dashboard BI / Planner Agent", "Float JSON", "Aggiornamento giornaliero"),
    ("Previsione produzione eolica", "Sistema Meteo Esterno", "Planning Agent", "CSV / API REST", "Orizzonte 7 giorni; utilizzato per pianificazione finestre manutenzione"),
    ("Disponibilità fornitori manutenzione", "Sistema Gestione Fornitori", "Planning Agent", "JSON / Manuale", "Attuale: raccolta telefonica manuale [DA DEFINIRE API]"),
    ("Storico interventi manutenzione", "ERP / Archivio cartaceo", "Knowledge Base (RAG)", "PDF / Word / Excel", "Digitalizzazione e indicizzazione necessaria"),
    ("Documentazione tecnica turbine", "Archivio documentale EGP", "Knowledge Base (RAG)", "PDF / Word", "Manuali OEM, procedure tecniche"),
    ("Dati RPM e temperatura generatore", "POM", "SAS / Knowledge Base", "JSON stream", "Indicatori complementari vibrazione"),
    ("Ordini di lavoro generati", "Planning Agent (TO-BE)", "ERP", "JSON / PDF", "Workflow approvazione HITL prima dell'emissione"),
    ("Verbale post-intervento", "Tecnico (Mobile App TO-BE)", "Knowledge Base / Dashboard", "JSON / PDF", "Feedback loop per aggiornamento modelli"),
]
for i, row_data in enumerate(data_mapping):
    if i + 1 < len(t4.rows):
        fill_row(t4, i + 1, list(row_data))

# ── TABLE 5 – ARCHITETTURA FUNZIONALE TO-BE ───────────────────────────────────
t5 = doc.tables[5]
arch_rows = [
    ("Orchestrator Agent", "Coordina il flusso end-to-end: raccoglie output degli agenti specializzati, consolida le analisi e gestisce l'interfaccia HITL con il supervisore", "LLM multi-agent orchestration (LangGraph / AutoGen)"),
    ("Alarm Analysis Agent", "Analizza gli alert predittivi da SAS/Iceberg, classifica per priorità e urgenza, correla con storico anomalie analoghe", "ML classification + LLM reasoning (GPT-4o / Mistral)"),
    ("Planning Agent", "Ottimizza il piano di manutenzione considerando vincoli multi-dimensionali: economici (prezzo energia, costo impresa), meteo, disponibilità risorse e ricambi", "Constraint optimization + LLM (function calling)"),
    ("Knowledge Agent", "Recupera e sintetizza informazioni tecniche da documentazione strutturata e non strutturata (manuali OEM, storici interventi, procedure)", "RAG + Vector DB (Azure AI Search / Qdrant) + LLM"),
    ("Spare Parts Agent", "Verifica disponibilità ricambi e fornitori alternativi; propone ordini preventivi per componenti critici a rischio", "DB query + LLM (function calling verso ERP)"),
    ("ETL Pipeline", "Acquisisce e normalizza dati da SAS, Iceberg e POM; alimenta gli agenti AI con dati strutturati e aggiornati", "Azure Data Factory / Apache Spark / Delta Lake"),
    ("Dashboard HITL (TO-BE)", "Interfaccia supervisore/tecnico per revisione report AI, approvazione piani manutenzione, KPI monitoring", "Power BI Embedded / React Web App / MS Teams Bot"),
    ("HITL Validation Module", "Gestisce il workflow di approvazione umana prima dell'esecuzione di azioni critiche (emissione OdL, comunicazione fornitori)", "Workflow engine (Power Automate / custom)"),
]
for i, row_data in enumerate(arch_rows):
    if i + 1 < len(t5.rows):
        fill_row(t5, i + 1, list(row_data))

# ── PUNTI APERTI 3.1b ─────────────────────────────────────────────────────────
_, p = find_para(doc, "- xxx: esposizione API interna/esterna [DA DEFINIRE]")
if p: set_para_text(p, "- SAS: modalità accesso API/esportazione dati predittivi verso Alarm Analysis Agent e ETL Pipeline [DA DEFINIRE con IT Owner SAS]")
_, p = find_para(doc, "- xxx: API disponibile? [DA DEFINIRE]")
if p: set_para_text(p, "- Iceberg: disponibilità API eventi per consumo real-time o near-real-time da parte dei agenti AI [DA DEFINIRE]")
_, p = find_para(doc, "- xxx: modalità lettura/scrittura [DA DEFINIRE]")
if p: set_para_text(p, "- POM: accesso dati vibrazione in streaming o batch verso pipeline ETL; diritti di lettura da sistema SCADA [DA DEFINIRE]")
_, p = find_para(doc, "- xxx: accesso API [DA DEFINIRE]")
if p: set_para_text(p, "- Sistema Gestione Fornitori / ERP Ricambi: accesso dati disponibilità e integrazione con Planning Agent e Spare Parts Agent [DA DEFINIRE]")
_, p = find_para(doc, "I seguenti accessi/API sono [DA DEFINIRE]")
if p: set_para_text(p, "I seguenti accessi/API sono [DA DEFINIRE] con IT Owner prima del kick-off del progetto: SAS, Iceberg, POM, Sistema Gestione Fornitori, ERP Ricambi.")

# ── TO-BE SOTTO-PROCESSI LABEL ────────────────────────────────────────────────
_, p = find_para(doc, "Sotto-processo A: xxx (TO-BE)")
if p: set_para_text(p, "Sotto-processo A: Acquisizione Automatizzata Dati, Analisi e Classificazione Allarmi (TO-BE)")
_, p = find_para(doc, "Sotto-processo B: xxx (TO-BE)")
if p: set_para_text(p, "Sotto-processo B: Pianificazione Ottimizzata Interventi di Manutenzione (TO-BE)")
_, p = find_para(doc, "Sotto-processo C: xxx (TO-BE)")
if p: set_para_text(p, "Sotto-processo C: Esecuzione, Monitoraggio e Feedback per Apprendimento Continuo (TO-BE)")

# ── TABLE 6 – TO-BE Sotto-processo A ──────────────────────────────────────────
t6 = doc.tables[6]
rows_tobe_a = [
    ("A1", "ETL Pipeline acquisisce dati da SAS, Iceberg e POM e li normalizza in un dataset unificato", "ETL Pipeline (automatico)", "Dati sensori raw (POM), score predittivi (SAS), eventi (Iceberg)", "Dataset normalizzato e strutturato", "SAS, Iceberg, POM, Azure Data Factory", "ETL automatico con regole di qualità dati"),
    ("A2", "Alarm Analysis Agent analizza il dataset unificato e classifica gli allarmi per priorità e urgenza", "Alarm Analysis Agent", "Dataset normalizzato", "Lista allarmi classificati (priorità 1-4, urgenza, componente)", "SAS, Iceberg", "ML classification + LLM reasoning"),
    ("A3", "Knowledge Agent interroga la KB tecnica per storico anomalie analoghe e procedure di riferimento", "Knowledge Agent", "ID componente, tipologia allarme, impianto", "Scheda tecnica componente, storico interventi analoghi, procedure OEM", "Knowledge Base (RAG)", "RAG + vector search su documentazione tecnica"),
    ("A4", "Alarm Analysis Agent correla allarmi con storico KB e arricchisce l'analisi con contesto tecnico", "Alarm Analysis Agent", "Allarmi classificati + output Knowledge Agent", "Analisi arricchita con contesto e raccomandazione preliminare", "Tutti", "LLM reasoning + contesto RAG"),
    ("A5", "Orchestrator Agent consolida le analisi e genera un report strutturato per il tecnico d'ufficio", "Orchestrator Agent", "Analisi arricchita da tutti gli agenti", "Report analitico completo (componenti, priorità, raccomandazioni)", "Tutti, Dashboard HITL", "LLM synthesis e formatting report"),
    ("A6", "HITL – Tecnico d'ufficio revisiona il report AI, valida o corregge le priorità assegnate", "Tecnico d'ufficio", "Report analitico AI", "Report validato con eventuali correzioni", "Dashboard HITL, Teams Bot", "Assist – Human decides e approva"),
    ("A7", "Report validato inviato a Planning Agent per avvio fase di pianificazione", "Orchestrator Agent", "Report validato", "Trigger avvio Sotto-processo B", "Orchestrator, Planning Agent", "Hand-off automatico tra agenti"),
]
for i, row_data in enumerate(rows_tobe_a):
    if i + 1 < len(t6.rows):
        fill_row(t6, i + 1, list(row_data))

# ── TABLE 7 – TO-BE Sotto-processo B ──────────────────────────────────────────
t7 = doc.tables[7]
rows_tobe_b = [
    ("B1", "Planning Agent riceve la lista interventi validati e avvia la raccolta dei vincoli di pianificazione", "Planning Agent", "Lista interventi prioritizzati (validata HITL)", "Raccolta vincoli avviata", "Planning Agent, API sistemi esterni", "Automatic constraint collection via function calling"),
    ("B2", "Planning Agent interroga il sistema meteo per previsioni dei prossimi 14 giorni sugli impianti target", "Planning Agent", "Lista impianti interessati", "Previsioni meteo 14 giorni per finestre manutenzione", "API meteo (es. ECMWF/OpenWeather)", "API call automatica + analisi LLM"),
    ("B3", "Spare Parts Agent verifica la disponibilità dei ricambi necessari per i componenti da sostituire", "Spare Parts Agent", "Lista componenti da sostituire con codici articolo", "Disponibilità ricambi, lead time, fornitori alternativi", "ERP Ricambi, Sistema Gestione Fornitori", "DB query + LLM per analisi disponibilità"),
    ("B4", "Planning Agent raccoglie la disponibilità dei fornitori e delle imprese appaltatrici", "Planning Agent", "Lista interventi e impianti", "Slot disponibilità fornitori principali e alternativi", "Sistema Gestione Fornitori [DA DEFINIRE API]", "API call o form digitale HITL"),
    ("B5", "Planning Agent integra tutti i vincoli e genera il piano di manutenzione ottimizzato", "Planning Agent", "Interventi, meteo, ricambi, fornitori, vincoli economici", "Piano di manutenzione ottimizzato con scheduling dettagliato", "Planning Agent (constraint optimizer)", "Multi-constraint optimization + LLM"),
    ("B6", "Orchestrator Agent presenta il piano al supervisore tramite Dashboard HITL per revisione e approvazione", "Orchestrator Agent", "Piano ottimizzato", "Piano presentato su dashboard con motivazioni AI", "Dashboard HITL, Teams Bot", "LLM explanation generation"),
    ("B7", "HITL – Supervisore impianto revisiona il piano AI, può modificarlo e infine lo approva", "Supervisore impianto", "Piano ottimizzato AI con spiegazioni", "Piano approvato (o modificato e approvato)", "Dashboard HITL", "Assist – Human decides e approva"),
    ("B8", "Piano approvato triggera la generazione degli ordini di lavoro e notifiche ai fornitori", "Orchestrator Agent", "Piano approvato", "OdL draft generati, notifiche bozza pronte", "ERP, Sistema Fornitori", "Automazione generazione OdL"),
]
for i, row_data in enumerate(rows_tobe_b):
    if i + 1 < len(t7.rows):
        fill_row(t7, i + 1, list(row_data))

# ── TABLE 8 – TO-BE Sotto-processo C ──────────────────────────────────────────
t8 = doc.tables[8]
rows_tobe_c = [
    ("C1", "Sistema notifica automaticamente fornitori e tecnici degli interventi pianificati (email/Teams)", "Orchestrator Agent / Sistema", "OdL approvati", "Notifiche inviate a fornitori e tecnici", "ERP, Email/Teams", "Notifiche automatiche post-approvazione HITL"),
    ("C2", "Esecuzione fisica degli interventi di manutenzione da parte di tecnici/fornitori", "Tecnico / Fornitore", "Ordini di lavoro", "Intervento eseguito", "On-field (Mobile App)", "N/A – attività umana fisica"),
    ("C3", "Tecnico compila il verbale post-intervento tramite Mobile App (strutturato)", "Tecnico manutentore", "Intervento completato", "Verbale digitale strutturato (esito, durata, anomalie riscontrate)", "Mobile App TO-BE", "Form guidato AI con suggerimenti compilazione"),
    ("C4", "Knowledge Agent acquisisce il verbale e aggiorna la Knowledge Base con nuove evidenze", "Knowledge Agent", "Verbale post-intervento strutturato", "KB aggiornata con nuovo case di manutenzione", "Knowledge Base (RAG)", "Document ingestion + embedding update"),
    ("C5", "Orchestrator Agent aggiorna lo stato degli interventi e calcola i KPI di manutenzione", "Orchestrator Agent", "Verbali completati, dati post-intervento impianto", "KPI aggiornati (MTTR, accuracy alarm, efficacia intervento)", "Dashboard HITL, SAS (feedback)", "KPI calculation + dashboard update"),
    ("C6", "Feedback loop: agenti migliorano i modelli predittivi con i dati di esito degli interventi", "Alarm Analysis Agent / SAS", "KPI e esiti interventi", "Modelli predittivi aggiornati, threshold calibrate", "SAS, Alarm Analysis Agent", "Continual learning / model retraining workflow"),
    ("C7", "Reporting periodico aggregato a management: performance impianti, trend guasti, efficienza manutenzione", "Orchestrator Agent", "Dati KPI aggregati (mensile)", "Report management con insights e raccomandazioni strategiche", "Dashboard HITL, Power BI", "LLM-generated insights da dati aggregati"),
]
for i, row_data in enumerate(rows_tobe_c):
    if i + 1 < len(t8.rows):
        fill_row(t8, i + 1, list(row_data))

# ── PROCESS CARDS TO-BE ───────────────────────────────────────────────────────
# Card 1 TO-BE
_, p = find_para(doc, "Card 1: xxx (TO-BE)")
if p: set_para_text(p, "Card 1: Acquisizione Dati e Classificazione Allarmi (TO-BE)")
_, p = find_para(doc, "- xxx", "Normal", _)
if p: set_para_text(p, "- Dati sensori (POM); Alert predittivi e score rischio (SAS, Iceberg); Documentazione tecnica (Knowledge Base)")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p: set_para_lines(p, [
    "1. ETL Pipeline normalizza automaticamente dati da SAS, Iceberg e POM",
    "2. Alarm Analysis Agent classifica gli allarmi per priorità (1-4), urgenza e tipologia componente",
    "3. Knowledge Agent recupera storico anomalie analoghe e procedure di riferimento",
    "4. Orchestrator Agent genera report consolidato con raccomandazioni",
    "5. HITL: Tecnico d'ufficio valida e approva il report tramite Dashboard"
])
_, p = find_para(doc, "- xxx", "Normal", _)
if p: set_para_text(p, "- Report analitico validato con lista componenti a rischio, priorità e raccomandazioni di intervento")
_, p = find_para(doc, "- xxx", "Normal", _)
if p: set_para_text(p, "- SAS, Iceberg, POM, Knowledge Base, Dashboard HITL, ETL Pipeline")

# Card 2 TO-BE
_, p = find_para(doc, "Card 2: xxx (TO-BE)")
if p: set_para_text(p, "Card 2: Pianificazione Ottimizzata Interventi (TO-BE)")
_, p = find_para(doc, "- xxx", "Normal", _)
if p: set_para_text(p, "- Lista interventi validati; Previsioni meteo 14 giorni; Disponibilità fornitori e ricambi; Vincoli economici (prezzo energia)")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p: set_para_lines(p, [
    "1. Planning Agent raccoglie automaticamente vincoli (meteo, fornitori, ricambi, costi)",
    "2. Spare Parts Agent verifica disponibilità ricambi e propone fornitori alternativi se necessario",
    "3. Planning Agent genera piano ottimizzato multi-vincolo",
    "4. HITL: Supervisore revisiona, eventualmente modifica e approva il piano"
])
_, p = find_para(doc, "- xxx", "Normal", _)
if p: set_para_text(p, "- Piano manutenzione ottimizzato approvato; OdL draft generati per emissione")
_, p = find_para(doc, "- xxx", "Normal", _)
if p: set_para_text(p, "- Planning Agent, Spare Parts Agent, Sistema Meteo, Gestione Fornitori, ERP Ricambi, Dashboard HITL")

# Card 3 TO-BE
_, p = find_para(doc, "Card 3: xxx (TO-BE)")
if p: set_para_text(p, "Card 3: Esecuzione, Monitoraggio e Feedback (TO-BE)")
_, p = find_para(doc, "- xxx", "Normal", _)
if p: set_para_text(p, "- Piano manutenzione approvato; Verbali post-intervento dei tecnici (strutturati via Mobile App)")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p: set_para_lines(p, [
    "1. Sistema notifica automaticamente fornitori e tecnici (post-approvazione HITL)",
    "2. Tecnici eseguono interventi e compilano verbale digitale via Mobile App",
    "3. Knowledge Agent acquisisce verbale e aggiorna KB con nuovo case",
    "4. Orchestrator calcola KPI (MTTR, accuracy alarm) e aggiorna dashboard",
    "5. Feedback loop: agenti calibrano modelli predittivi con esiti reali"
])
_, p = find_para(doc, "- xxx", "Normal", _)
if p: set_para_text(p, "- KB aggiornata; KPI manutenzione calcolati e disponibili su dashboard; modelli predittivi migliorati")
_, p = find_para(doc, "- xxx", "Normal", _)
if p: set_para_text(p, "- Knowledge Base, Orchestrator Agent, ERP, Mobile App, Dashboard HITL, SAS (feedback)")

# ── 3.4 COSA NON FA L'AI ──────────────────────────────────────────────────────
_, p = find_para(doc, "3.4 Cosa NON fa l'AI", "Heading 3")
_, p_cosa = find_para(doc, "xxx", "Normal", _)
if p_cosa:
    set_para_lines(p_cosa, [
        "❌ Non esegue fisicamente gli interventi di manutenzione (rimane responsabilità di tecnici/fornitori)",
        "❌ Non emette autonomamente ordini di lavoro senza esplicita approvazione del supervisore umano (HITL obbligatorio)",
        "❌ Non sostituisce il giudizio tecnico del supervisore per decisioni critiche di sicurezza sul campo",
        "❌ Non accede direttamente ai sistemi di controllo SCADA degli impianti wind né può modificarne parametri",
        "❌ Non prende decisioni autonome su curtailment, fermata impianti o modifica set-point di produzione",
        "❌ Non comunica autonomamente con fornitori o imprese appaltatrici senza validazione umana",
        "❌ Non modifica i modelli predittivi SAS senza revisione e approvazione del team tecnico",
        "❌ Non sostituisce le competenze specialistiche dei tecnici manutentori nella diagnosi fisica di guasti"
    ])

# ── 3.5 COMPONENTE AI ─────────────────────────────────────────────────────────
_, p = find_para(doc, "3.5 Componente ai", "Heading 3")
_, p_ai = find_para(doc, "Xxx", "Normal", _)
if p_ai:
    set_para_lines(p_ai, [
        "Architettura multi-agent AI composta da 4 agenti specializzati coordinati da un Orchestrator:",
        "",
        "(1) Orchestrator Agent – coordina il flusso end-to-end, consolida gli output degli agenti specializzati "
        "e gestisce l'interfaccia HITL con tecnici e supervisori.",
        "Pattern tassonomico: Orchestrator / Coordinator Agent",
        "",
        "(2) Alarm Analysis Agent – analizza gli alert predittivi da SAS/Iceberg, classifica per priorità "
        "e urgenza, correla con storico anomalie da Knowledge Base.",
        "Pattern tassonomico: Analyzer / Classifier Agent",
        "",
        "(3) Planning Agent – ottimizza il piano di manutenzione considerando vincoli multi-dimensionali "
        "(economici, meteo, disponibilità risorse, ricambi). Include Spare Parts sub-agent.",
        "Pattern tassonomico: Planner / Optimizer Agent",
        "",
        "(4) Knowledge Agent – recupera e sintetizza informazioni tecniche da documentazione strutturata "
        "e non strutturata tramite RAG; aggiorna la KB con feedback post-intervento.",
        "Pattern tassonomico: Retriever / RAG Agent"
    ])

# Fix second xxx in AI section
_, p_ai2 = find_para(doc, "**in questa sezione va indicata la componente AI", "Normal", _)
if p_ai2:
    set_para_text(p_ai2, "")

# ── 4. DELTA AS-IS vs TO-BE ───────────────────────────────────────────────────
_, p = find_para(doc, "Operatività")
if p: set_para_text(p, "Operatività – Confronto AS-IS / TO-BE")
_, p = find_para(doc, "AS-IS:", "Normal")
if p: set_para_text(p, "AS-IS:")
_, p = find_para(doc, "- xxx", "Normal", _)
if p: set_para_lines(p, [
    "- Analisi manuale giornaliera delle dashboard da parte di tecnici d'ufficio (~4h/ciclo)",
    "- Comunicazione informale (telefono/email) tra ufficio e supervisore d'impianto",
    "- Pianificazione manuale interventi senza ottimizzazione formale multi-vincolo",
    "- Knowledge tecnica non strutturata; dipendenza da esperienza individuale del supervisore"
])

# ── 4.1 IMPATTI OPERATIVI ─────────────────────────────────────────────────────
_, h41 = find_para(doc, "4.1 Impatti Operativi", "Heading 3")
_, p = find_para(doc, "xxx", "Normal", _)
if p: set_para_lines(p, [
    "- Riduzione stimata ~60% del tempo di analisi allarmi (da ~4h a ~1.5h per ciclo di monitoraggio)",
    "- Riduzione MTTR stimata del 30-40%: da ~15 giorni a ~9 giorni (dalla previsione all'intervento completato)",
    "- Liberazione di ~1 FTE/impianto dedicato all'analisi manuale delle dashboard",
    "- Aumento del numero di allarmi gestiti per operatore: da ~20 a ~80 allarmi/settimana/FTE",
    "- Miglioramento qualità decisioni pianificazione grazie a ottimizzazione multi-vincolo automatica"
])

# ── 4.2 INVARIANTI ────────────────────────────────────────────────────────────
_, h42 = find_para(doc, "4.2 Invarianti", "Heading 3")
inv_count = 0
for i, para in enumerate(doc.paragraphs):
    if i <= _:
        continue
    if para.style.name.startswith("Heading"):
        break
    if para.text.strip() == "- xxx":
        if inv_count == 0:
            set_para_text(para, "- La decisione finale di approvazione degli interventi rimane in carico al supervisore umano (HITL obbligatorio)")
        elif inv_count == 1:
            set_para_text(para, "- I sistemi SAS, Iceberg e POM rimangono le fonti primarie di dati (nessuna sostituzione di sistemi esistenti)")
        elif inv_count == 2:
            set_para_text(para, "- I tecnici eseguono fisicamente tutti gli interventi di manutenzione sul campo")
        elif inv_count == 3:
            set_para_text(para, "- Le procedure di sicurezza e i protocolli normativi vigenti per l'esercizio di impianti eolici rimangono invariati")
        inv_count += 1

# ── 4.3 NUOVI REQUISITI ABILITANTI ───────────────────────────────────────────
_, h43 = find_para(doc, "4.3 Nuovi Requisiti", "Heading 3")
req_count = 0
for i, para in enumerate(doc.paragraphs):
    if i <= _:
        continue
    if para.style.name == "Heading 3" or para.style.name == "Heading 2":
        break
    if para.text.strip() == "- xxx":
        if req_count == 0:
            set_para_text(para, "- Definizione e implementazione API/interfacce verso SAS, Iceberg e POM per alimentazione pipeline ETL")
        elif req_count == 1:
            set_para_text(para, "- Creazione della Knowledge Base tecnica: digitalizzazione e indicizzazione di manuali OEM e storici interventi")
        elif req_count == 2:
            set_para_text(para, "- Sviluppo interfaccia HITL (Dashboard TO-BE / Teams Bot) per supervisori e tecnici d'ufficio")
        elif req_count == 3:
            set_para_text(para, "- Formazione del personale sull'utilizzo del sistema AI e sui nuovi flussi operativi con HITL")
        req_count += 1
    if para.text.strip() == "TO-BE:":
        break

# ── TO-BE bullets ─────────────────────────────────────────────────────────────
_, p = find_para(doc, "TO-BE:", "Normal")
if p: set_para_text(p, "TO-BE:")
_, p = find_para(doc, "- xxx", "Normal", _)
if p: set_para_lines(p, [
    "- Pipeline ETL automatizzata acquisisce e normalizza continuamente dati da SAS, Iceberg e POM",
    "- Agenti AI (Alarm Analysis + Knowledge) analizzano, classificano e prioritizzano gli allarmi in autonomia",
    "- Planning Agent genera piani manutenzione ottimizzati multi-vincolo in pochi minuti",
    "- Supervisore revisiona e approva tramite dashboard dedicata (HITL): decisione informata in < 30 min",
    "- Knowledge Base strutturata rende accessibile la conoscenza tecnica a tutti gli operatori"
])

# ── TABLE 9 – ROADMAP ─────────────────────────────────────────────────────────
t9 = doc.tables[9]
roadmap = [
    ("Fase 1 – Raccolta Requisiti e Accessi\n(M1–M2)", "Definizione architettura dettagliata; mappatura API verso SAS, Iceberg, POM; raccolta dati storici e documentazione tecnica da indicizzare", "Architettura tecnica validata; accordi IT per accessi API; dataset storico iniziale per KB", "2 mesi"),
    ("Fase 2 – Data Access e Knowledge Base\n(M2–M4)", "Implementazione pipeline ETL; creazione e indicizzazione Knowledge Base (manuali OEM, storici interventi); validazione qualità dati", "ETL funzionante su ambiente di sviluppo; KB popolata con ≥80% documentazione disponibile", "2 mesi"),
    ("Fase 3 – AI PoC\n(M3–M6)", "Sviluppo e test Alarm Analysis Agent e Planning Agent in sandbox; validazione accuracy classificazione allarmi su dati storici; prima versione Knowledge Agent", "PoC demo con accuracy > 85% su dataset storico; planning agent genera piani coerenti su casi test", "3 mesi"),
    ("Fase 4 – Integration e HITL Interface\n(M5–M8)", "Integrazione agenti con sistemi reali (SAS, Iceberg, POM); sviluppo Dashboard HITL e Teams Bot; test operativi con supervisori e tecnici d'ufficio", "Sistema integrato su ambiente pre-produzione; Dashboard HITL validata da 2-3 supervisori pilota", "3 mesi"),
    ("Fase 5 – Pilota e Validazione\n(M7–M9)", "Pilota operativo su 1-2 impianti selezionati; monitoraggio KPI (MTTR, accuracy, FTE); go/no-go per scale-up su tutti gli impianti wind", "Riduzione MTTR ≥ 20% verificata su impianti pilota; adozione sistema > 80% supervisori; report go/no-go scale-up", "3 mesi"),
    ("Fase 6 – Scale-up\n(M9–M12)", "Estensione graduale a tutti gli impianti wind in scope; formazione personale; messa a regime monitoraggio e supporto continuativo", "Sistema in produzione su tutti gli impianti; piano di miglioramento continuo definito", "3 mesi"),
]
for i, row_data in enumerate(roadmap):
    if i + 1 < len(t9.rows):
        fill_row(t9, i + 1, list(row_data))

# ── TABLE 10 – KPI QUANTITATIVI ───────────────────────────────────────────────
t10 = doc.tables[10]
kpi_quant = [
    ("MTTR Manutenzione Predittiva (giorni dalla previsione all'intervento completato)", "~15 giorni", "< 9 giorni (–40%)", "Tracking timestamp: da alert SAS a chiusura OdL su ERP"),
    ("Accuracy classificazione allarmi AI (% allarmi correttamente prioritizzati)", "~70% (analisi manuale)", "> 90%", "Confronto classificazione AI vs. esito reale dell'intervento (feedback loop)"),
    ("Volume allarmi gestiti per FTE (allarmi/settimana/operatore)", "~20 allarmi/settimana/FTE", "~80 allarmi/settimana/FTE (+300%)", "Conteggio allarmi processati su dashboard / FTE assegnati"),
    ("Tempo medio analisi singolo allarme", "~4 ore/allarme (analisi manuale)", "< 45 minuti (AI + HITL)", "Timestamp apertura → chiusura analisi su Dashboard HITL"),
    ("Tasso falsi positivi allarmi AI", "N/A (baseline manuale soggettiva)", "< 10%", "% interventi generati da AI risultati non necessari post-esecuzione"),
    ("Lead time pianificazione intervento", "~3 giorni (raccolta vincoli manuale)", "< 4 ore (AI-generated + HITL approval)", "Timestamp ricezione lista interventi → piano approvato"),
    ("Disponibilità ricambi al momento dell'intervento", "~75% (verifica manuale asincrona)", "> 95%", "Tracciamento OdL: interventi eseguiti senza blocchi ricambi / totale interventi"),
    ("Copertura Knowledge Base (% documentazione indicizzata)", "~0% (non strutturata)", "> 85%", "Ratio documenti indicizzati in KB / totale documenti tecnici disponibili"),
    ("Tempo compilazione verbale post-intervento", "~30 min (cartaceo)", "< 10 min (Mobile App)", "Timestamp apertura → invio verbale digitale tramite Mobile App"),
    ("Riduzione fermate impreviste (guasti non anticipati da predittiva)", "Baseline da storico EGP (n. fermate/anno)", "–25% fermate impreviste", "Conteggio fermate impreviste anno post-vs.-pre deployment AI"),
]
for i, row_data in enumerate(kpi_quant):
    if i + 1 < len(t10.rows):
        fill_row(t10, i + 1, list(row_data))

# ── TABLE 11 – KPI QUALITATIVI ────────────────────────────────────────────────
t11 = doc.tables[11]
kpi_qual = [
    ("Carico cognitivo operatori (self-reported)", "Riduzione carico analisi manuale; target: majorità operatori lo valuta ridotto (scala Likert, survey semestrale)", "Survey semestrale operatori d'ufficio e supervisori"),
    ("Qualità knowledge retrieval (% query KB risolte senza escalation)", "Target: > 80% delle query Knowledge Agent risolte autonomamente senza intervento manuale", "Tracking query KB: risolte autonomamente / totale query"),
    ("Adozione sistema AI da parte dei supervisori", "Target: > 90% supervisori utilizza regolarmente la Dashboard HITL entro 6 mesi dal go-live", "Log di accesso e utilizzo Dashboard HITL"),
    ("Qualità dei piani AI (% piani approvati senza modifiche dal supervisore)", "Target: > 70% dei piani generati dall'AI approvati senza modifiche sostanziali", "Tracking HITL: piani approvati as-is / totale piani"),
    ("Completezza verbali post-intervento digitali", "Target: > 95% verbali compilati in formato strutturato tramite Mobile App entro 2h dalla fine intervento", "Tracking invio verbali digitali / totale interventi eseguiti"),
    ("Soddisfazione utenti sistema AI (NPS interno)", "Target: NPS > 30 entro 12 mesi dal go-live del sistema", "Survey NPS trimestrale a tutti gli utenti del sistema"),
    ("Aggiornamento continuo modelli (frequenza retraining)", "Target: retraining automatico modelli con batch mensile di nuovi dati; nessuna degradazione accuracy > 5%", "Monitoring automatico accuracy con alerting su degradazione"),
    ("Tracciabilità decisioni AI (% decisioni con spiegazione leggibile)", "Target: 100% delle raccomandazioni AI presentate con motivazione comprensibile al supervisore", "Audit log Dashboard HITL: % decisioni con explanation presente"),
]
for i, row_data in enumerate(kpi_qual):
    if i + 1 < len(t11.rows):
        fill_row(t11, i + 1, list(row_data))

# ── SAVE ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Blueprint saved to:\n{OUTPUT}")
