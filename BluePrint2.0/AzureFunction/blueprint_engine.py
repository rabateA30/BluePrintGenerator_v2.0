"""
blueprint_engine.py — Core logic for BlueprintGenerator Azure Function
Ports the workflow from BlueprintGenerator.agent.md (Phases 2-5) to server-side Python.

Phase 2 — Template analysis (python-docx)
Phase 3 — Source content extraction (python-docx, PyMuPDF, python-pptx)
Phase 4 — Content synthesis via Azure OpenAI (gpt-4o)
Phase 5 — Document generation (python-docx) + upload to SharePoint
"""

import os
import copy
import io
import logging
import tempfile
import requests
from typing import Optional

# ── Third-party ────────────────────────────────────────────────────────────────
from docx import Document
from docx.oxml.ns import qn
import fitz          # PyMuPDF
from pptx import Presentation as PptxDoc
from openai import AzureOpenAI

# ── Configuration from environment ────────────────────────────────────────────
AZURE_OPENAI_ENDPOINT   = os.environ.get("AZURE_OPENAI_ENDPOINT", "")
AZURE_OPENAI_KEY        = os.environ.get("AZURE_OPENAI_KEY", "")
AZURE_OPENAI_DEPLOYMENT = os.environ.get("AZURE_OPENAI_DEPLOYMENT", "gpt-4o")

AZURE_TENANT_ID     = os.environ.get("AZURE_TENANT_ID", "")
AZURE_CLIENT_ID     = os.environ.get("AZURE_CLIENT_ID", "")
AZURE_CLIENT_SECRET = os.environ.get("AZURE_CLIENT_SECRET", "")
SHAREPOINT_SITE_URL = os.environ.get("SHAREPOINT_SITE_URL", "")

BLUEPRINT_TEMPLATE_URL = os.environ.get("BLUEPRINT_TEMPLATE_URL", "")
# If running locally with template available on disk:
BLUEPRINT_TEMPLATE_LOCAL = os.environ.get("BLUEPRINT_TEMPLATE_LOCAL", "")

# ── OpenAI client ──────────────────────────────────────────────────────────────
_oai_client: Optional[AzureOpenAI] = None

def _get_oai_client() -> AzureOpenAI:
    global _oai_client
    if _oai_client is None:
        _oai_client = AzureOpenAI(
            azure_endpoint=AZURE_OPENAI_ENDPOINT,
            api_key=AZURE_OPENAI_KEY,
            api_version="2024-12-01-preview",
        )
    return _oai_client


# ── Microsoft Graph helpers ────────────────────────────────────────────────────

def _get_graph_token() -> str:
    """Obtain an app-only access token for Microsoft Graph."""
    url = f"https://login.microsoftonline.com/{AZURE_TENANT_ID}/oauth2/v2.0/token"
    resp = requests.post(url, data={
        "grant_type":    "client_credentials",
        "client_id":     AZURE_CLIENT_ID,
        "client_secret": AZURE_CLIENT_SECRET,
        "scope":         "https://graph.microsoft.com/.default",
    }, timeout=30)
    resp.raise_for_status()
    return resp.json()["access_token"]


def _graph_get(path: str, token: str) -> dict:
    resp = requests.get(
        f"https://graph.microsoft.com/v1.0{path}",
        headers={"Authorization": f"Bearer {token}"},
        timeout=60,
    )
    resp.raise_for_status()
    return resp.json()


def _graph_download(path: str, token: str) -> bytes:
    resp = requests.get(
        f"https://graph.microsoft.com/v1.0{path}",
        headers={"Authorization": f"Bearer {token}"},
        timeout=120,
    )
    resp.raise_for_status()
    return resp.content


def _graph_upload(site_id: str, parent_item_id: str, filename: str,
                  content: bytes, token: str) -> str:
    """Upload a file to SharePoint and return its web URL."""
    resp = requests.put(
        f"https://graph.microsoft.com/v1.0/sites/{site_id}/drive/items/"
        f"{parent_item_id}:/{filename}:/content",
        headers={
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/octet-stream",
        },
        data=content,
        timeout=120,
    )
    resp.raise_for_status()
    return resp.json().get("webUrl", "")


# ── Public: list SharePoint folders ───────────────────────────────────────────

def list_sharepoint_folders() -> list[dict]:
    """Return list of project folders available on SharePoint for blueprint generation."""
    if not AZURE_CLIENT_ID:
        # Demo mode — return mock list
        return [
            {"name": "Global - CMR Virtual Assistant", "file_count": 4, "file_types": "PPTX, PDF"},
            {"name": "HSEQ Inspection ES", "file_count": 3, "file_types": "PPTX, DOCX"},
            {"name": "BuyAssist ITA", "file_count": 2, "file_types": "PDF, DOCX"},
        ]

    token = _get_graph_token()
    # Get SharePoint site id
    site_path = SHAREPOINT_SITE_URL.replace("https://", "").replace("/sites/", ":")
    site_info = _graph_get(f"/sites/{site_path}", token)
    site_id = site_info["id"]

    # List top-level folders in the drive root
    items = _graph_get(f"/sites/{site_id}/drive/root/children", token)
    folders = []
    for item in items.get("value", []):
        if "folder" in item:
            children = _graph_get(
                f"/sites/{site_id}/drive/items/{item['id']}/children", token
            )
            files = [c for c in children.get("value", []) if "file" in c]
            exts = set()
            for f in files:
                ext = f["name"].rsplit(".", 1)[-1].upper() if "." in f["name"] else ""
                if ext in ("PDF", "PPTX", "DOCX", "XLSX"):
                    exts.add(ext)
            folders.append({
                "name": item["name"],
                "file_count": len(files),
                "file_types": ", ".join(sorted(exts)) or "—",
            })
    return folders


# ── Phase 3: Document extraction ──────────────────────────────────────────────

def _extract_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _extract_pptx(content: bytes) -> str:
    prs = PptxDoc(io.BytesIO(content))
    parts = []
    for slide_num, slide in enumerate(prs.slides, 1):
        slide_texts = []
        for shape in slide.shapes:
            if hasattr(shape, "text") and shape.text.strip():
                slide_texts.append(shape.text.strip())
        if slide_texts:
            parts.append(f"[Slide {slide_num}]\n" + "\n".join(slide_texts))
    return "\n\n".join(parts)


def _extract_pdf(content: bytes) -> str:
    doc = fitz.open(stream=content, filetype="pdf")
    parts = []
    for page_num, page in enumerate(doc, 1):
        text = page.get_text().strip()
        if text:
            parts.append(f"[Page {page_num}]\n{text}")
    return "\n\n".join(parts)


def _extract_file(filename: str, content: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "docx":
        return _extract_docx(content)
    if ext == "pptx":
        return _extract_pptx(content)
    if ext == "pdf":
        return _extract_pdf(content)
    return ""  # Skip unsupported formats


# ── Phase 5: python-docx helpers ──────────────────────────────────────────────

def find_para(doc: Document, text_fragment: str, style: str = None, start: int = 0):
    """Return (index, para) for first paragraph containing text_fragment after start."""
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if text_fragment.lower() in p.text.lower():
            if style is None or p.style.name == style:
                return i, p
    return None, None


def set_para_text(para, new_text: str):
    """Replace paragraph text preserving first run's formatting."""
    if not para.runs:
        para.add_run(new_text)
        return
    rpr_copy = None
    first_rpr = para.runs[0]._r.find(qn("w:rPr"))
    if first_rpr is not None:
        rpr_copy = copy.deepcopy(first_rpr)
    for run in para.runs:
        run.text = ""
    para.runs[0].text = new_text
    if rpr_copy is not None:
        existing = para.runs[0]._r.find(qn("w:rPr"))
        if existing is not None:
            para.runs[0]._r.remove(existing)
        para.runs[0]._r.insert(0, rpr_copy)


def fill_cell(cell, text: str):
    """Replace table cell text preserving formatting."""
    for para in cell.paragraphs:
        if para.runs:
            set_para_text(para, "")
    if cell.paragraphs:
        set_para_text(cell.paragraphs[0], text)
    else:
        cell.add_paragraph(text)


# ── Phase 4: Content synthesis via Azure OpenAI ───────────────────────────────

BLUEPRINT_SYSTEM_PROMPT = """Sei un Blueprint Architect specializzato in Business Process Design e AI Augmentation per il gruppo Enel.
Compili documenti Blueprint conformi al framework Blueprint 1.0 Enel-GICT.
Regole:
- Usa SOLO informazioni presenti nei documenti sorgente forniti.
- Se mancano informazioni per una sezione, scrivi "[Da definire]".
- Non inventare dati, nomi, sistemi o processi.
- Usa linguaggio professionale, tecnico e conciso.
- Rispetta la struttura e la terminologia Enel."""


def _synthesize_section(section_name: str, section_purpose: str,
                         extracted_content: str, output_language: str) -> str:
    """Call Azure OpenAI to synthesize content for one blueprint section."""
    client = _get_oai_client()
    lang_instruction = {
        "Italian": "Rispondi in italiano.",
        "English": "Reply in English.",
        "Spanish": "Responde en español.",
    }.get(output_language, "Reply in English.")

    prompt = f"""{lang_instruction}

Sezione Blueprint da compilare: **{section_name}**
Scopo della sezione: {section_purpose}

Documenti sorgente del progetto:
---
{extracted_content[:12000]}
---

Genera il contenuto per questa sezione del Blueprint. Sii conciso e strutturato.
Se le informazioni non sono disponibili nei documenti, scrivi "[Da definire]".
"""
    response = client.chat.completions.create(
        model=AZURE_OPENAI_DEPLOYMENT,
        messages=[
            {"role": "system", "content": BLUEPRINT_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        max_tokens=1500,
        temperature=0.2,
    )
    return response.choices[0].message.content.strip()


# Section definitions — mapping template sections to synthesis instructions
SECTIONS = [
    ("0. Stakeholder & Partecipanti",
     "Tabella con ruoli: Business Owner, Process Owner, AI Owner, SME, Sponsor, IT System Owner. "
     "Identifica nomi e unità org dai documenti. Se non presenti, usa [Da definire]."),
    ("1.1 Scopo",
     "Descrizione chiara di COSA fa il processo e QUALI sistemi sono in scope (1-3 frasi)."),
    ("1.2 Finalità",
     "Lista di 3-4 obiettivi del processo: continuità, classificazione, attivazione attori, target MTTR."),
    ("1.3 Perimetro",
     "Due blocchi: IN SCOPE (3 voci) e OUT OF SCOPE (2 voci). Usa '- ' come prefisso per ogni voce."),
    ("1.4 Vincoli chiave",
     "Vincoli categorizzati in: Normativi (regolatori), Tecnici (integrazione/sistemi), Organizzativi (governance/training)."),
    ("2.1 Sistemi coinvolti AS-IS",
     "Tabella: Sistema | Ruolo nel processo | Tipo. Identifica tutti i sistemi/tool citati nei documenti."),
    ("2.2 Sequenza operativa AS-IS",
     "Flusso del processo AS-IS in sub-processi A/B/C con tabelle Step/Attività/Attore/Input/Output/Sistemi."),
    ("2.3 Process Cards AS-IS",
     "Per ogni macro-attività: Input, Attività (numerata), Output, Pain Point, HITL/Note (sempre 'N/A — processo manuale AS-IS' per AS-IS)."),
    ("2.4 Pain Points",
     "Pain points generali (5 voci: FTE, MTTR, bottleneck, qualità, scalabilità) + per-sistema (1 bullet per sistema)."),
    ("2.5 Verifiche Specifiche",
     "Checklist per sistema: Campi da Verificare + Condizioni di Validità per ogni sistema coinvolto."),
    ("3.1 Architettura Funzionale TO-BE",
     "Tabella: Componente | Funzione | Tecnologia. Includi: Diagnostic Engine, Pattern Recognition, Ticket Agent, Knowledge Base, Decision Tree, Dashboard, ETL."),
    ("3.2 Sequenza operativa TO-BE",
     "Flusso TO-BE con 7 colonne: Step, Attività, Attore, Input, Output, Sistemi, AI+HITL."),
    ("3.3 Process Cards TO-BE",
     "Per ogni macro-attività AI: Input, Attività AI (numerata), Output, Sistemi."),
    ("3.4 Cosa l'AI NON fa",
     "Lista di 6-8 elementi espliciti con simbolo negativo: no invio autonomo ticket, no scrittura SCADA, no comunicazioni ISO autonome, no decisioni di curtailment, ecc."),
    ("4. Delta AS-IS vs TO-BE",
     "Strutturato in: AS-IS bullets, 4.1 Impatti Operativi, 4.2 Invarianti (Non Cambia), 4.3 Nuovi Requisiti Abilitanti, TO-BE bullets."),
    ("5. Roadmap",
     "5 fasi M1-M12: Raccolta Requisiti, Data Access, AI PoC, Integration, Pilota — con obiettivi, output e durate."),
    ("6. KPI",
     "Quantitativi (4 KPI: MTTR, accuracy ticket, volume allarmi, tempo analisi — con baseline AS-IS e target). "
     "Qualitativi (4 KPI: carico cognitivo, accuracy AI, KB coverage, adoption)."),
]


# ── Main pipeline ──────────────────────────────────────────────────────────────

def run_blueprint_pipeline(source_folder: str, project_name: str,
                            output_language: str, output_filename: str) -> dict:
    """
    Full pipeline:
    1. Download source files from SharePoint
    2. Extract text content
    3. Synthesize sections via Azure OpenAI
    4. Fill Blueprint template
    5. Upload output to SharePoint
    6. Return result dict
    """
    logging.info(f"Pipeline start: project={project_name}, lang={output_language}")

    # ── Step 1: Get template ───────────────────────────────────────────────────
    if BLUEPRINT_TEMPLATE_LOCAL and os.path.exists(BLUEPRINT_TEMPLATE_LOCAL):
        with open(BLUEPRINT_TEMPLATE_LOCAL, "rb") as f:
            template_bytes = f.read()
    elif BLUEPRINT_TEMPLATE_URL:
        token = _get_graph_token()
        logging.info("Downloading template from Blob...")
        resp = requests.get(BLUEPRINT_TEMPLATE_URL, timeout=60)
        resp.raise_for_status()
        template_bytes = resp.content
    else:
        raise RuntimeError("No blueprint template configured. Set BLUEPRINT_TEMPLATE_LOCAL or BLUEPRINT_TEMPLATE_URL.")

    # ── Step 2: Download source files from SharePoint ─────────────────────────
    token = _get_graph_token()
    site_path = SHAREPOINT_SITE_URL.replace("https://", "").replace("/sites/", ":")
    site_info = _graph_get(f"/sites/{site_path}", token)
    site_id = site_info["id"]

    # Find the folder
    folder_items = _graph_get(f"/sites/{site_id}/drive/root/children", token)
    target_folder = None
    for item in folder_items.get("value", []):
        if item["name"].lower() == source_folder.lower() and "folder" in item:
            target_folder = item
            break

    if not target_folder:
        raise ValueError(f"Folder '{source_folder}' not found on SharePoint.")

    children = _graph_get(f"/sites/{site_id}/drive/items/{target_folder['id']}/children", token)
    source_files = [
        c for c in children.get("value", [])
        if "file" in c and c["name"].rsplit(".", 1)[-1].lower() in ("pdf", "pptx", "docx", "xlsx")
    ]

    if not source_files:
        raise ValueError(f"No supported files (PDF, PPTX, DOCX) found in folder '{source_folder}'.")

    # ── Step 3: Extract content ────────────────────────────────────────────────
    logging.info(f"Extracting content from {len(source_files)} files...")
    all_content_parts = []
    for f_item in source_files:
        logging.info(f"  Extracting: {f_item['name']}")
        file_bytes = _graph_download(
            f"/sites/{site_id}/drive/items/{f_item['id']}/content", token
        )
        extracted = _extract_file(f_item["name"], file_bytes)
        if extracted:
            all_content_parts.append(f"=== {f_item['name']} ===\n{extracted}")

    combined_content = "\n\n".join(all_content_parts)
    logging.info(f"Total extracted content: {len(combined_content)} chars")

    # ── Step 4: Synthesize all sections via Azure OpenAI ──────────────────────
    logging.info("Synthesizing sections via Azure OpenAI...")
    filled_sections: list[dict] = []
    open_questions: list[str] = []

    for section_name, section_purpose in SECTIONS:
        logging.info(f"  Synthesizing: {section_name}")
        content = _synthesize_section(section_name, section_purpose, combined_content, output_language)
        has_real_content = "[Da definire]" not in content or len(content) > 50
        filled_sections.append({
            "section": section_name,
            "content": content,
            "filled": has_real_content,
        })
        if not has_real_content or content.strip() == "[Da definire]":
            open_questions.append(section_name)

    # ── Step 5: Fill template ──────────────────────────────────────────────────
    logging.info("Filling Blueprint template...")
    doc = Document(io.BytesIO(template_bytes))

    # Title
    if doc.paragraphs:
        set_para_text(doc.paragraphs[0], f"Blueprint – {project_name}")

    # Fill each synthesized section by searching for its heading
    section_map = {s["section"]: s["content"] for s in filled_sections}
    for section_name, content in section_map.items():
        # Search for the section heading in the document
        short_name = section_name.split("—")[0].strip().split(".")[0].strip()
        _, para = find_para(doc, short_name)
        if para is None:
            continue
        # Find the next paragraph after the heading and fill it
        idx = doc.paragraphs.index(para)
        if idx + 1 < len(doc.paragraphs):
            set_para_text(doc.paragraphs[idx + 1], content)

    # Save to bytes
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_bytes = output_buffer.getvalue()

    # ── Step 6: Upload to SharePoint ──────────────────────────────────────────
    logging.info(f"Uploading '{output_filename}' to SharePoint...")
    output_url = _graph_upload(
        site_id=site_id,
        parent_item_id=target_folder["id"],
        filename=output_filename,
        content=output_bytes,
        token=token,
    )
    logging.info(f"Upload complete: {output_url}")

    return {
        "output_url": output_url,
        "sections_filled": sum(1 for s in filled_sections if s["filled"]),
        "open_questions": open_questions,
    }
