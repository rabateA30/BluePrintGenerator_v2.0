"""
Blueprint Generator – IT O&M Gestione Penali (EGP TGX)
Progetto:  IT - O&M - Gestione Penali — EGP TGX Italia
Template:  Blueprint_1.4_vuoto.docx
Output:    Blueprint_Penali_OM.docx
Lingua:    Italiano

NOTA: il materiale sorgente si trova nella cartella SharePoint
      "IT - O&M - Gestione Penali" (canale Grids e EGP TGX).
      Tutte le sezioni marcate con [INFO RICHIESTA: ...] devono
      essere confermate e completate dal BSN / team EGP TGX prima
      della pubblicazione del Blueprint.

Esecuzione:
    python generate_blueprint_penali.py [TEMPLATE_PATH] [OUTPUT_PATH]
"""

import copy
import os
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree  # noqa: F401


# ── path resolution ────────────────────────────────────────────────────────────

def _resolve_path(cli_index, env_name, default_value):
    if len(sys.argv) > cli_index and sys.argv[cli_index]:
        return sys.argv[cli_index]
    return os.environ.get(env_name, default_value)


TEMPLATE = _resolve_path(1, "BLUEPRINT_TEMPLATE", "Blueprint_1.4_vuoto.docx")
OUTPUT   = _resolve_path(2, "BLUEPRINT_OUTPUT",   "Blueprint_Penali_OM.docx")


# ── helpers ───────────────────────────────────────────────────────────────────

def find_para(doc, fragment, style=None, start=0):
    if start is None:
        start = 0
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if fragment in p.text:
            if style is None or p.style.name == style:
                return i, p
    return None, None


def _get_rpr(para):
    for run in para.runs:
        rpr = run._r.find(qn("w:rPr"))
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None


def set_para_text(para, new_text):
    rpr = _get_rpr(para)
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    r_elem = OxmlElement("w:r")
    if rpr is not None:
        r_elem.append(rpr)
    t_elem = OxmlElement("w:t")
    t_elem.text = new_text
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_elem.append(t_elem)
    para._p.append(r_elem)


def set_para_lines(para, lines):
    rpr = _get_rpr(para)
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    for idx, line in enumerate(lines):
        r_elem = OxmlElement("w:r")
        if rpr is not None:
            r_elem.append(copy.deepcopy(rpr))
        t_elem = OxmlElement("w:t")
        t_elem.text = line
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r_elem.append(t_elem)
        para._p.append(r_elem)
        if idx < len(lines) - 1:
            r_br = OxmlElement("w:r")
            if rpr is not None:
                r_br.append(copy.deepcopy(rpr))
            br = OxmlElement("w:br")
            r_br.append(br)
            para._p.append(r_br)


def fill_cell(cell, text):
    para = cell.paragraphs[0]
    set_para_text(para, text)


def fill_row(table, row_idx, values):
    row = table.rows[row_idx]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            fill_cell(row.cells[ci], val)


# ── document ──────────────────────────────────────────────────────────────────

doc = Document(TEMPLATE)


# ── TITOLO ────────────────────────────────────────────────────────────────────
_, p = find_para(doc, "Blueprint – xxx", "Heading 1")
if p:
    set_para_text(p, "Blueprint – IT O&M Gestione Penali: AI per la Gestione Automatizzata delle Penali Contrattuali EGP TGX")


# ── SOMMARIO ESECUTIVO ────────────────────────────────────────────────────────
_, p = find_para(doc, "Processi Identificati")
if p:
    set_para_text(p, "Processi Identificati: Gestione Penali O&M – Ricezione, Verifica, Contestazione e Monitoraggio delle Penali Contrattuali EGP TGX Italia")

_, p = find_para(doc, "Contesto Generale")
if p:
    set_para_lines(p, [
        "Contesto Generale:",
        "Il processo riguarda la gestione delle penali contrattuali nell'ambito delle attivita' di "
        "Operations & Maintenance degli impianti EGP TGX in Italia. Il processo include la ricezione "
        "delle notifiche di penale da committenti/enti regolatori, la verifica della correttezza "
        "del calcolo, l'eventuale contestazione documentata, la gestione del contenzioso e il "
        "monitoraggio del KPI di penale rispetto ai valori contrattuali.",
        "",
        "[INFO RICHIESTA AL BSN/EGP TGX: indicare il volume medio mensile di notifiche di penale, "
        "i tipologie principali (availability, performance ratio, curtailment, SLA fornitori) e "
        "il valore economico annuo gestito dal processo.]",
        "",
        "Proposta: architettura AI multi-agent (Orchestrator, Penalty Classifier Agent, "
        "Calculation Verifier Agent, Dispute Generator Agent, Knowledge Agent) per automatizzare "
        "il ciclo ricezione-verifica-contestazione, con il referente O&M come responsabile finale "
        "delle decisioni di contestazione (HITL).",
        "",
        "Proprietario processo: EGP/TGX – O&M Italia | "
        "[INFO RICHIESTA: specificare unita' organizzativa e numero di FTE coinvolti nel processo AS-IS]"
    ])


# ── PROCESSO 1 – TITOLO ───────────────────────────────────────────────────────
_, p = find_para(doc, "Processo 1: xxx", "Heading 1")
if p:
    set_para_text(p, "Processo 1: IT O&M Gestione Penali – Gestione Automatizzata delle Penali Contrattuali EGP TGX Italia")


# ── 1.1 SCOPO ─────────────────────────────────────────────────────────────────
_, p = find_para(doc, "1.1 Scopo", "Heading 3")
scopo_start = _ if _ is not None else 0
count = 0
for i, para in enumerate(doc.paragraphs):
    if i <= scopo_start:
        continue
    if para.style.name == "Heading 3":
        break
    if para.text.strip() == "xxx":
        if count == 0:
            set_para_text(para,
                "Il processo gestisce il ciclo di vita delle penali contrattuali relative agli "
                "impianti O&M EGP TGX in Italia: dalla ricezione delle notifiche di penale "
                "(emesse da committenti, TSO/DSO o enti regolatori) alla verifica di correttezza "
                "del calcolo, alla generazione documentata delle contestazioni, fino alla "
                "chiusura del contenzioso e al pagamento o recupero degli importi.")
        elif count == 1:
            set_para_text(para,
                "Obiettivo principale: ridurre l'esposizione economica alle penali attraverso "
                "la verifica sistematica e tempestiva dei calcoli, la contestazione automatizzata "
                "e documentata delle penali errate, e il monitoraggio proattivo dei KPI "
                "contrattuali per prevenire future penali.\n"
                "[INFO RICHIESTA AL BSN: confermare gli obiettivi prioritari – "
                "riduzione valore penali pagate, velocita' di contestazione, monitoraggio "
                "preventivo, o tutti e tre]")
        elif count == 2:
            set_para_text(para,
                "[INFO RICHIESTA ALL'IT OWNER: elencare i sistemi in scope – es. SAP/ERP "
                "(contratti e contabilita'), SCADA / sistemi di monitoraggio impianti, "
                "piattaforma documentale contratti, sistema di ticketing, strumenti Office "
                "(Excel, SharePoint) attualmente utilizzati per la gestione penali.]")
        count += 1
        if count >= 3:
            break


# ── 1.2 FINALITA' ─────────────────────────────────────────────────────────────
_, h12 = find_para(doc, "1.2 Finalit", "Heading 3")
finalita_start = _ if _ is not None else 0
count = 0
for i, para in enumerate(doc.paragraphs):
    if i <= finalita_start:
        continue
    if para.style.name == "Heading 3":
        break
    if para.text.strip() == "xxx":
        if count == 0:
            set_para_text(para,
                "- Ridurre il valore complessivo delle penali accettate attraverso la verifica "
                "automatizzata e sistematica della correttezza dei calcoli applicati dai "
                "committenti/enti emittenti, recuperando importi non dovuti.")
        elif count == 1:
            set_para_text(para,
                "- Accelerare i tempi di risposta alle notifiche di penale generando "
                "automaticamente le bozze di contestazione documentate, riducendo il "
                "rischio di decadenza dei termini contrattuali.\n"
                "[INFO RICHIESTA AL BSN: confermare la finestra temporale contrattuale "
                "per la contestazione delle penali (es. 15/30 giorni dalla notifica)]")
        elif count == 2:
            set_para_text(para,
                "- Monitorare proattivamente i KPI contrattuali (availability, performance "
                "ratio, curtailment) per anticipare e prevenire l'insorgenza di nuove penali "
                "attraverso alert predittivi e azioni correttive preventive.")
        count += 1
        if count >= 3:
            break


# ── 1.3 PERIMETRO ─────────────────────────────────────────────────────────────
_, h13 = find_para(doc, "1.3 Perimetro", "Heading 3")
h13_idx = _
if h13_idx is None:
    raise ValueError("Sezione '1.3 Perimetro' non trovata nel template.")
in_scope_count = 0
out_scope_count = 0
in_out_mode = None
for i, para in enumerate(doc.paragraphs):
    if i <= h13_idx:
        continue
    if para.style.name == "Heading 3" or para.style.name == "Heading 2":
        break
    txt = para.text.strip()
    if txt == "IN SCOPE:":
        in_out_mode = "in"
    elif txt == "OUT OF SCOPE:":
        in_out_mode = "out"
    elif txt == "- xxx":
        if in_out_mode == "in":
            if in_scope_count == 0:
                set_para_text(para, "- Ricezione, classificazione e registrazione delle notifiche di penale "
                              "da committenti, TSO/DSO e enti regolatori per gli impianti EGP TGX Italia")
            elif in_scope_count == 1:
                set_para_text(para, "- Verifica automatizzata della correttezza del calcolo delle penali "
                              "tramite confronto con dati SCADA/operativi e clausole contrattuali")
            elif in_scope_count == 2:
                set_para_text(para, "[INFO RICHIESTA AL BSN: confermare se e' in scope la gestione delle "
                              "penali verso i fornitori O&M (sub-appaltatori) oltre a quelle ricevute "
                              "da committenti/TSO]")
            in_scope_count += 1
        elif in_out_mode == "out":
            if out_scope_count == 0:
                set_para_text(para, "- Esecuzione fisica degli interventi tecnici correttivi sui impianti "
                              "(rimane in carico al team O&M di campo)")
            elif out_scope_count == 1:
                set_para_text(para, "[INFO RICHIESTA AL BSN: indicare esplicitamente cosa e' out of scope "
                              "– es. gestione penali di natura fiscale/tributaria, controversie legali "
                              "in fase giudiziale, penali relative a business unit diverse da EGP TGX Italia]")
            out_scope_count += 1


# ── 1.4 VINCOLI CHIAVE ────────────────────────────────────────────────────────
_, h14 = find_para(doc, "1.4 Vincoli chiave", "Heading 3")
_, pnorm = find_para(doc, "xxx", "List Paragraph", _)
if pnorm:
    set_para_text(pnorm,
        "Conformita' ai termini contrattuali e alle clausole penali definite nei contratti di "
        "servizio O&M; rispetto delle scadenze contrattuali per la contestazione delle penali "
        "(tipicamente 15-30 giorni dalla notifica); conformita' alle normative vigenti per "
        "la gestione dei dati di produzione e dei documenti contrattuali.\n"
        "[INFO RICHIESTA AL LEGAL/COMPLIANCE: confermare i termini contrattuali applicabili "
        "e i vincoli normativi specifici (ARERA, GME, contratti GSE, PPA)]")
_, ptec_val = find_para(doc, "xxx", "Normal", _)
if ptec_val:
    set_para_text(ptec_val,
        "[INFO RICHIESTA ALL'IT OWNER: definire disponibilita' di API verso SAP/ERP, "
        "sistemi SCADA/PI, piattaforma documentale contratti e sistema di ticketing. "
        "Indicare se esistono vincoli di accesso ai dati operativi degli impianti "
        "(es. aggregazione minima, latenza, GDPR/NDA con fornitori).]")
_, porg_val = find_para(doc, "xxx", "Normal", _)
if porg_val:
    set_para_text(porg_val,
        "[INFO RICHIESTA AL BUSINESS OWNER: definire il modello di governance per "
        "l'approvazione delle contestazioni generate dall'AI; identificare i referenti "
        "O&M e legali coinvolti nel processo di validazione (HITL); pianificare il "
        "change management per l'adozione del sistema da parte del team penali.]")


# ── TABLE 0 – SISTEMI COINVOLTI AS-IS ────────────────────────────────────────
t0 = doc.tables[0]
systems = [
    ("SAP / ERP Contratti e Contabilita'",
     "[INFO RICHIESTA ALL'IT OWNER: descrivere il modulo SAP utilizzato per la gestione "
     "dei contratti O&M e la contabilizzazione delle penali (es. SAP PM, SAP FI, "
     "modulo specifico Enel)]",
     "ERP / Contract Management"),
    ("SCADA / Sistema PI / Monitoraggio Impianti",
     "[INFO RICHIESTA: specificare la piattaforma di monitoraggio dei dati di produzione "
     "utilizzata (es. OSIsoft PI, sistema proprietario SCADA, piattaforma Enel); "
     "indicare la granularita' temporale dei dati disponibili (15 min, orario)]",
     "SCADA / Data Historian"),
    ("Piattaforma Documentale Contratti",
     "[INFO RICHIESTA: specificare dove sono archiviati i contratti di servizio "
     "e le clausole penali – es. SharePoint, Documentum, sistema gestione contratti Enel]",
     "Document Management / SharePoint"),
    ("Sistema di Ticketing / Gestione Reclami",
     "[INFO RICHIESTA: specificare se esiste un sistema di ticketing per il tracciamento "
     "delle contestazioni (es. ServiceNow, Jira, sistema interno Enel) o se la gestione "
     "avviene via email/Excel]",
     "Ticketing / Workflow"),
    ("Excel / Fogli di Calcolo",
     "Strumenti Office utilizzati per il calcolo manuale delle penali, il confronto "
     "con i dati di produzione e la preparazione delle contestazioni",
     "Spreadsheet (manual)"),
    ("Email / Posta Elettronica",
     "Canale principale per la ricezione delle notifiche di penale da committenti/TSO "
     "e per l'invio delle contestazioni documentate",
     "Communication (email)"),
]
for i, (s, r, t) in enumerate(systems):
    if i + 1 < len(t0.rows):
        fill_row(t0, i + 1, [s, r, t])


# ── AS-IS SOTTO-PROCESSI LABEL ────────────────────────────────────────────────
_, pa = find_para(doc, "Sotto-processo A: xxx", "Normal")
if pa:
    set_para_text(pa, "Sotto-processo A: Ricezione, Registrazione e Classificazione delle Notifiche di Penale")
_, pb = find_para(doc, "Sotto-processo B: xxx", "Normal")
if pb:
    set_para_text(pb, "Sotto-processo B: Verifica Correttezza Calcolo e Preparazione Contestazione")
_, pc = find_para(doc, "Sotto-processo C: xxx", "Normal")
if pc:
    set_para_text(pc, "Sotto-processo C: Gestione Contenzioso, Chiusura e Monitoraggio KPI Penali")


# ── TABLE 1 – AS-IS Sotto-processo A ─────────────────────────────────────────
t1 = doc.tables[1]
rows_a = [
    ("A1",
     "Ricezione della notifica di penale via email da parte del committente/TSO/DSO; "
     "il referente O&M legge la notifica e ne valuta l'entita'",
     "Referente O&M / Team Penali",
     "Email notifica penale (committente/TSO)",
     "Notifica ricevuta e presa in carico",
     "Email"),
    ("A2",
     "[INFO RICHIESTA: descrivere come viene registrata la notifica di penale – "
     "esiste un registro centralizzato (Excel, sistema ticketing) o la gestione "
     "e' distribuita per singolo referente/impianto?]",
     "Referente O&M",
     "Email notifica penale",
     "Penale registrata nel registro (Excel/sistema)",
     "[INFO: strumento di registrazione]"),
    ("A3",
     "Classificazione manuale della tipologia di penale: "
     "availability, performance ratio, curtailment, SLA fornitore, altro",
     "Referente O&M",
     "Notifica penale registrata",
     "Penale classificata per tipologia e impianto",
     "Email / Excel"),
    ("A4",
     "[INFO RICHIESTA: descrivere come viene identificato il contratto di riferimento "
     "e le clausole applicabili – il referente consulta manualmente il documento "
     "contrattuale o esiste un sistema di accesso rapido alle clausole penali?]",
     "Referente O&M / Legale",
     "Notifica penale classificata",
     "Contratto e clausole penali identificati",
     "[INFO: sistema documentale contratti]"),
    ("A5",
     "Valutazione preliminare della plausibilita' della penale "
     "in base all'esperienza del referente O&M",
     "Referente O&M",
     "Contratto e notifica penale",
     "Prima valutazione di congruita' (accettare / approfondire / contestare)",
     "Email / Excel"),
    ("A6",
     "[INFO RICHIESTA: descrivere il processo di escalation in caso di penali "
     "di importo elevato – soglie di escalation, figure coinvolte (management, legale)]",
     "Referente O&M / Management",
     "Prima valutazione",
     "Decisione su iter da seguire",
     "[INFO: canale di escalation]"),
]
for i, row_data in enumerate(rows_a):
    if i + 1 < len(t1.rows):
        fill_row(t1, i + 1, list(row_data))


# ── TABLE 2 – AS-IS Sotto-processo B ─────────────────────────────────────────
t2 = doc.tables[2]
rows_b = [
    ("B1",
     "Estrazione manuale dei dati di produzione/disponibilita' dell'impianto "
     "dal sistema SCADA/PI per il periodo di riferimento della penale",
     "Referente O&M / IT",
     "Periodo di riferimento penale",
     "Dataset dati produzione/disponibilita' per il periodo",
     "[INFO: SCADA / sistema PI]"),
    ("B2",
     "[INFO RICHIESTA: descrivere come viene effettuato il ricalcolo della penale – "
     "esiste un foglio Excel standardizzato o ogni referente usa metodi propri? "
     "Indicare la formula contrattuale applicata (es. formula availability penalty)]",
     "Referente O&M",
     "Dati produzione + clausole contrattuali",
     "Calcolo penale rielaborato",
     "Excel / [INFO: strumento di calcolo]"),
    ("B3",
     "Confronto tra il valore di penale notificato dal committente "
     "e il valore ricalcolato dal referente O&M",
     "Referente O&M",
     "Penale committente + ricalcolo O&M",
     "Delta identificato (congruente / discrepanza)",
     "Excel"),
    ("B4",
     "[INFO RICHIESTA: descrivere il processo di preparazione della contestazione – "
     "chi redige la lettera di contestazione, quali documenti vengono allegati "
     "(dati SCADA, log eventi, report tecnici), come viene approvata prima dell'invio?]",
     "Referente O&M / Legale",
     "Ricalcolo + documentazione tecnica",
     "Bozza lettera contestazione con allegati",
     "[INFO: strumento di redazione]"),
    ("B5",
     "[INFO RICHIESTA: specificare i controlli di qualita' applicati alla "
     "contestazione prima dell'invio – revisione legale obbligatoria, soglie importo, "
     "approvazione management]",
     "[INFO: attore di approvazione]",
     "Bozza contestazione",
     "Contestazione approvata",
     "[INFO: processo di approvazione]"),
]
for i, row_data in enumerate(rows_b):
    if i + 1 < len(t2.rows):
        fill_row(t2, i + 1, list(row_data))


# ── TABLE 3 – AS-IS Sotto-processo C ─────────────────────────────────────────
t3 = doc.tables[3]
rows_c = [
    ("C1",
     "Invio della contestazione via email al committente/TSO "
     "entro il termine contrattuale previsto",
     "Referente O&M",
     "Contestazione approvata",
     "Contestazione inviata e protocollata",
     "Email"),
    ("C2",
     "[INFO RICHIESTA: descrivere il processo di follow-up del contenzioso – "
     "come viene tracciato lo stato (es. in attesa risposta, accettata, rifiutata, "
     "negoziazione in corso)? Esiste un registro aggiornato o e' gestito via email?]",
     "Referente O&M",
     "Contestazione inviata",
     "Stato contenzioso tracciato",
     "Email / Excel / [INFO: sistema ticketing]"),
    ("C3",
     "Gestione delle risposte del committente: accettazione della contestazione "
     "(storno penale), rifiuto con motivazione, proposta di accordo parziale",
     "Referente O&M / Legale",
     "Risposta committente",
     "Decisione su chiusura o proseguimento contenzioso",
     "Email"),
    ("C4",
     "[INFO RICHIESTA: descrivere il processo di chiusura della penale – "
     "come viene aggiornato il sistema SAP/ERP in caso di storno? "
     "Chi autorizza il pagamento in caso di penale accettata?]",
     "Referente O&M / Finance",
     "Esito contenzioso",
     "Penale chiusa (stornata o pagata) nel sistema",
     "SAP / ERP"),
    ("C5",
     "[INFO RICHIESTA: specificare se esiste un reporting periodico sulle penali "
     "gestite (es. report mensile management) e chi lo predispone, "
     "con quali strumenti e con quale frequenza]",
     "Referente O&M / Management",
     "Registro penali aggiornato",
     "Report periodico per management",
     "Excel / [INFO: strumento di reporting]"),
    ("C6",
     "Monitoraggio dei KPI contrattuali (availability, PR, curtailment) "
     "per prevenire l'insorgenza di nuove penali nel periodo successivo",
     "Referente O&M",
     "Dati SCADA / report operativi",
     "Allerta su impianti a rischio penale nel periodo corrente",
     "[INFO: SCADA / sistema PI / Excel]"),
    ("C7",
     "[INFO RICHIESTA: descrivere come vengono identificate e implementate "
     "le azioni correttive per gli impianti che hanno generato penali ripetute – "
     "chi decide, quali sistemi vengono coinvolti, come si chiude il loop]",
     "[INFO: attori coinvolti]",
     "Analisi cause penale + KPI impianto",
     "Piano azioni correttive",
     "[INFO: sistema O&M / ERP]"),
]
for i, row_data in enumerate(rows_c):
    if i + 1 < len(t3.rows):
        fill_row(t3, i + 1, list(row_data))


# ── PROCESS CARDS AS-IS ───────────────────────────────────────────────────────
_, p = find_para(doc, "Card 1: xxx")
if p:
    set_para_text(p, "Card 1: Ricezione, Registrazione e Prima Valutazione Penali")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "[INFO RICHIESTA: elencare gli input del processo – "
        "es. email notifica penale, bollettino TSO, comunicazione GME/ARERA, "
        "notifica contrattuale da committente PPA]")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "1. Referente O&M riceve la notifica di penale via email",
        "2. [INFO: descrivere il processo di registrazione – Excel centralizzato, ticketing, altro]",
        "3. Classificazione manuale per tipologia (availability, PR, curtailment, SLA fornitore)",
        "4. Prima valutazione di congruita' basata sull'esperienza del referente"
    ])
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "[INFO: output del processo – es. penale registrata e classificata nel registro, "
        "decisione su iter (accettare / verificare / contestare)]")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "[INFO RICHIESTA: descrivere i pain point – "
        "es. nessun sistema centralizzato di ricezione, rischio di perdere notifiche, "
        "classificazione soggettiva, dipendenza da singoli referenti; "
        "indicare il volume mensile di notifiche gestite]")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "[INFO: quantificare il tempo medio per la ricezione e registrazione di una penale (ore/penale)]")

_, p = find_para(doc, "Card 2: xxx")
if p:
    set_para_text(p, "Card 2: Verifica Calcolo e Preparazione Contestazione")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "[INFO: input – notifica penale registrata, dati SCADA/PI per il periodo, clausole contrattuali]")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "1. Estrazione manuale dati di produzione/disponibilita' dal SCADA/PI per il periodo di riferimento",
        "2. [INFO: descrivere il ricalcolo – foglio Excel standard o metodo ad hoc per tipologia penale?]",
        "3. Confronto penale notificata vs. ricalcolo O&M; identificazione eventuale discrepanza",
        "4. [INFO: descrivere la redazione della contestazione – chi la scrive, chi la approva]"
    ])
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "[INFO: output – contestazione redatta con allegati tecnici o accettazione motivata della penale]")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "[INFO RICHIESTA: indicare i pain point – es. estrazione dati SCADA laboriosa, "
        "nessuno strumento di ricalcolo automatico, processo non standardizzato, "
        "elevato rischio di errore umano; tempo medio per la verifica di una penale (ore)]")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "N/A — processo manuale AS-IS")

_, p = find_para(doc, "Card 3: xxx")
if p:
    set_para_text(p, "Card 3: Gestione Contenzioso e Monitoraggio KPI")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "[INFO: input – contestazione inviata, risposta committente, registro penali, dati KPI impianto]")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "1. Invio contestazione via email entro i termini contrattuali",
        "2. [INFO: descrivere il tracking del contenzioso – Excel aggiornato manualmente, ticketing, email]",
        "3. Gestione risposta committente (accettazione storno / rifiuto / negoziazione)",
        "4. Chiusura nel sistema SAP/ERP (storno o pagamento penale)"
    ])
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "[INFO: output – penale chiusa (stornata o pagata), report periodico management; "
        "specificare il formato del report e la frequenza]")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "[INFO RICHIESTA: pain point specifici del contenzioso – "
        "es. nessun tracking automatico dello stato, rischio di scadenza termini, "
        "mancanza di visibilita' su penali aggregate per impianto/portfolio, "
        "monitoraggio KPI preventivo assente o manuale]")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "N/A — processo manuale AS-IS")


# ── PAIN POINT GENERALI ───────────────────────────────────────────────────────
_, p = find_para(doc, "1. Effort FTE elevato")
if p:
    set_para_text(p,
        "1. Effort FTE elevato: estrazione dati SCADA, ricalcolo in Excel e redazione delle "
        "contestazioni richiedono un elevato numero di ore per penale; "
        "[INFO RICHIESTA: quantificare il numero di FTE dedicati e le ore/penale medie]")
_, p = find_para(doc, "2. Tempi di attraversamento lunghi")
if p:
    set_para_text(p,
        "2. Tempi di attraversamento lunghi: dal ricevimento della notifica alla trasmissione "
        "della contestazione trascorrono [INFO: indicare i giorni medi attuali]; "
        "il rischio di superare i termini contrattuali per la contestazione e' concreto")
_, p = find_para(doc, "3. Colli di bottiglia")
if p:
    set_para_text(p,
        "3. Colli di bottiglia: il processo e' accentrato su pochi referenti esperti, "
        "la conoscenza delle clausole contrattuali e' distribuita informalmente, "
        "e il monitoraggio preventivo dei KPI non e' sistematico")


# ── TABLE 4 – DATA MAPPING ────────────────────────────────────────────────────
t4 = doc.tables[4]
data_mapping = [
    ("Notifiche di penale (tipo, importo, periodo di riferimento, impianto)",
     "Email committente / TSO / ARERA",
     "Penalty Classifier Agent / Registro Penali",
     "Email / PDF",
     "[INFO RICHIESTA: specificare il formato strutturato o meno delle notifiche – "
     "PDF libero, file strutturato, API del committente]"),
    ("Dati di produzione e disponibilita' impianto (per periodo di riferimento penale)",
     "[INFO: SCADA / Sistema PI / Data Historian]",
     "Calculation Verifier Agent",
     "CSV / REST API / PI Web API",
     "[INFO: specificare granularita' dati, latenza, disponibilita' storica (anni)]"),
    ("Clausole contrattuali e formule di calcolo penali",
     "[INFO: piattaforma documentale contratti / SharePoint]",
     "Knowledge Agent / Calculation Verifier Agent",
     "PDF / Word / JSON strutturato",
     "[INFO RICHIESTA: specificare se i contratti sono in formato strutturato estraibile "
     "o richiedono OCR/parsing; indicare il numero di contratti in scope]"),
    ("Storico penali gestite (registro storico: tipologia, importo, esito contestazione)",
     "[INFO: Excel / sistema ticketing / SAP]",
     "Knowledge Agent / Dashboard HITL",
     "Excel / CSV / API SAP",
     "[INFO RICHIESTA: specificare anni di storico disponibile e formato; "
     "indicare se esiste gia' un registro strutturato o solo email archiviate]"),
    ("Dati KPI contrattuali (availability, performance ratio, curtailment per impianto)",
     "[INFO: SCADA / sistema reporting operativo]",
     "Penalty Monitor Agent / Dashboard HITL",
     "CSV / REST API",
     "[INFO: specificare frequenza di aggiornamento dei KPI operativi e soglie contrattuali]"),
    ("Ordini di lavoro e interventi manutenzione (correlazione penale-evento tecnico)",
     "SAP / ERP Manutenzione",
     "Dispute Generator Agent / Knowledge Agent",
     "JSON / API SAP",
     "[INFO RICHIESTA: specificare se esiste l'integrazione tra OdL di manutenzione "
     "e registrazione cause penale (es. penale availability causata da fermo manutenzione programmata)]"),
    ("Risposte committente (accettazione/rifiuto contestazione)",
     "Email committente",
     "Registro Penali / Dashboard HITL",
     "Email / PDF",
     "[INFO: specificare se le risposte arrivano via email libera o tramite portale strutturato]"),
    ("Dati meteorologici e curtailment TSO",
     "[INFO: fonte dati meteo / log curtailment TSO]",
     "Calculation Verifier Agent",
     "CSV / API REST",
     "[INFO RICHIESTA: rilevante per penali di availability/PR – specificare la fonte "
     "dei dati di curtailment TSO e meteo utilizzata per le exclusion clause]"),
]
for i, row_data in enumerate(data_mapping):
    if i + 1 < len(t4.rows):
        fill_row(t4, i + 1, list(row_data))


# ── TABLE 5 – ARCHITETTURA FUNZIONALE TO-BE ───────────────────────────────────
t5 = doc.tables[5]
arch_rows = [
    ("Orchestrator Agent",
     "Coordina il flusso end-to-end tra gli agenti specializzati; gestisce l'interfaccia "
     "HITL con il referente O&M; consolida le analisi e presenta le raccomandazioni per "
     "la decisione di accettare o contestare la penale",
     "LLM multi-agent orchestration (LangGraph / AutoGen)"),
    ("Penalty Classifier Agent",
     "Acquisisce e classifica automaticamente le notifiche di penale per tipologia "
     "(availability, PR, curtailment, SLA fornitore), impianto, committente e periodo; "
     "registra nel sistema centralizzato e allerta il referente O&M",
     "LLM NLP classification + document parsing (GPT-4o / [INFO: modello selezionato])"),
    ("Calculation Verifier Agent",
     "Verifica automaticamente la correttezza del calcolo della penale: "
     "estrae i dati di produzione/disponibilita' dal SCADA/PI, applica la formula "
     "contrattuale, confronta con il valore notificato e quantifica la discrepanza",
     "Formula engine + LLM reasoning + API SCADA/PI"),
    ("Dispute Generator Agent",
     "Genera automaticamente la bozza di lettera di contestazione con allegati tecnici "
     "(dati SCADA, calcolo rielaborato, log eventi, exclusion clause applicate); "
     "la bozza viene presentata al referente O&M per revisione e approvazione HITL",
     "LLM text generation + document assembly (GPT-4o + template contrattuale)"),
    ("Penalty Monitor Agent",
     "Monitora proattivamente i KPI contrattuali (availability, performance ratio, "
     "curtailment) per tutti gli impianti in scope; genera alert predittivi "
     "quando un impianto si avvicina alle soglie di penale contrattuali",
     "Rule-based monitoring + ML trend forecasting + LLM alert generation"),
    ("Knowledge Agent",
     "Recupera e sintetizza informazioni da contratti (clausole penali, exclusion clause), "
     "storico penali gestite, casi di contestazione vinti/persi, procedure O&M; "
     "supporta il Calculation Verifier e il Dispute Generator con contesto specifico",
     "RAG + Vector DB (Azure AI Search / Qdrant) + LLM"),
    ("ETL Pipeline",
     "Acquisisce e normalizza dati da SCADA/PI, SAP/ERP, piattaforma documentale, "
     "email notifiche (con OCR/parsing per PDF); alimenta gli agenti AI con dati strutturati",
     "Azure Data Factory / [INFO: piattaforma dati Enel] + document parser"),
    ("Dashboard HITL (TO-BE)",
     "Interfaccia per il referente O&M: revisione classificazione penale, validazione "
     "ricalcolo, approvazione contestazione, tracking stato contenzioso, "
     "monitoraggio KPI portfolio impianti",
     "Power BI Embedded / React / Teams Bot; [INFO: preferenza del team EGP TGX]"),
    ("HITL Validation Module",
     "Gestisce il workflow di approvazione umana prima dell'invio di qualsiasi "
     "contestazione (obbligo HITL per ogni azione verso committente/TSO); "
     "traccia audit log completo di tutte le decisioni",
     "Power Automate / workflow engine; [INFO: integrazione con SAP esistente]"),
]
for i, row_data in enumerate(arch_rows):
    if i + 1 < len(t5.rows):
        fill_row(t5, i + 1, list(row_data))


# ── TO-BE SOTTO-PROCESSI LABEL ────────────────────────────────────────────────
_, p = find_para(doc, "Sotto-processo A: xxx (TO-BE)")
if p:
    set_para_text(p,
        "Sotto-processo A: Ricezione Automatizzata, Classificazione e Verifica Calcolo Penali (TO-BE)")
_, p = find_para(doc, "Sotto-processo B: xxx (TO-BE)")
if p:
    set_para_text(p,
        "Sotto-processo B: Generazione Contestazione e Workflow Approvazione HITL (TO-BE)")
_, p = find_para(doc, "Sotto-processo C: xxx (TO-BE)")
if p:
    set_para_text(p,
        "Sotto-processo C: Gestione Contenzioso, Monitoraggio Preventivo e Feedback Loop (TO-BE)")


# ── TABLE 6 – TO-BE Sotto-processo A ──────────────────────────────────────────
t6 = doc.tables[6]
rows_tobe_a = [
    ("A1",
     "ETL Pipeline monitora la casella email dedicata (o portale committente) "
     "e acquisisce automaticamente le notifiche di penale; le normalizza in formato strutturato",
     "ETL Pipeline (automatico)",
     "Email/notifica penale (PDF, email, portale)",
     "Notifica strutturata acquisita nel sistema",
     "Email gateway, [INFO: portale committente/TSO]",
     "Document parsing + OCR per PDF non strutturati; alert su acquisizione fallita"),
    ("A2",
     "Penalty Classifier Agent classifica la notifica per tipologia, impianto, "
     "committente, periodo di riferimento e importo; registra nel sistema centralizzato "
     "e allerta il referente O&M su Dashboard/Teams",
     "Penalty Classifier Agent",
     "Notifica strutturata",
     "Penale classificata e registrata; alert referente O&M",
     "Registro centralizzato, Teams Bot",
     "LLM NLP classification con confidence score; alert immediato su Team"),
    ("A3",
     "Calculation Verifier Agent estrae automaticamente i dati di produzione "
     "e disponibilita' per il periodo di riferimento dal SCADA/PI",
     "Calculation Verifier Agent",
     "Periodo di riferimento, ID impianto, tipologia penale",
     "Dataset produzione/disponibilita' per il periodo",
     "[INFO: SCADA / PI Web API]",
     "API automatica verso SCADA/PI; validazione qualita' dati estratti"),
    ("A4",
     "Knowledge Agent recupera dal contratto le clausole penali applicabili, "
     "la formula di calcolo e le exclusion clause pertinenti (meteo, curtailment TSO)",
     "Knowledge Agent",
     "Tipologia penale, committente, ID contratto",
     "Clausole contrattuali, formula calcolo, exclusion clause",
     "Knowledge Base (RAG) contratti",
     "RAG su clausole contrattuali; cross-check con storico penali analoghe"),
    ("A5",
     "Calculation Verifier Agent applica la formula contrattuale ai dati "
     "SCADA estratti, calcola il valore corretto e quantifica la discrepanza "
     "rispetto al valore notificato; applica automaticamente le exclusion clause",
     "Calculation Verifier Agent",
     "Dati SCADA + clausole contrattuali + exclusion clause",
     "Ricalcolo penale O&M, delta vs. notifica, breakdown esclusioni applicate",
     "Calculation Engine + LLM reasoning",
     "LLM per applicazione exclusion clause; calcolo deterministico per importo penale"),
    ("A6",
     "Orchestrator Agent consolida l'analisi (classificazione + ricalcolo + esclusioni) "
     "e genera il report strutturato per il referente O&M con raccomandazione "
     "(contestare totalmente / parzialmente / accettare) e motivazione",
     "Orchestrator Agent",
     "Output Calculation Verifier + Knowledge Agent",
     "Report analitico con raccomandazione motivata",
     "Dashboard HITL, Teams Bot",
     "LLM synthesis con spiegazione della raccomandazione"),
    ("A7",
     "HITL – Referente O&M revisiona il report AI, valida il ricalcolo, "
     "puo' aggiungere contesto operativo non rilevabile automaticamente "
     "(es. accordi verbali con committente, situazioni straordinarie impianto)",
     "Referente O&M",
     "Report analitico AI con raccomandazione",
     "Decisione validata: contestare (totalmente/parzialmente) o accettare",
     "Dashboard HITL",
     "Assist – Human decides e approva; audit log obbligatorio"),
]
for i, row_data in enumerate(rows_tobe_a):
    if i + 1 < len(t6.rows):
        fill_row(t6, i + 1, list(row_data))


# ── TABLE 7 – TO-BE Sotto-processo B ──────────────────────────────────────────
t7 = doc.tables[7]
rows_tobe_b = [
    ("B1",
     "Dispute Generator Agent riceve la decisione di contestazione validata "
     "e avvia la generazione automatica della lettera di contestazione",
     "Dispute Generator Agent",
     "Decisione HITL: contestare (totale/parziale) + motivazioni",
     "Generazione lettera contestazione avviata",
     "Dispute Generator Agent",
     "Trigger automatico post-approvazione HITL"),
    ("B2",
     "Knowledge Agent recupera dalla KB il template di contestazione appropriato "
     "per tipologia di penale e committente, lo storico contestazioni analoghe "
     "e i precedenti vinti/persi per orientare il tono e la strategia",
     "Knowledge Agent",
     "Tipologia penale, committente, storico contestazioni",
     "Template contestazione + best practice da casi analoghi",
     "Knowledge Base (RAG)",
     "RAG su storico contestazioni e best practice"),
    ("B3",
     "Dispute Generator Agent genera la bozza di lettera di contestazione "
     "compilando: dati tecnici, ricalcolo documentato, exclusion clause applicate, "
     "riferimenti normativi/contrattuali, richiesta di storno importo",
     "Dispute Generator Agent",
     "Template + dati tecnici + ricalcolo + exclusion clause",
     "Bozza lettera contestazione completa con allegati tecnici",
     "LLM text generation + document assembly",
     "LLM generation con template strutturato; allegati tecnici auto-generati"),
    ("B4",
     "Orchestrator Agent presenta la bozza contestazione al referente O&M "
     "su Dashboard HITL per revisione, modifica e approvazione finale "
     "prima di qualsiasi invio al committente",
     "Orchestrator Agent",
     "Bozza contestazione completa",
     "Bozza presentata su Dashboard per revisione HITL",
     "Dashboard HITL",
     "HITL obbligatorio – nessun invio automatico al committente"),
    ("B5",
     "HITL – Referente O&M revisiona la bozza, apporta eventuali modifiche "
     "(tono, dati tecnici, strategia argomentativa) e approva la versione finale",
     "Referente O&M",
     "Bozza contestazione AI",
     "Contestazione approvata (o modificata e approvata)",
     "Dashboard HITL",
     "Human revises and approves; modifica manuale supportata da UI"),
    ("B6",
     "[INFO RICHIESTA: specificare se e' richiesto un secondo livello di approvazione "
     "(es. legale, management) per contestazioni oltre una soglia di importo; "
     "definire il workflow di approvazione multilivello se applicabile]",
     "[INFO: attore secondo livello]",
     "Contestazione approvata da referente O&M",
     "[INFO: contestazione approvata anche dal secondo livello]",
     "[INFO: sistema di workflow approvazione]",
     "[INFO: specificare soglie importo e livelli di approvazione]"),
    ("B7",
     "Sistema invia la contestazione approvata al committente via email "
     "o portale dedicato, registra l'invio con timestamp e allerta il referente "
     "sulla scadenza per la risposta",
     "Orchestrator Agent",
     "Contestazione approvata",
     "Contestazione inviata; reminder scadenza risposta impostato",
     "Email gateway / [INFO: portale committente]",
     "Invio automatico solo post-approvazione HITL; audit log completo"),
]
for i, row_data in enumerate(rows_tobe_b):
    if i + 1 < len(t7.rows):
        fill_row(t7, i + 1, list(row_data))


# ── TABLE 8 – TO-BE Sotto-processo C ──────────────────────────────────────────
t8 = doc.tables[8]
rows_tobe_c = [
    ("C1",
     "Sistema monitora automaticamente i termini di risposta del committente "
     "e notifica il referente O&M con escalation automatica in caso di scadenza imminente",
     "Orchestrator Agent",
     "Contestazione inviata, scadenza risposta",
     "Alert proattivo a referente O&M; escalation automatica",
     "Dashboard HITL, Teams Bot",
     "Reminder automatici 7/3/1 giorni prima della scadenza"),
    ("C2",
     "Penalty Classifier Agent acquisisce la risposta del committente "
     "(accettazione storno / rifiuto / proposta parziale) e aggiorna lo stato "
     "del contenzioso nel registro centralizzato",
     "Penalty Classifier Agent",
     "Risposta committente (email / portale)",
     "Stato contenzioso aggiornato; notifica al referente O&M",
     "Registro centralizzato, Dashboard HITL",
     "NLP classification risposta committente; update automatico stato"),
    ("C3",
     "HITL – Referente O&M gestisce l'esito del contenzioso: accettazione dello storno, "
     "gestione del rifiuto (valutazione escalation legale), approvazione accordo parziale",
     "Referente O&M",
     "Risposta committente classificata",
     "Decisione finale su chiusura/proseguimento contenzioso",
     "Dashboard HITL",
     "Human decides – AI supporta con analisi precedenti analoghi"),
    ("C4",
     "Orchestrator Agent aggiorna SAP/ERP (storno penale o registrazione pagamento) "
     "sulla base della decisione HITL validata dal referente O&M",
     "Orchestrator Agent",
     "Decisione HITL validata",
     "SAP/ERP aggiornato; penale chiusa nel sistema",
     "SAP / ERP",
     "Integrazione automatica SAP solo post-approvazione HITL"),
    ("C5",
     "Knowledge Agent acquisisce l'esito del contenzioso (vinto/perso/parziale) "
     "e aggiorna la Knowledge Base con il case documentato per apprendimento futuro",
     "Knowledge Agent",
     "Esito contenzioso + documentazione tecnica",
     "KB aggiornata con nuovo case penale risolto",
     "Knowledge Base (RAG)",
     "Document ingestion + embedding update per continual learning"),
    ("C6",
     "Penalty Monitor Agent monitora continuamente i KPI contrattuali "
     "di tutti gli impianti in scope e genera alert predittivi quando "
     "un impianto si avvicina alla soglia di penale nel mese/trimestre corrente",
     "Penalty Monitor Agent",
     "KPI operativi real-time (SCADA/PI)",
     "Alert predittivo per impianti a rischio penale",
     "SCADA/PI, Dashboard HITL, Teams Bot",
     "ML trend forecasting + regole contrattuali; alert proattivi"),
    ("C7",
     "Orchestrator Agent genera report periodici aggregati "
     "(mensile/trimestrale) per management: valore penali ricevute/contestate/vinte, "
     "performance portfolio, impianti critici, savings da contestazioni",
     "Orchestrator Agent",
     "Dati aggregati registro penali + KPI portfolio",
     "Report management con insights e trend; dashboard executive",
     "Dashboard HITL, Power BI",
     "LLM-generated insights da dati aggregati"),
]
for i, row_data in enumerate(rows_tobe_c):
    if i + 1 < len(t8.rows):
        fill_row(t8, i + 1, list(row_data))


# ── PROCESS CARDS TO-BE ───────────────────────────────────────────────────────
_, p = find_para(doc, "Card 1: xxx (TO-BE)")
if p:
    set_para_text(p, "Card 1: Ricezione Automatizzata, Classificazione e Verifica Calcolo (TO-BE)")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "- Notifiche penale (email/portale committente); Dati SCADA/PI per il periodo; "
        "Contratti e clausole penali (KB); Dati meteo/curtailment TSO (exclusion clause)")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "1. ETL Pipeline acquisisce automaticamente le notifiche di penale (email/portale)",
        "2. Penalty Classifier Agent classifica e registra; allerta il referente O&M",
        "3. Calculation Verifier Agent estrae dati SCADA e applica formula contrattuale",
        "4. Knowledge Agent recupera exclusion clause e confronta con storico casi analoghi",
        "5. HITL: Referente O&M valida il ricalcolo e decide se contestare"
    ])
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "- Report analitico validato con raccomandazione motivata (contestare totalmente / "
        "parzialmente / accettare) e ricalcolo documentato")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "SCADA/PI, Knowledge Base contratti, Penalty Classifier Agent, "
        "Calculation Verifier Agent, Dashboard HITL")

_, p = find_para(doc, "Card 2: xxx (TO-BE)")
if p:
    set_para_text(p, "Card 2: Generazione Contestazione e Approvazione HITL (TO-BE)")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "- Decisione HITL di contestare; Ricalcolo documentato; Exclusion clause; "
        "Template contestazione da KB; Storico contestazioni analoghe")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "1. Dispute Generator Agent genera bozza contestazione con allegati tecnici",
        "2. Knowledge Agent recupera best practice da casi analoghi vinti",
        "3. HITL: Referente O&M revisiona e approva la bozza (modifica possibile)",
        "4. [INFO: secondo livello approvazione se importo > soglia – da configurare]",
        "5. Sistema invia la contestazione approvata al committente con timestamp"
    ])
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "- Contestazione inviata al committente; Reminder scadenza risposta impostato; "
        "Audit log completo della decisione HITL")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "Knowledge Base (RAG), Dispute Generator Agent, Dashboard HITL, Email/Portale committente")

_, p = find_para(doc, "Card 3: xxx (TO-BE)")
if p:
    set_para_text(p, "Card 3: Gestione Contenzioso, Chiusura e Monitoraggio Preventivo (TO-BE)")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- Risposta committente; Registro penali aggiornato; KPI operativi real-time impianti")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "1. Penalty Classifier Agent acquisisce risposta committente e aggiorna stato",
        "2. HITL: Referente O&M gestisce esito (storno / rifiuto / accordo parziale)",
        "3. SAP/ERP aggiornato automaticamente post-approvazione HITL",
        "4. Knowledge Agent aggiorna KB con case documentato per futuro apprendimento",
        "5. Penalty Monitor Agent monitora KPI impianti e genera alert predittivi preventivi"
    ])
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "- KB aggiornata; SAP/ERP allineato; KPI portfolio monitorati; "
        "Report management periodico generato; Alert preventivi per impianti a rischio")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "Knowledge Base, Orchestrator Agent, SAP/ERP, SCADA/PI, Dashboard HITL, Power BI")


# ── 3.4 COSA NON FA L'AI ──────────────────────────────────────────────────────
_, p = find_para(doc, "3.4 Cosa NON fa l'AI", "Heading 3")
_, p_cosa = find_para(doc, "xxx", "Normal", _)
if p_cosa:
    set_para_lines(p_cosa, [
        "Non invia autonomamente contestazioni al committente senza esplicita approvazione del referente O&M (HITL obbligatorio)",
        "Non aggiorna SAP/ERP (storno o pagamento penale) senza validazione umana",
        "Non prende decisioni su escalation legale o apertura contenzioso formale in autonomia",
        "Non interpreta clausole contrattuali ambigue senza revisione del referente O&M o del legale",
        "Non negozia direttamente con il committente o il TSO in autonomia",
        "Non accede a sistemi di pagamento o gestisce flussi finanziari in autonomia",
        "Non modifica i parametri operativi degli impianti per evitare penali (rimane in carico al team O&M di campo)",
        "Non divulga dati contrattuali o produttivi a soggetti terzi non autorizzati"
    ])


# ── 3.5 COMPONENTE AI ─────────────────────────────────────────────────────────
_, p = find_para(doc, "3.5 Componente ai", "Heading 3")
_, p_ai = find_para(doc, "Xxx", "Normal", _)
if p_ai:
    set_para_lines(p_ai, [
        "Architettura multi-agent AI composta da agenti specializzati coordinati da un Orchestrator:",
        "",
        "(1) Orchestrator Agent – coordina il flusso end-to-end e gestisce l'interfaccia HITL.",
        "Pattern tassonomico: Orchestrator / Coordinator Agent",
        "",
        "(2) Penalty Classifier Agent – acquisisce e classifica notifiche di penale per tipologia, "
        "impianto e committente; genera alert al referente O&M.",
        "Pattern tassonomico: Classifier / Intake Agent",
        "",
        "(3) Calculation Verifier Agent – estrae dati SCADA, applica formula contrattuale, "
        "verifica correttezza calcolo e quantifica discrepanza.",
        "Pattern tassonomico: Verifier / Calculator Agent",
        "",
        "(4) Dispute Generator Agent – genera automaticamente la bozza di contestazione "
        "con allegati tecnici; utilizza template contrattuali e best practice da storico.",
        "Pattern tassonomico: Generator / Writer Agent",
        "",
        "(5) Penalty Monitor Agent – monitora proattivamente i KPI contrattuali degli impianti "
        "e genera alert predittivi prima dell'insorgenza di penali.",
        "Pattern tassonomico: Monitor / Early Warning Agent",
        "",
        "(6) Knowledge Agent – recupera e sintetizza clausole contrattuali, exclusion clause, "
        "storico contestazioni e best practice tramite RAG.",
        "Pattern tassonomico: Retriever / RAG Agent"
    ])


# ── DELTA AS-IS vs TO-BE ──────────────────────────────────────────────────────
_, p = find_para(doc, "AS-IS:", "Normal")
if p:
    set_para_text(p, "AS-IS:")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "- Ricezione manuale delle notifiche via email; nessun sistema centralizzato di acquisizione e tracking",
        "- Verifica calcolo penali su Excel: processo non standardizzato, elevato rischio di errore umano",
        "- Knowledge contrattuale distribuita informalmente tra referenti; "
        "[INFO: specificare il numero di referenti coinvolti]",
        "- Nessun monitoraggio preventivo sistematico dei KPI contrattuali; penali rilevate a posteriori",
        "- [INFO RICHIESTA: aggiungere ulteriori limitazioni specifiche emerse dai documenti sorgente nella "
        "cartella IT - O&M - Gestione Penali]"
    ])
_, p = find_para(doc, "TO-BE:", "Normal")
if p:
    set_para_text(p, "TO-BE:")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "- ETL Pipeline acquisisce e registra automaticamente le notifiche di penale; "
        "Penalty Classifier Agent categorizza e allerta in tempo reale",
        "- Calculation Verifier Agent verifica automaticamente il calcolo tramite dati SCADA/PI "
        "e clausole contrattuali; errori di calcolo identificati sistematicamente",
        "- Dispute Generator Agent genera bozze contestazioni documentate; "
        "il referente O&M rivede e approva tramite Dashboard HITL",
        "- Penalty Monitor Agent monitora proattivamente i KPI e genera alert preventivi "
        "prima che si materializzino nuove penali",
        "- Knowledge Base strutturata rende accessibili clausole contrattuali e "
        "best practice a tutti i referenti O&M"
    ])


# ── TABLE 9 – ROADMAP ─────────────────────────────────────────────────────────
t9 = doc.tables[9]
roadmap = [
    ("Fase 1 – Raccolta Requisiti e Accessi\n(M1-M2)",
     "Allineamento con BSN e team EGP TGX per definire perimetro definitivo; "
     "mappatura API verso SAP/ERP, SCADA/PI e piattaforma documentale; "
     "raccolta e inventario contratti in scope; "
     "[INFO: specificare il numero di contratti e committenti in scope]",
     "Architettura tecnica validata; accordi IT per accessi API SCADA e SAP; "
     "inventario contratti e clausole penali disponibili",
     "2 mesi — [INFO RICHIESTA AL PM: confermare]"),
    ("Fase 2 – Data Pipeline e Knowledge Base\n(M2-M4)",
     "Implementazione pipeline ETL per acquisizione notifiche penale (email gateway), "
     "estrazione dati SCADA/PI e integrazione SAP; "
     "digitalizzazione e indicizzazione Knowledge Base contratti "
     "(clausole penali, formule calcolo, exclusion clause, storico contestazioni)",
     "ETL funzionante su ambiente di sviluppo; "
     "KB popolata con contratti in scope e storico penali disponibile",
     "2 mesi — [INFO RICHIESTA AL PM: confermare]"),
    ("Fase 3 – AI PoC\n(M3-M6)",
     "Sviluppo e test Penalty Classifier Agent e Calculation Verifier Agent "
     "su dati storici di penali gestite; "
     "validazione accuracy classificazione e correttezza ricalcoli; "
     "[INFO: specificare se Penalty Monitor Agent e' incluso nel PoC]",
     "PoC demo con accuracy classificazione > 90% e match calcolo > 95% "
     "su dataset storico; Dispute Generator Agent genera bozze valutate positivamente "
     "dal team O&M",
     "3 mesi — [INFO RICHIESTA AL PM: confermare]"),
    ("Fase 4 – Integration e HITL Interface\n(M5-M8)",
     "Integrazione agenti con sistemi reali (SAP, SCADA/PI, email gateway, "
     "piattaforma documentale); "
     "sviluppo Dashboard HITL e Teams Bot; "
     "test operativi con referenti O&M pilota",
     "Sistema integrato su ambiente pre-produzione; "
     "Dashboard HITL validata da [INFO: n.] referenti O&M pilota",
     "3 mesi — [INFO RICHIESTA AL PM: confermare]"),
    ("Fase 5 – Pilota e Validazione\n(M7-M9)",
     "Pilota operativo su [INFO: indicare subset di committenti/impianti pilota]; "
     "monitoraggio KPI (savings contestazioni, riduzione tempi, accuracy AI); "
     "go/no-go per scale-up al portfolio completo EGP TGX Italia",
     "Savings da contestazioni AI verificati su penali pilota; "
     "adozione sistema > 80% referenti O&M coinvolti; report go/no-go scale-up",
     "3 mesi — [INFO RICHIESTA AL PM: confermare]"),
    ("Fase 6 – Scale-up\n(M9-M12)",
     "Estensione a tutti i contratti e committenti in scope EGP TGX Italia; "
     "attivazione Penalty Monitor Agent su tutti gli impianti; "
     "formazione referenti O&M; messa a regime del sistema; "
     "[INFO: specificare il numero totale di contratti/impianti in scope per lo scale-up]",
     "Sistema in produzione su portfolio completo EGP TGX Italia; "
     "piano di ottimizzazione continua e aggiornamento KB definito",
     "3 mesi — [INFO RICHIESTA AL PM: confermare]"),
]
for i, row_data in enumerate(roadmap):
    if i + 1 < len(t9.rows):
        fill_row(t9, i + 1, list(row_data))


# ── TABLE 10 – KPI QUANTITATIVI ───────────────────────────────────────────────
t10 = doc.tables[10]
kpi_quant = [
    ("Valore penali contestate e vinte / annuo (€)",
     "[INFO RICHIESTA AL BSN: richiedere il valore attuale AS-IS di penali contestate "
     "e stornate per anno (€) e il tasso di successo delle contestazioni manuali]",
     "[INFO: definire target TO-BE – es. +30% savings da contestazioni grazie a "
     "verifica sistematica e tempestiva]",
     "Report SAP: penali stornate / totale penali contestate (€ e %)"),
    ("Tempo medio di risposta alle notifiche di penale (giorni dalla ricezione all'invio contestazione)",
     "[INFO RICHIESTA AL BSN: richiedere il tempo medio AS-IS in giorni "
     "dalla ricezione della notifica all'invio della contestazione]",
     "< 3 giorni lavorativi (AI + HITL) [INFO: confermare target con Business Owner]",
     "Timestamp ricezione notifica -> data invio contestazione nel sistema"),
    ("Tasso di errori di calcolo rilevati nelle penali ricevute (%)",
     "[INFO RICHIESTA AL BSN: richiedere una stima della percentuale attuale di penali "
     "con errori di calcolo (rilevati manualmente) e del tasso di contestazione attuale]",
     "> 95% accuracy identificazione errori calcolo da parte del Calculation Verifier Agent",
     "Confronto ricalcolo AI vs. esito accettazione committente su campione storico"),
    ("Tempo medio per la verifica e ricalcolo di una penale (ore/penale)",
     "[INFO RICHIESTA AL BSN: richiedere il tempo medio AS-IS per la verifica manuale "
     "di una penale (estrazione SCADA + ricalcolo + redazione contestazione, ore)]",
     "< 30 minuti (AI-assisted + HITL) [INFO: confermare target]",
     "Timestamp avvio verifica -> approvazione HITL contestazione su Dashboard"),
    ("Numero di penali gestite per FTE per mese",
     "[INFO RICHIESTA: richiedere la capacita' AS-IS (penali/mese/FTE) per il team penali]",
     "[INFO: definire target TO-BE – tipicamente 3-5x rispetto AS-IS con AI]",
     "Penali processate nel mese / FTE assegnati al processo penali"),
    ("Copertura monitoraggio preventivo KPI (% impianti monitorati con alert proattivi)",
     "0% (monitoraggio preventivo assente o manuale sporadico)",
     "100% impianti in scope con alert preventivi attivi",
     "Rapporto impianti con Penalty Monitor Agent attivo / totale impianti in scope"),
    ("Tasso di rispetto scadenze contrattuali contestazione (%)",
     "[INFO RICHIESTA: richiedere la percentuale AS-IS di contestazioni inviate entro i "
     "termini contrattuali; indicare se si sono verificati casi di decadenza]",
     "100% contestazioni inviate entro i termini contrattuali",
     "Contestazioni inviate entro scadenza / totale contestazioni avviate"),
    ("Copertura Knowledge Base contratti (% clausole penali indicizzate)",
     "0% (clausole penali non strutturate e non indicizzate)",
     "> 95% clausole penali di contratti in scope indicizzate in KB",
     "Ratio clausole penali indicizzate / totale clausole penali nei contratti in scope"),
]
for i, row_data in enumerate(kpi_quant):
    if i + 1 < len(t10.rows):
        fill_row(t10, i + 1, list(row_data))


# ── TABLE 11 – KPI QUALITATIVI ────────────────────────────────────────────────
t11 = doc.tables[11]
kpi_qual = [
    ("Soddisfazione referenti O&M (self-reported)",
     "Riduzione del carico di lavoro manuale per verifica e contestazione penali; "
     "target: maggioranza referenti lo valuta ridotto (scala Likert, survey semestrale) "
     "[INFO: definire target score con Business Owner]",
     "Survey semestrale ai referenti O&M del team penali EGP TGX"),
    ("Adozione sistema AI da parte dei referenti O&M",
     "Target: > 90% referenti utilizza regolarmente la Dashboard HITL "
     "entro 6 mesi dal go-live [INFO: confermare target con EGP TGX]",
     "Log di accesso e utilizzo Dashboard HITL (sessioni/settimana per utente)"),
    ("Qualita' contestazioni generate dall'AI (% accettate senza modifiche sostanziali)",
     "Target: > 70% delle bozze di contestazione generate dall'AI accettate dal "
     "referente senza modifiche sostanziali entro 12 mesi dal go-live "
     "[INFO: confermare target con i referenti pilota]",
     "Tracking HITL: bozze approvate as-is / totale bozze generate"),
    ("Completezza e accuratezza della Knowledge Base contrattuale",
     "Target: referenti O&M valutano il Knowledge Agent utile e accurato per "
     "il recupero di clausole e best practice contestazione "
     "(scala Likert > 4/5) [INFO: calibrare target]",
     "Survey trimestrale ai referenti O&M + tasso query KB risolte senza escalation"),
    ("Tracciabilita' e auditabilita' del processo penali",
     "Target: 100% delle decisioni di accettare/contestare una penale documentate "
     "con motivazione nel sistema (obbligatorio per audit e controllo interno)",
     "Audit log Dashboard HITL: % decisioni con motivazione registrata"),
    ("Efficacia monitoraggio preventivo KPI",
     "Target: > 80% degli impianti che hanno superato la soglia di penale "
     "nel trimestre precedente hanno ricevuto un alert predittivo in anticipo "
     "[INFO: calibrare dopo deployment Penalty Monitor Agent]",
     "Confronto: impianti con alert preventivo emesso / impianti che hanno generato "
     "penale nel periodo di riferimento"),
    ("Soddisfazione management su reporting portfolio penali",
     "Target: management valuta il report AI-generated accurato e utile per "
     "le decisioni di portfolio (scala Likert > 4/5) "
     "[INFO: definire target con il management EGP TGX]",
     "Survey trimestrale al management O&M EGP TGX"),
    ("[INFO RICHIESTA AL BSN: specificare ulteriori KPI qualitativi rilevanti "
     "per la Gestione Penali O&M EGP TGX emersi dai documenti sorgente]",
     "[INFO: definire target]",
     "[INFO: definire metodo di misurazione]"),
]
for i, row_data in enumerate(kpi_qual):
    if i + 1 < len(t11.rows):
        fill_row(t11, i + 1, list(row_data))


# ── SAVE ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"\n Blueprint IT O&M Gestione Penali salvato in: {OUTPUT}")
print("\nSezioni completate con dati contestuali EGP TGX - Gestione Penali.")
print()
print("=" * 70)
print("OPEN QUESTIONS — SEZIONI CON [INFO RICHIESTA] DA COMPLETARE:")
print("(Le sezioni seguenti NON potevano essere compilate senza il")
print(" materiale sorgente presente nella cartella SharePoint")
print(" 'IT - O&M - Gestione Penali')")
print("=" * 70)
open_questions = [
    {
        "section": "Sommario Esecutivo — Contesto Generale",
        "reason": "Volume mensile di notifiche penale, tipologie principali e valore economico annuo gestito non disponibili senza il materiale sorgente",
        "suggestion": "Richiedere al BSN/EGP TGX: n. notifiche/mese, breakdown per tipologia (availability, PR, curtailment, SLA fornitore) e valore economico annuo (€)"
    },
    {
        "section": "1.1 Scopo — Sistemi in scope",
        "reason": "Sistemi IT specifici utilizzati per la gestione penali (modulo SAP, piattaforma SCADA/PI, sistema documentale) non noti senza analisi dei documenti sorgente",
        "suggestion": "Richiedere all'IT Owner: nome e versione dei sistemi SAP, SCADA/PI e documentale in uso; disponibilita' di API per integrazione"
    },
    {
        "section": "1.2 Finalita' — Termini contrattuali contestazione",
        "reason": "Finestra temporale contrattuale per la contestazione delle penali non specificata nel materiale disponibile",
        "suggestion": "Richiedere al BSN/Legale: termine contrattuale (es. 15/30 giorni) per ciascuna tipologia di committente/contratto in scope"
    },
    {
        "section": "1.3 Perimetro — Penali verso fornitori O&M",
        "reason": "Non chiaro se il perimetro include anche la gestione delle penali applicate ai sub-appaltatori O&M",
        "suggestion": "Confermare con BSN: il processo include solo le penali ricevute da committenti/TSO, o anche le penali emesse verso i fornitori di manutenzione?"
    },
    {
        "section": "Tabla 0 (Sistemi) — SAP modulo, SCADA/PI, sistema documentale",
        "reason": "Dettaglio dei sistemi IT specifici (modulo SAP, nome piattaforma SCADA/PI, sistema documentale contratti) non disponibile senza i documenti sorgente",
        "suggestion": "Richiedere all'IT Owner: descrizione completa dell'architettura IT AS-IS (sistema SAP, SCADA/PI, documentale, ticketing)"
    },
    {
        "section": "Tabla 1 (AS-IS Sotto-processo A) — Registro e sistema di tracking",
        "reason": "Non noto se esiste un registro centralizzato delle penali (Excel, SAP, ticketing) o se la gestione e' distribuita via email per singolo referente",
        "suggestion": "Analizzare il materiale sorgente nella cartella SharePoint per identificare il tool di registrazione attuale (es. Excel allegato, screenshot sistema)"
    },
    {
        "section": "Tabla 2 (AS-IS Sotto-processo B) — Metodo di ricalcolo",
        "reason": "Non nota la metodologia di ricalcolo manuale delle penali (foglio Excel standard vs. ad hoc per tipologia) senza i documenti sorgente",
        "suggestion": "Recuperare il/i file Excel di calcolo penali dalla cartella SharePoint e analizzare le formule applicate per ciascuna tipologia"
    },
    {
        "section": "Tabla 4 (Data Mapping) — Formato notifiche committente",
        "reason": "Formato delle notifiche di penale (PDF libero, strutturato, API portale committente) sconosciuto senza analisi dei documenti sorgente",
        "suggestion": "Analizzare esempi di notifiche penale nella cartella SharePoint per definire la strategia di parsing (strutturato vs. OCR/LLM extraction)"
    },
    {
        "section": "Tabla 9 (Roadmap) — Numero contratti e impianti in scope",
        "reason": "Numero di contratti, committenti e impianti da includere nel pilota e nello scale-up non definiti",
        "suggestion": "Richiedere al BSN e al PM: n. contratti attivi, n. committenti/TSO, n. impianti in scope; proposta impianti/committenti per il pilota (Fase 5)"
    },
    {
        "section": "Tabla 10 (KPI Quantitativi) — Tutti i valori baseline AS-IS",
        "reason": "Nessun dato quantitativo baseline disponibile (valore penali/anno, tempo medio verifica, tasso contestazioni vinte) senza il materiale sorgente",
        "suggestion": "Richiedere al BSN una scheda KPI AS-IS con: valore penali ricevute/anno (€), valore penali stornate/anno (€), tempo medio verifica (ore/penale), n. FTE dedicati"
    },
    {
        "section": "Tabla 10 (KPI Quantitativi) — Ruoli stakeholder (Business Owner, Data Owner, IT Owner, Product Owner)",
        "reason": "Nomi e unita' organizzative dei ruoli chiave del progetto non specificati",
        "suggestion": "Richiedere al BSN: nome e unita' organizzativa per Business Owner, Data Owner (responsabile dati SCADA/contratti), IT Owner (GICT/IT EGP), Product Owner (AISA Factory)"
    },
]
for i, oq in enumerate(open_questions, 1):
    print(f"\n[{i:02d}] Sezione: {oq['section']}")
    print(f"     Motivo: {oq['reason']}")
    print(f"     Suggerimento: {oq['suggestion']}")
print()
print(f"Totale open questions: {len(open_questions)}")
print()
print("PROSSIMI PASSI:")
print("  1. Condividere questo Blueprint bozza con il BSN EGP TGX per validazione")
print("  2. Fornire accesso alla cartella SharePoint 'IT - O&M - Gestione Penali'")
print("     per completare le sezioni con [INFO RICHIESTA]")
print("  3. Raccogliere i valori baseline KPI AS-IS dal team O&M")
print("  4. Completare e approvare il Blueprint prima del kickoff tecnico")
