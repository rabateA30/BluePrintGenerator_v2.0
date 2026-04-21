"""
Blueprint Generator – AIDRO (AI + Idroelettrico / AI-Driven Remote Operations)
Progetto:  GDS+OP-36 AIDRO — EGP TGX
Template:  Blueprint_1.4_vuoto.docx
Output:    Blueprint_AIDRO.docx
Lingua:    Italiano

NOTA: il file sorgente (GDS+OP-36 AIDRO.msg) non è leggibile in formato testo.
      Tutte le sezioni marcate con [INFO RICHIESTA: ...] devono essere confermate
      e completate dal BSN / team EGP TGX prima della pubblicazione.

Esecuzione:
    python generate_blueprint_AIDRO.py [TEMPLATE_PATH] [OUTPUT_PATH]
"""

import copy
import os
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree  # noqa: F401


# ── path resolution ────────────────────────────────────────────────────────────

def _resolve_path(cli_index, env_name, default_value):
    if len(sys.argv) > cli_index and sys.argv[cli_index]:
        return sys.argv[cli_index]
    return os.environ.get(env_name, default_value)


TEMPLATE = _resolve_path(1, "BLUEPRINT_TEMPLATE", "Blueprint_1.4_vuoto.docx")
OUTPUT   = _resolve_path(2, "BLUEPRINT_OUTPUT",   "Blueprint_AIDRO.docx")


# ── helpers ───────────────────────────────────────────────────────────────────

def find_para(doc, fragment, style=None, start=0):
    if start is None:
        start = 0
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if fragment in p.text:
            if style is None or p.style.name == style:
                return i, p
    return None, None


def _get_rpr(para):
    for run in para.runs:
        rpr = run._r.find(qn("w:rPr"))
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None


def set_para_text(para, new_text):
    rpr = _get_rpr(para)
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    r_elem = OxmlElement("w:r")
    if rpr is not None:
        r_elem.append(rpr)
    t_elem = OxmlElement("w:t")
    t_elem.text = new_text
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_elem.append(t_elem)
    para._p.append(r_elem)


def set_para_lines(para, lines):
    rpr = _get_rpr(para)
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    for idx, line in enumerate(lines):
        r_elem = OxmlElement("w:r")
        if rpr is not None:
            r_elem.append(copy.deepcopy(rpr))
        t_elem = OxmlElement("w:t")
        t_elem.text = line
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r_elem.append(t_elem)
        para._p.append(r_elem)
        if idx < len(lines) - 1:
            r_br = OxmlElement("w:r")
            if rpr is not None:
                r_br.append(copy.deepcopy(rpr))
            br = OxmlElement("w:br")
            r_br.append(br)
            para._p.append(r_br)


def fill_cell(cell, text):
    para = cell.paragraphs[0]
    set_para_text(para, text)


def fill_row(table, row_idx, values):
    row = table.rows[row_idx]
    for ci, val in enumerate(values):
        if ci < len(row.cells):
            fill_cell(row.cells[ci], val)


# ── document ──────────────────────────────────────────────────────────────────

doc = Document(TEMPLATE)


# ── TITOLO ────────────────────────────────────────────────────────────────────
_, p = find_para(doc, "Blueprint – xxx", "Heading 1")
if p:
    set_para_text(p, "Blueprint – AIDRO: AI per la Gestione Operativa degli Impianti Idroelettrici")


# ── SOMMARIO ESECUTIVO ────────────────────────────────────────────────────────
_, p = find_para(doc, "Processi Identificati")
if p:
    set_para_text(p, "Processi Identificati: AIDRO – Ottimizzazione Operativa e Manutenzione Predittiva Impianti Idroelettrici con AI")

_, p = find_para(doc, "Contesto Generale")
if p:
    set_para_lines(p, [
        "Contesto Generale:",
        "Il progetto AIDRO (GDS+OP-36) nasce nell'ambito dell'AI Scale Up Accelerator del canale Grids e "
        "EGP TGX. L'obiettivo e' introdurre capacita' AI per ottimizzare la gestione operativa e la manutenzione "
        "predittiva degli impianti idroelettrici EGP, automatizzando l'analisi dei dati di produzione e "
        "degli alert di anomalia e ottimizzando la pianificazione degli interventi.",
        "",
        "[INFO RICHIESTA AL BSN/EGP TGX: indicare il numero di impianti in scope, la portata installata "
        "complessiva (MW) e i sistemi informativi attualmente utilizzati per il monitoraggio.]",
        "",
        "Proposta: architettura multi-agent AI (Orchestrator, Hydro Monitor Agent, Planning Agent, "
        "Knowledge Agent, Water Resource Agent) per automatizzare il ciclo monitoraggio-analisi-pianificazione, "
        "mantenendo il supervisore di impianto come responsabile finale delle decisioni operative (HITL).",
        "",
        "Proprietario processo: EGP/TGX – O&M Idroelettrico | "
        "[INFO RICHIESTA: specificare unita' organizzativa e numero di FTE coinvolti nel processo AS-IS]"
    ])


# ── PROCESSO 1 – TITOLO ───────────────────────────────────────────────────────
_, p = find_para(doc, "Processo 1: xxx", "Heading 1")
if p:
    set_para_text(p, "Processo 1: AIDRO – Gestione Operativa e Manutenzione Predittiva Impianti Idroelettrici")


# ── 1.1 SCOPO ─────────────────────────────────────────────────────────────────
_, p = find_para(doc, "1.1 Scopo", "Heading 3")
scopo_start = _ if _ else 0
count = 0
for i, para in enumerate(doc.paragraphs):
    if i <= scopo_start:
        continue
    if para.style.name == "Heading 3":
        break
    if para.text.strip() == "xxx":
        if count == 0:
            set_para_text(para,
                "Il processo gestisce il monitoraggio operativo e la manutenzione predittiva degli impianti "
                "idroelettrici EGP, acquisendo dati di produzione, livello invaso, portata e stato dei "
                "componenti (turbine, generatori, paratoie) dai sistemi di controllo SCADA e dai sensori di campo.")
        elif count == 1:
            set_para_text(para,
                "Obiettivo principale: identificare tempestivamente anomalie e componenti a rischio, "
                "ottimizzare la pianificazione degli interventi di manutenzione e supportare le decisioni "
                "operative sulla regolazione della produzione idroelettrica.\n"
                "[INFO RICHIESTA AL BSN: confermare obiettivi primari e secondari del progetto AIDRO]")
        elif count == 2:
            set_para_text(para,
                "[INFO RICHIESTA ALL'IT OWNER: elencare i sistemi in scope – es. SCADA idroelettrico, "
                "sistema di monitoraggio invasi, piattaforma predittiva, ERP manutenzione, "
                "sistemi di gestione risorsa idrica (idrologia).]")
        count += 1
        if count >= 3:
            break


# ── 1.2 FINALITA' ─────────────────────────────────────────────────────────────
_, h12 = find_para(doc, "1.2 Finalit", "Heading 3")
count = 0
for i, para in enumerate(doc.paragraphs):
    if i <= _:
        continue
    if para.style.name == "Heading 3":
        break
    if para.text.strip() == "xxx":
        if count == 0:
            set_para_text(para,
                "- Ridurre il Mean Time To Repair (MTTR) della manutenzione degli impianti idroelettrici "
                "attraverso analisi automatizzata degli allarmi e pianificazione ottimizzata degli interventi.")
        elif count == 1:
            set_para_text(para,
                "- Ottimizzare la gestione della risorsa idrica integrando previsioni meteo-idrologiche, "
                "prezzi di mercato e disponibilita' delle macchine per massimizzare la produzione.\n"
                "[INFO RICHIESTA AL BSN: confermare se l'ottimizzazione della risorsa idrica e' in scope]")
        elif count == 2:
            set_para_text(para,
                "- Strutturare e rendere accessibile la knowledge tecnica degli impianti "
                "(manuali OEM, storici interventi, procedure operative) tramite Knowledge Base con AI generativa.")
        count += 1
        if count >= 3:
            break


# ── 1.3 PERIMETRO ─────────────────────────────────────────────────────────────
_, h13 = find_para(doc, "1.3 Perimetro", "Heading 3")
h13_idx = _
in_scope_count = 0
out_scope_count = 0
in_out_mode = None
for i, para in enumerate(doc.paragraphs):
    if i <= h13_idx:
        continue
    if para.style.name == "Heading 3" or para.style.name == "Heading 2":
        break
    txt = para.text.strip()
    if txt == "IN SCOPE:":
        in_out_mode = "in"
    elif txt == "OUT OF SCOPE:":
        in_out_mode = "out"
    elif txt == "- xxx":
        if in_out_mode == "in":
            if in_scope_count == 0:
                set_para_text(para, "- Monitoraggio predittivo dei componenti critici degli impianti idroelettrici "
                              "(turbine, generatori, paratoie, cuscinetti)")
            elif in_scope_count == 1:
                set_para_text(para, "- Pianificazione ottimizzata degli interventi di manutenzione tramite AI multi-agent "
                              "(considerando vincoli idrologici, meteo, disponibilita' risorse e ricambi)")
            elif in_scope_count == 2:
                set_para_text(para, "[INFO RICHIESTA AL BSN: indicare se e' in scope l'ottimizzazione della "
                              "programmazione della produzione (dispatch) o solo la manutenzione]")
            in_scope_count += 1
        elif in_out_mode == "out":
            if out_scope_count == 0:
                set_para_text(para, "- Esecuzione fisica degli interventi di manutenzione e delle operazioni "
                              "di campo (rimane in carico a tecnici/fornitori)")
            elif out_scope_count == 1:
                set_para_text(para, "[INFO RICHIESTA AL BSN: indicare esplicitamente cosa e' out of scope "
                              "– es. gestione concessioni idriche, rapporti con enti regolatori, "
                              "manutenzione civile delle opere (dighe, gallerie)]")
            out_scope_count += 1


# ── 1.4 VINCOLI CHIAVE ────────────────────────────────────────────────────────
_, h14 = find_para(doc, "1.4 Vincoli chiave", "Heading 3")
_, pnorm = find_para(doc, "xxx", "List Paragraph", _)
if pnorm:
    set_para_text(pnorm,
        "Conformita' alle normative per la sicurezza delle dighe e degli invasi (D.Lgs. 1 ottobre 2021 n. 152; "
        "D.P.R. 1/11/1959 n. 1363 vigilanza dighe); rispetto dei requisiti di concessione idrica e delle "
        "prescrizioni degli enti di controllo (RSDI, Autorita' di Bacino).\n"
        "[INFO RICHIESTA AL LEGAL/COMPLIANCE: confermare vincoli normativi specifici applicabili]")
_, ptec_val = find_para(doc, "xxx", "Normal", _)
if ptec_val:
    set_para_text(ptec_val,
        "[INFO RICHIESTA ALL'IT OWNER: definire disponibilita' di API verso sistemi SCADA idroelettrico, "
        "piattaforma di monitoraggio invasi, sistemi idrometeorologici e ERP manutenzione. "
        "Indicare se esistono vincoli di cybersecurity per l'accesso remoto agli impianti.]")
_, porg_val = find_para(doc, "xxx", "Normal", _)
if porg_val:
    set_para_text(porg_val,
        "[INFO RICHIESTA AL BUSINESS OWNER: definire il modello di governance per l'approvazione AI-generated; "
        "identificare i supervisori di impianto coinvolti nel pilota; pianificare il change management per "
        "l'adozione del sistema HITL da parte degli operatori.]")


# ── TABLE 0 – SISTEMI COINVOLTI AS-IS ────────────────────────────────────────
t0 = doc.tables[0]
systems = [
    ("SCADA Idroelettrico",
     "[INFO RICHIESTA ALL'IT OWNER: descrivere il sistema SCADA per il controllo degli impianti idroelettrici "
     "(es. sistema proprietario OEM, piattaforma Enel, altro) e le sue capacita' di monitoraggio]",
     "SCADA / Control System"),
    ("Sistema di Monitoraggio Invasi",
     "[INFO RICHIESTA: specificare il sistema/piattaforma per il monitoraggio dei livelli idrometrici, "
     "portate e volumi invasati; indicare se e' un sistema proprietario o di terze parti]",
     "Hydrology Monitoring Platform"),
    ("Piattaforma Predittiva / Analytics",
     "[INFO RICHIESTA: specificare se esiste una piattaforma di analytics predittivi per i componenti "
     "(analoga a SAS nel progetto eolico) o se questo e' un gap da colmare con AIDRO]",
     "Analytics / ML Platform"),
    ("ERP Manutenzione",
     "Sistema gestione ordini di lavoro, ricambi e pianificazione risorse per la manutenzione degli impianti",
     "ERP / CMMS"),
    ("Sistema Previsioni Idro-Meteo",
     "[INFO RICHIESTA: specificare la fonte delle previsioni meteo e idrologiche utilizzate – "
     "es. servizio interno Enel, ARPA, fornitori esterni (ECMWF, OpenWeather)]",
     "Hydro-Meteo Forecasting"),
    ("Email / Telefono / Strumenti Office",
     "Canali di comunicazione informale tra personale d'ufficio e supervisori d'impianto per la "
     "notifica di anomalie e la pianificazione degli interventi",
     "Communication (informal)"),
]
for i, (s, r, t) in enumerate(systems):
    fill_row(t0, i + 1, [s, r, t])


# ── AS-IS SOTTO-PROCESSI LABEL ────────────────────────────────────────────────
_, pa = find_para(doc, "Sotto-processo A: xxx", "Normal")
if pa:
    set_para_text(pa, "Sotto-processo A: Acquisizione Dati e Monitoraggio Operativo Impianti Idroelettrici")
_, pb = find_para(doc, "Sotto-processo B: xxx", "Normal")
if pb:
    set_para_text(pb, "Sotto-processo B: Analisi Anomalie e Comunicazione al Supervisore")
_, pc = find_para(doc, "Sotto-processo C: xxx", "Normal")
if pc:
    set_para_text(pc, "Sotto-processo C: Pianificazione e Coordinamento Interventi di Manutenzione")


# ── TABLE 1 – AS-IS Sotto-processo A ─────────────────────────────────────────
t1 = doc.tables[1]
rows_a = [
    ("A1",
     "Sensori di campo acquisiscono dati operativi (livello invaso, portata, produzione, "
     "temperatura e vibrazione turbine/generatori) e li trasmettono al sistema SCADA",
     "Sensori / Impianto (automatico)",
     "Dati raw sensori (livello, portata, T, vibrazione)",
     "Stream dati acquisito e memorizzato nel SCADA",
     "[INFO: specificare sistema SCADA]"),
    ("A2",
     "[INFO RICHIESTA: descrivere come il SCADA elabora i dati e genera gli alert di anomalia – "
     "automaticamente tramite soglie configurate, o richiede intervento manuale?]",
     "SCADA (automatico/manuale)",
     "Stream dati sensori",
     "Alert e segnalazioni di anomalia",
     "SCADA"),
    ("A3",
     "Il personale d'ufficio accede al sistema di monitoraggio (dashboard/SCADA) "
     "per il controllo giornaliero dello stato degli impianti",
     "Tecnico d'ufficio",
     "Dashboard/SCADA aggiornata",
     "Sessione di monitoraggio avviata",
     "SCADA, Dashboard"),
    ("A4",
     "Analisi manuale degli alert e delle anomalie rilevate dal sistema; valutazione "
     "della priorita' di intervento basata su esperienza personale",
     "Tecnico d'ufficio",
     "Alert e segnalazioni di anomalia",
     "Lista componenti/impianti da gestire con priorita' assegnata",
     "Dashboard SCADA"),
    ("A5",
     "[INFO RICHIESTA: descrivere come il tecnico d'ufficio comunica le anomalie al supervisore "
     "di impianto – telefono, email, sistema ticketing?]",
     "Tecnico d'ufficio",
     "Lista anomalie prioritizzate",
     "Comunicazione inviata al supervisore",
     "[INFO: canale di comunicazione]"),
    ("A6",
     "Il supervisore riceve la comunicazione e avvia la valutazione per la pianificazione dell'intervento",
     "Supervisore impianto",
     "Comunicazione tecnico d'ufficio",
     "Informazioni ricevute, avvio pianificazione",
     "[INFO: canale di comunicazione]"),
]
for i, row_data in enumerate(rows_a):
    if i + 1 < len(t1.rows):
        fill_row(t1, i + 1, list(row_data))


# ── TABLE 2 – AS-IS Sotto-processo B ─────────────────────────────────────────
t2 = doc.tables[2]
rows_b = [
    ("B1",
     "[INFO RICHIESTA: descrivere il processo di analisi delle anomalie e di valutazione "
     "del rischio per i componenti degli impianti idroelettrici nel processo AS-IS]",
     "Tecnico d'ufficio",
     "[INFO: input del processo di analisi]",
     "[INFO: output del processo di analisi]",
     "[INFO: sistemi utilizzati]"),
    ("B2",
     "[INFO RICHIESTA: specificare se esiste una piattaforma predittiva per i componenti "
     "degli impianti idroelettrici o se l'analisi e' completamente manuale]",
     "[INFO: attore]",
     "[INFO: input]",
     "[INFO: output]",
     "[INFO: sistemi]"),
    ("B3",
     "Valutazione della disponibilita' idrica e delle previsioni di portata per il "
     "periodo di interesse (influenza sulla pianificazione della manutenzione)",
     "Tecnico d'ufficio / Operatore idrologia",
     "Previsioni meteo-idrologiche",
     "Finestre operative identificate per la manutenzione (invaso in regola, portata minima garantita)",
     "[INFO: sistema previsioni idro-meteo]"),
    ("B4",
     "[INFO RICHIESTA: descrivere come vengono gestite le comunicazioni con i fornitori "
     "di manutenzione e come viene verificata la disponibilita' dei ricambi]",
     "[INFO: attore]",
     "[INFO: input]",
     "[INFO: output]",
     "[INFO: sistemi]"),
    ("B5",
     "[INFO RICHIESTA: specificare se esistono check di sicurezza specifici "
     "(es. verifiche dighe, autorizzazioni per lavori in quota, permessi di lavoro) "
     "che condizionano la pianificazione degli interventi]",
     "[INFO: attore]",
     "[INFO: input]",
     "[INFO: output]",
     "[INFO: sistemi]"),
]
for i, row_data in enumerate(rows_b):
    if i + 1 < len(t2.rows):
        fill_row(t2, i + 1, list(row_data))


# ── TABLE 3 – AS-IS Sotto-processo C ─────────────────────────────────────────
t3 = doc.tables[3]
rows_c = [
    ("C1",
     "Supervisore verifica la disponibilita' delle imprese e dei tecnici specializzati per "
     "le date ipotizzate, considerando i vincoli di accesso agli impianti (remoti/montagna)",
     "Supervisore impianto",
     "Lista componenti da manutenere",
     "Disponibilita' risorse verificata",
     "Telefono / Email"),
    ("C2",
     "Supervisore verifica le previsioni idro-meteo per identificare finestre operative "
     "idonee (meteo, livello invaso, portata minima garantita per il rilascio d'acqua)",
     "Supervisore impianto",
     "Previsioni meteo e idrologiche",
     "Finestre operative idonee identificate",
     "[INFO: sistema previsioni]"),
    ("C3",
     "Supervisore verifica la disponibilita' dei ricambi necessari per gli interventi",
     "Supervisore impianto",
     "Lista componenti da sostituire",
     "Disponibilita' ricambi verificata",
     "ERP / Telefono fornitori"),
    ("C4",
     "[INFO RICHIESTA: descrivere il processo di pianificazione degli interventi – "
     "esiste un sistema di CMMS/ERP dedicato o la pianificazione avviene su foglio Excel/manuale?]",
     "Supervisore impianto",
     "Vincoli identificati",
     "Piano manutenzione bozza",
     "[INFO: strumento di pianificazione]"),
    ("C5",
     "Coordinamento logistica per l'accesso agli impianti: permessi, dotazioni di sicurezza "
     "specifiche per impianti idroelettrici (lavori in quota, ambienti confinati, vicino all'acqua)",
     "Supervisore impianto",
     "Piano manutenzione bozza",
     "Logistica confermata",
     "Telefono / Email"),
    ("C6",
     "Emissione ordini di lavoro formali per i tecnici e i fornitori",
     "Supervisore impianto",
     "Piano manutenzione approvato",
     "Ordini di lavoro emessi",
     "ERP / [INFO: sistema OdL]"),
    ("C7",
     "Esecuzione degli interventi di manutenzione da parte di tecnici/fornitori specializzati",
     "Tecnico / Fornitore",
     "Ordini di lavoro",
     "Verbale intervento compilato",
     "Mobile App / Carta"),
]
for i, row_data in enumerate(rows_c):
    if i + 1 < len(t3.rows):
        fill_row(t3, i + 1, list(row_data))


# ── PROCESS CARDS AS-IS ───────────────────────────────────────────────────────
_, p = find_para(doc, "Card 1: xxx")
if p:
    set_para_text(p, "Card 1: Monitoraggio Operativo e Rilevamento Anomalie Impianti Idroelettrici")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "[INFO RICHIESTA: elencare gli input del processo di monitoraggio – "
        "es. dati livello invaso, portata, vibrazione turbine, temperatura generatori]")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "1. [INFO: descrivere come il SCADA acquisisce e visualizza i dati operativi]",
        "2. [INFO: descrivere come vengono generati gli alert – soglie, regole, modelli predittivi?]",
        "3. Tecnico d'ufficio analizza manualmente gli alert sulla dashboard",
        "4. Prioritizzazione manuale degli interventi basata su esperienza"
    ])
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "[INFO: output del processo – es. lista anomalie prioritizzate, comunicazione al supervisore]")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "[INFO RICHIESTA: descrivere i pain point del processo AS-IS – "
        "es. analisi manuale lenta, mancanza di correlazione automatica, "
        "dipendenza da esperienza individuale; indicare ore/giorno per il monitoraggio]")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "[INFO: quantificare l'intervento umano richiesto – es. ore/giorno per il monitoraggio manuale]")

_, p = find_para(doc, "Card 2: xxx")
if p:
    set_para_text(p, "Card 2: Analisi Anomalie e Comunicazione al Supervisore")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "[INFO: input – es. lista anomalie da SCADA, storico manutenzione]")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "1. Tecnico d'ufficio esamina le anomalie rilevate e ne valuta la criticita'",
        "2. Verifica manuale dello storico guasti per componenti analoghi",
        "3. [INFO: descrivere come viene assegnata la priorita' – regole fisse, esperienza, altro]",
        "4. [INFO: descrivere il canale di comunicazione con il supervisore – telefono, email, sistema]"
    ])
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "[INFO: output – lista componenti da manutenere comunicata al supervisore]")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "[INFO RICHIESTA: indicare i pain point – es. analisi soggettiva, "
        "comunicazione non tracciata, tempo medio di analisi per singola anomalia]")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "N/A — processo manuale AS-IS")

_, p = find_para(doc, "Card 3: xxx")
if p:
    set_para_text(p, "Card 3: Pianificazione e Coordinamento Interventi di Manutenzione")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "[INFO: input – lista componenti da manutenere, previsioni idro-meteo, disponibilita' risorse]")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "1. Supervisore raccoglie informazioni su disponibilita' fornitori, meteo, ricambi (processo manuale)",
        "2. [INFO: descrivere come vengono integrati i vincoli idrologici nella pianificazione]",
        "3. Coordinamento logistica (accessi impianto, permessi, dotazioni sicurezza specialistiche)",
        "4. Emissione ordini di lavoro e comunicazione a tecnici/fornitori"
    ])
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "[INFO: output – piano manutenzione e ordini di lavoro emessi; "
        "specificare il formato – carta, Excel, ERP]")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "[INFO RICHIESTA: pain point specifici della pianificazione per impianti idroelettrici – "
        "es. dipendenza dalla disponibilita' idrica, difficolta' di accesso agli impianti remoti, "
        "vincoli normativi sulle finestre operative]")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "N/A — processo manuale AS-IS")


# ── PAIN POINT GENERALI ───────────────────────────────────────────────────────
_, p = find_para(doc, "1. Effort FTE elevato")
if p:
    set_para_text(p,
        "1. Effort FTE elevato: monitoraggio manuale degli impianti richiede personale dedicato; "
        "[INFO RICHIESTA: quantificare il numero di FTE coinvolti e le ore/giorno di analisi manuale]")
_, p = find_para(doc, "2. Tempi di attraversamento lunghi")
if p:
    set_para_text(p,
        "2. Tempi di attraversamento lunghi: dalla rilevazione dell'anomalia all'intervento manutentivo "
        "trascorrono [INFO: indicare il MTTR medio attuale in giorni]; l'accesso agli impianti remoti "
        "aumenta ulteriormente i tempi operativi")
_, p = find_para(doc, "3. Colli di bottiglia")
if p:
    set_para_text(p,
        "3. Colli di bottiglia: la pianificazione e' centralizzata sul supervisore, "
        "aggravata dalla necessita' di coordinare vincoli idrologici specifici degli impianti idroelettrici")


# ── TABLE 4 – DATA MAPPING ────────────────────────────────────────────────────
t4 = doc.tables[4]
data_mapping = [
    ("Dati operativi turbina (vibrazione, temperatura, potenza)",
     "[INFO: SCADA/sistema sensori]",
     "[INFO: piattaforma analytics/predittiva TO-BE]",
     "JSON stream / CSV",
     "[INFO: specificare frequenza di aggiornamento e protocollo di trasmissione]"),
    ("Livello invaso e portata in ingresso/uscita",
     "[INFO: sistema monitoraggio invasi]",
     "Water Resource Agent / Planning Agent",
     "JSON / REST API",
     "[INFO: orizzonte storico disponibile; frequenza di aggiornamento]"),
    ("Previsioni idro-meteorologiche (portata, precipitazioni)",
     "[INFO: sistema previsioni idro-meteo]",
     "Water Resource Agent / Planning Agent",
     "CSV / API REST",
     "[INFO: orizzonte previsionale (ore/giorni); fonte dati (ARPA, servizio interno Enel, altro)]"),
    ("Alert e anomalie SCADA",
     "[INFO: SCADA idroelettrico]",
     "Hydro Monitor Agent",
     "JSON / OPC-UA",
     "[INFO: specificare la struttura del messaggio di alert e i codici di anomalia]"),
    ("Storico ordini di lavoro e interventi",
     "ERP Manutenzione",
     "Knowledge Base (RAG)",
     "JSON / Excel / PDF",
     "Digitalizzazione e indicizzazione necessaria; [INFO: quantificare anni di storico disponibile]"),
    ("Documentazione tecnica impianti (manuali OEM, procedure)",
     "[INFO: archivio documentale EGP/TGX]",
     "Knowledge Base (RAG)",
     "PDF / Word",
     "[INFO: specificare la piattaforma di archiviazione – SharePoint, file server, altro]"),
    ("Dati prezzi mercato energia (PUN, profilo prezzo)",
     "GME / [INFO: fonte prezzi energia]",
     "Planning Agent / Water Resource Agent",
     "CSV / API REST",
     "[INFO: rilevante se l'ottimizzazione della produzione e' in scope – confermare con BSN]"),
    ("Disponibilita' fornitori e ricambi",
     "ERP / [INFO: sistema gestione fornitori]",
     "Planning Agent",
     "JSON / Manuale",
     "[INFO RICHIESTA: specificare se esiste un'API verso il sistema gestione fornitori "
     "o se la disponibilita' viene raccolta manualmente]"),
    ("Verbali post-intervento",
     "Tecnico manutentore (Mobile App TO-BE / Carta AS-IS)",
     "Knowledge Base / Dashboard",
     "JSON / PDF",
     "Feedback loop per aggiornamento modelli e KB; [INFO: attualmente su carta o sistema digitale?]"),
]
for i, row_data in enumerate(data_mapping):
    if i + 1 < len(t4.rows):
        fill_row(t4, i + 1, list(row_data))


# ── TABLE 5 – ARCHITETTURA FUNZIONALE TO-BE ───────────────────────────────────
t5 = doc.tables[5]
arch_rows = [
    ("Orchestrator Agent",
     "Coordina il flusso end-to-end tra gli agenti specializzati; gestisce l'interfaccia HITL "
     "con supervisori e tecnici d'ufficio; consolida le analisi e presenta le raccomandazioni",
     "LLM multi-agent orchestration (LangGraph / AutoGen)"),
    ("Hydro Monitor Agent",
     "Analizza gli alert SCADA e i dati operativi degli impianti idroelettrici; classifica "
     "le anomalie per priorita', urgenza e impatto sulla produzione; correla con storico KB",
     "ML classification + LLM reasoning (GPT-4o / [INFO: modello selezionato])"),
    ("Planning Agent",
     "Ottimizza il piano di manutenzione considerando vincoli multi-dimensionali: "
     "idrologici (disponibilita' acqua, deflusso minimo vitale), meteo, economici, "
     "disponibilita' fornitori e ricambi, finestre operative",
     "Constraint optimization + LLM (function calling)"),
    ("Water Resource Agent",
     "[INFO RICHIESTA: specificare se questo agente e' in scope – potrebbe ottimizzare "
     "la produzione idroelettrica in funzione delle previsioni idrologiche e dei prezzi di mercato]",
     "[INFO: tecnologia – es. optimization solver + LLM + API GME]"),
    ("Knowledge Agent",
     "Recupera e sintetizza informazioni tecniche da documentazione strutturata e non strutturata "
     "(manuali OEM turbine/generatori/paratoie, storici interventi, procedure specifiche impianto); "
     "aggiorna la KB con feedback post-intervento",
     "RAG + Vector DB (Azure AI Search / Qdrant) + LLM"),
    ("ETL Pipeline",
     "Acquisisce e normalizza dati da SCADA, sistemi di monitoraggio invasi, previsioni "
     "idro-meteo e ERP; alimenta gli agenti AI con dati strutturati e aggiornati",
     "Azure Data Factory / Apache Spark / [INFO: piattaforma dati Enel esistente]"),
    ("Dashboard HITL (TO-BE)",
     "Interfaccia per supervisori e tecnici: revisione report AI, approvazione piani manutenzione, "
     "monitoraggio KPI operativi, visualizzazione stato impianti",
     "Power BI Embedded / React / Teams Bot; [INFO: preferenza del team EGP TGX]"),
    ("HITL Validation Module",
     "Gestisce il workflow di approvazione umana prima dell'esecuzione di azioni critiche "
     "(emissione OdL, comunicazione fornitori, modifiche operative)",
     "Power Automate / workflow engine; [INFO: integrazione con ERP esistente]"),
]
for i, row_data in enumerate(arch_rows):
    if i + 1 < len(t5.rows):
        fill_row(t5, i + 1, list(row_data))


# ── TO-BE SOTTO-PROCESSI LABEL ────────────────────────────────────────────────
_, p = find_para(doc, "Sotto-processo A: xxx (TO-BE)")
if p:
    set_para_text(p,
        "Sotto-processo A: Monitoraggio Automatizzato, Analisi e Classificazione Anomalie (TO-BE)")
_, p = find_para(doc, "Sotto-processo B: xxx (TO-BE)")
if p:
    set_para_text(p,
        "Sotto-processo B: Pianificazione Ottimizzata Interventi con Vincoli Idrologici (TO-BE)")
_, p = find_para(doc, "Sotto-processo C: xxx (TO-BE)")
if p:
    set_para_text(p,
        "Sotto-processo C: Esecuzione, Monitoraggio e Feedback per Apprendimento Continuo (TO-BE)")


# ── TABLE 6 – TO-BE Sotto-processo A ──────────────────────────────────────────
t6 = doc.tables[6]
rows_tobe_a = [
    ("A1",
     "ETL Pipeline acquisisce in automatico dati da SCADA, sistema monitoraggio invasi e "
     "previsioni idro-meteo; normalizza e struttura il dataset per gli agenti AI",
     "ETL Pipeline (automatico)",
     "Dati operativi raw (SCADA, sensori, livello invaso, previsioni)",
     "Dataset normalizzato e strutturato",
     "SCADA, sistema monitoraggio invasi, sistema previsioni, [INFO: altri sistemi]",
     "ETL automatico con regole di qualita' dati e alerting su anomalie pipeline"),
    ("A2",
     "Hydro Monitor Agent analizza il dataset e classifica le anomalie per priorita' "
     "e urgenza, correlando con storico guasti analoghi da Knowledge Base",
     "Hydro Monitor Agent",
     "Dataset normalizzato",
     "Lista anomalie classificate (priorita' 1-4, urgenza, componente, impianto)",
     "SCADA, [INFO: piattaforma analytics]",
     "ML classification + LLM reasoning"),
    ("A3",
     "Knowledge Agent interroga la KB tecnica per storico anomalie analoghe, procedure "
     "di intervento e documentazione tecnica dei componenti",
     "Knowledge Agent",
     "ID componente, tipologia anomalia, impianto",
     "Scheda tecnica, storico interventi analoghi, procedure OEM",
     "Knowledge Base (RAG)",
     "RAG + vector search su documentazione tecnica impianti idroelettrici"),
    ("A4",
     "Hydro Monitor Agent correla le anomalie con il contesto idrologico attuale "
     "(livello invaso, portata, previsioni) per valutare l'impatto sulla produzione",
     "Hydro Monitor Agent",
     "Anomalie classificate + dati idrologici correnti + output Knowledge Agent",
     "Analisi arricchita con impatto sulla produzione e raccomandazione preliminare",
     "Tutti",
     "LLM reasoning + contesto idrologico e tecnico"),
    ("A5",
     "Orchestrator Agent consolida le analisi e genera un report strutturato "
     "con anomalie, priorita', raccomandazioni e impatto atteso sulla produzione",
     "Orchestrator Agent",
     "Analisi consolidata da tutti gli agenti",
     "Report analitico completo per il tecnico d'ufficio",
     "Dashboard HITL, Teams Bot",
     "LLM synthesis e formatting report"),
    ("A6",
     "HITL – Tecnico d'ufficio revisiona il report AI, valida o corregge le priorita', "
     "puo' aggiungere contesto operativo non rilevabile automaticamente",
     "Tecnico d'ufficio",
     "Report analitico AI",
     "Report validato con eventuali correzioni",
     "Dashboard HITL",
     "Assist – Human decides e approva"),
    ("A7",
     "Report validato inviato a Planning Agent per avvio fase di pianificazione ottimizzata",
     "Orchestrator Agent",
     "Report validato",
     "Trigger avvio Sotto-processo B",
     "Orchestrator, Planning Agent",
     "Hand-off automatico tra agenti"),
]
for i, row_data in enumerate(rows_tobe_a):
    if i + 1 < len(t6.rows):
        fill_row(t6, i + 1, list(row_data))


# ── TABLE 7 – TO-BE Sotto-processo B ──────────────────────────────────────────
t7 = doc.tables[7]
rows_tobe_b = [
    ("B1",
     "Planning Agent riceve la lista interventi validati e avvia la raccolta automatica "
     "dei vincoli di pianificazione specifici per impianti idroelettrici",
     "Planning Agent",
     "Lista interventi prioritizzati (validata HITL)",
     "Raccolta vincoli avviata",
     "Planning Agent, API sistemi esterni",
     "Automatic constraint collection via function calling"),
    ("B2",
     "Water Resource Agent verifica le previsioni idrologiche e identifica le finestre "
     "operative compatibili con i vincoli di concessione (deflusso minimo vitale, livello invaso)",
     "Water Resource Agent",
     "Previsioni idro-meteo + dati invaso + vincoli concessione",
     "Finestre operative idonee nei prossimi 14 giorni",
     "[INFO: sistema previsioni idro-meteo], [INFO: DB vincoli concessione]",
     "Ottimizzazione multi-vincolo idrologico + LLM"),
    ("B3",
     "Planning Agent verifica la disponibilita' dei ricambi necessari per i componenti da sostituire",
     "Planning Agent",
     "Lista componenti da sostituire con codici articolo",
     "Disponibilita' ricambi, lead time, fornitori alternativi",
     "ERP Manutenzione, [INFO: sistema gestione fornitori]",
     "DB query + LLM per analisi disponibilita'"),
    ("B4",
     "Planning Agent raccoglie la disponibilita' dei fornitori specializzati "
     "per impianti idroelettrici (remoti, richiedono certificazioni specifiche)",
     "Planning Agent",
     "Lista interventi e impianti target",
     "Slot disponibilita' fornitori principali e alternativi",
     "[INFO: sistema gestione fornitori] — [INFO RICHIESTA: API disponibile?]",
     "API call o form digitale HITL"),
    ("B5",
     "Planning Agent integra tutti i vincoli (idrologici, meteo, ricambi, fornitori, "
     "economici) e genera il piano di manutenzione ottimizzato",
     "Planning Agent",
     "Tutti i vincoli raccolti",
     "Piano di manutenzione ottimizzato con scheduling dettagliato",
     "Planning Agent (constraint optimizer)",
     "Multi-constraint optimization + LLM; [INFO: specificare i pesi dei vincoli]"),
    ("B6",
     "Orchestrator Agent presenta il piano al supervisore tramite Dashboard HITL "
     "con spiegazione delle motivazioni AI per ciascuna decisione (explainability obbligatoria)",
     "Orchestrator Agent",
     "Piano ottimizzato",
     "Piano presentato su dashboard con motivazioni e impatto atteso",
     "Dashboard HITL, Teams Bot",
     "LLM explanation generation"),
    ("B7",
     "HITL – Supervisore impianto revisiona il piano AI, valuta la coerenza con la "
     "conoscenza operativa locale, puo' modificarlo e infine lo approva",
     "Supervisore impianto",
     "Piano ottimizzato AI con spiegazioni",
     "Piano approvato (o modificato e approvato)",
     "Dashboard HITL",
     "Assist – Human decides e approva"),
    ("B8",
     "Piano approvato triggera la generazione degli ordini di lavoro e delle "
     "notifiche ai fornitori (post-HITL approval, mai in autonomia)",
     "Orchestrator Agent",
     "Piano approvato",
     "OdL draft generati, notifiche bozza pronte",
     "ERP, [INFO: sistema notifiche fornitori]",
     "Automazione generazione OdL post-approvazione umana"),
]
for i, row_data in enumerate(rows_tobe_b):
    if i + 1 < len(t7.rows):
        fill_row(t7, i + 1, list(row_data))


# ── TABLE 8 – TO-BE Sotto-processo C ──────────────────────────────────────────
t8 = doc.tables[8]
rows_tobe_c = [
    ("C1",
     "Sistema notifica automaticamente fornitori e tecnici degli interventi pianificati "
     "(email/Teams), solo dopo approvazione HITL del supervisore",
     "Orchestrator Agent / Sistema",
     "OdL approvati",
     "Notifiche inviate a fornitori e tecnici",
     "ERP, Email/Teams",
     "Notifiche automatiche solo post-approvazione HITL"),
    ("C2",
     "Esecuzione fisica degli interventi di manutenzione da parte di tecnici/fornitori "
     "specializzati (impianti idroelettrici: richiedono competenze e certificazioni specifiche)",
     "Tecnico / Fornitore specializzato",
     "Ordini di lavoro",
     "Intervento eseguito",
     "On-field (Mobile App)",
     "N/A – attivita' umana fisica; [INFO: specificare certificazioni richieste per impianti idro]"),
    ("C3",
     "Tecnico compila il verbale post-intervento tramite Mobile App (strutturato) "
     "includendo dati tecnici (componente sostituito, misure pre/post, anomalie riscontrate)",
     "Tecnico manutentore",
     "Intervento completato",
     "Verbale digitale strutturato",
     "Mobile App TO-BE",
     "Form guidato AI con suggerimenti compilazione"),
    ("C4",
     "Knowledge Agent acquisisce il verbale e aggiorna la Knowledge Base "
     "con il nuovo case di manutenzione idroelettrica",
     "Knowledge Agent",
     "Verbale post-intervento strutturato",
     "KB aggiornata con nuovo case",
     "Knowledge Base (RAG)",
     "Document ingestion + embedding update"),
    ("C5",
     "Orchestrator Agent aggiorna lo stato degli interventi, calcola i KPI "
     "e aggiorna la dashboard operativa",
     "Orchestrator Agent",
     "Verbali completati, dati post-intervento",
     "KPI aggiornati (MTTR, accuracy anomalia, disponibilita' impianto)",
     "Dashboard HITL",
     "KPI calculation + dashboard update automatico"),
    ("C6",
     "Feedback loop: Hydro Monitor Agent aggiorna i modelli predittivi "
     "con i dati di esito degli interventi (continual learning)",
     "Hydro Monitor Agent / [INFO: piattaforma analytics]",
     "KPI e esiti interventi",
     "Modelli predittivi aggiornati",
     "[INFO: piattaforma analytics], Hydro Monitor Agent",
     "[INFO: specificare la strategia di continual learning – batch mensile, on-demand]"),
    ("C7",
     "Reporting periodico aggregato a management: performance impianti, "
     "trend guasti, efficienza manutenzione, produzione ottimizzata",
     "Orchestrator Agent",
     "Dati KPI aggregati (mensile)",
     "Report management con insights e raccomandazioni strategiche",
     "Dashboard HITL, Power BI",
     "LLM-generated insights da dati aggregati"),
]
for i, row_data in enumerate(rows_tobe_c):
    if i + 1 < len(t8.rows):
        fill_row(t8, i + 1, list(row_data))


# ── PROCESS CARDS TO-BE ───────────────────────────────────────────────────────
_, p = find_para(doc, "Card 1: xxx (TO-BE)")
if p:
    set_para_text(p, "Card 1: Monitoraggio e Classificazione Anomalie Impianti Idroelettrici (TO-BE)")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "- Dati operativi SCADA; Alert anomalie; Previsioni idrologiche; Knowledge Base tecnica impianti")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "1. ETL Pipeline normalizza automaticamente i dati da tutti i sistemi sorgente",
        "2. Hydro Monitor Agent classifica le anomalie per priorita' e impatto sulla produzione",
        "3. Knowledge Agent recupera storico anomalie analoghe e procedure tecniche",
        "4. Orchestrator Agent genera report consolidato con raccomandazioni",
        "5. HITL: Tecnico d'ufficio valida e approva il report tramite Dashboard"
    ])
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "- Report analitico validato con lista anomalie prioritizzate, raccomandazioni di intervento "
        "e impatto stimato sulla produzione idroelettrica")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "SCADA, sistema monitoraggio invasi, Knowledge Base, ETL Pipeline, Dashboard HITL")

_, p = find_para(doc, "Card 2: xxx (TO-BE)")
if p:
    set_para_text(p, "Card 2: Pianificazione Ottimizzata con Vincoli Idrologici (TO-BE)")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "- Lista interventi validati; Previsioni idro-meteo 14 giorni; "
        "Disponibilita' fornitori e ricambi; Vincoli concessione idrica")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "1. Water Resource Agent verifica finestre operative compatibili con vincoli idrologici",
        "2. Planning Agent raccoglie disponibilita' fornitori e ricambi",
        "3. Planning Agent genera piano ottimizzato multi-vincolo (idro, meteo, economico, risorse)",
        "4. HITL: Supervisore revisiona, eventualmente modifica e approva il piano"
    ])
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "- Piano manutenzione ottimizzato approvato; OdL draft generati per emissione post-HITL")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "Planning Agent, Water Resource Agent, [INFO: sistema previsioni idro-meteo], "
        "ERP Manutenzione, Dashboard HITL")

_, p = find_para(doc, "Card 3: xxx (TO-BE)")
if p:
    set_para_text(p, "Card 3: Esecuzione, Monitoraggio e Feedback Loop (TO-BE)")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p, "- Piano manutenzione approvato; Verbali post-intervento strutturati (Mobile App)")
_, p = find_para(doc, "1. xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "1. Sistema notifica automaticamente fornitori e tecnici (post-approvazione HITL)",
        "2. Tecnici eseguono interventi e compilano verbale digitale strutturato via Mobile App",
        "3. Knowledge Agent acquisisce verbale e aggiorna KB con nuovo case tecnico",
        "4. Orchestrator calcola KPI e aggiorna dashboard",
        "5. Feedback loop: agenti calibrano modelli predittivi con esiti reali"
    ])
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "- KB aggiornata; KPI operativi calcolati; modelli predittivi migliorati; report management generato")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_text(p,
        "Knowledge Base, Orchestrator Agent, ERP, Mobile App, Dashboard HITL, "
        "[INFO: piattaforma analytics/predittiva]")


# ── 3.4 COSA NON FA L'AI ──────────────────────────────────────────────────────
_, p = find_para(doc, "3.4 Cosa NON fa l'AI", "Heading 3")
_, p_cosa = find_para(doc, "xxx", "Normal", _)
if p_cosa:
    set_para_lines(p_cosa, [
        "Non esegue fisicamente gli interventi di manutenzione sugli impianti idroelettrici",
        "Non emette autonomamente ordini di lavoro senza esplicita approvazione del supervisore umano (HITL obbligatorio)",
        "Non prende decisioni autonome sulla gestione dell'invaso o sulla regolazione delle paratoie",
        "Non modifica autonomamente i parametri di produzione degli impianti (set-point, curtailment)",
        "Non gestisce le comunicazioni con gli enti regolatori (RSDI, Autorita' di Bacino) in autonomia",
        "Non sostituisce il giudizio tecnico del supervisore per decisioni critiche di sicurezza sul campo",
        "Non accede direttamente ai sistemi di controllo SCADA ne' puo' modificarne parametri operativi",
        "Non comunica autonomamente con fornitori o imprese appaltatrici senza validazione umana"
    ])


# ── 3.5 COMPONENTE AI ─────────────────────────────────────────────────────────
_, p = find_para(doc, "3.5 Componente ai", "Heading 3")
_, p_ai = find_para(doc, "Xxx", "Normal", _)
if p_ai:
    set_para_lines(p_ai, [
        "Architettura multi-agent AI composta da agenti specializzati coordinati da un Orchestrator:",
        "",
        "(1) Orchestrator Agent – coordina il flusso end-to-end e gestisce l'interfaccia HITL.",
        "Pattern tassonomico: Orchestrator / Coordinator Agent",
        "",
        "(2) Hydro Monitor Agent – analizza alert SCADA e dati operativi; classifica anomalie "
        "per priorita' e urgenza; correla con contesto idrologico e storico KB.",
        "Pattern tassonomico: Analyzer / Classifier Agent",
        "",
        "(3) Planning Agent – ottimizza il piano di manutenzione con vincoli idrologici "
        "(deflusso minimo vitale, livello invaso), meteo, economici, disponibilita' risorse.",
        "Pattern tassonomico: Planner / Optimizer Agent",
        "",
        "(4) Water Resource Agent – [INFO RICHIESTA: confermare se in scope] – ottimizza "
        "la programmazione della produzione idroelettrica.",
        "Pattern tassonomico: Optimizer Agent (produzione)",
        "",
        "(5) Knowledge Agent – recupera e sintetizza informazioni tecniche tramite RAG "
        "su documentazione impianti (manuali OEM, storici interventi).",
        "Pattern tassonomico: Retriever / RAG Agent"
    ])


# ── DELTA AS-IS vs TO-BE ──────────────────────────────────────────────────────
_, p = find_para(doc, "AS-IS:", "Normal")
if p:
    set_para_text(p, "AS-IS:")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "- Monitoraggio manuale degli impianti tramite dashboard SCADA; "
        "[INFO: specificare ore/giorno dedicate al monitoraggio]",
        "- Comunicazione informale tra personale d'ufficio e supervisori di impianto",
        "- Pianificazione manuale senza ottimizzazione formale dei vincoli idrologici",
        "- Knowledge tecnica non strutturata; dipendenza da esperienza individuale",
        "- [INFO RICHIESTA: aggiungere altre limitazioni specifiche emerse dalla BSN email AIDRO]"
    ])
_, p = find_para(doc, "TO-BE:", "Normal")
if p:
    set_para_text(p, "TO-BE:")
_, p = find_para(doc, "- xxx", "Normal", _)
if p:
    set_para_lines(p, [
        "- ETL Pipeline normalizza automaticamente i dati da SCADA, invasi e previsioni idro-meteo",
        "- Hydro Monitor Agent classifica le anomalie e genera report arricchiti con contesto tecnico e idrologico",
        "- Planning Agent ottimizza i piani manutenzione considerando vincoli idrologici, meteo ed economici",
        "- Supervisore revisiona e approva tramite dashboard dedicata (HITL): decisione informata in tempi ridotti",
        "- Knowledge Base strutturata rende accessibile la conoscenza tecnica a tutti gli operatori"
    ])


# ── TABLE 9 – ROADMAP ─────────────────────────────────────────────────────────
t9 = doc.tables[9]
roadmap = [
    ("Fase 1 – Raccolta Requisiti e Accessi\n(M1-M2)",
     "Allineamento con BSN e team EGP TGX per definire perimetro definitivo; "
     "mappatura API verso SCADA, sistema monitoraggio invasi ed ERP manutenzione; "
     "raccolta documentazione tecnica per KB",
     "Architettura tecnica validata; accordi IT per accessi API; "
     "inventario documentazione disponibile per la Knowledge Base",
     "2 mesi — [INFO RICHIESTA AL PM: confermare]"),
    ("Fase 2 – Data Access e Knowledge Base\n(M2-M4)",
     "Implementazione pipeline ETL verso SCADA e sistemi sorgente; "
     "digitalizzazione e indicizzazione Knowledge Base "
     "(manuali OEM turbine/generatori/paratoie, storici interventi); "
     "validazione qualita' dati idrologici",
     "ETL funzionante su ambiente di sviluppo; "
     "KB popolata con 80%+ della documentazione disponibile",
     "2 mesi — [INFO RICHIESTA AL PM: confermare]"),
    ("Fase 3 – AI PoC\n(M3-M6)",
     "Sviluppo e test Hydro Monitor Agent e Planning Agent in sandbox; "
     "validazione accuracy classificazione anomalie su dati storici; "
     "[INFO: specificare se Water Resource Agent e' incluso nel PoC]",
     "PoC demo con accuracy > 85% su dataset storico; "
     "planning agent genera piani coerenti con vincoli idrologici su casi test",
     "3 mesi — [INFO RICHIESTA AL PM: confermare]"),
    ("Fase 4 – Integration e HITL Interface\n(M5-M8)",
     "Integrazione agenti con sistemi reali (SCADA, invasi, ERP manutenzione); "
     "sviluppo Dashboard HITL e Teams Bot; "
     "test operativi con supervisori e tecnici d'ufficio pilota",
     "Sistema integrato su ambiente pre-produzione; "
     "Dashboard HITL validata da [INFO: n.] supervisori pilota",
     "3 mesi — [INFO RICHIESTA AL PM: confermare]"),
    ("Fase 5 – Pilota e Validazione\n(M7-M9)",
     "Pilota operativo su [INFO: indicare n. impianti pilota] impianti idroelettrici selezionati; "
     "monitoraggio KPI (MTTR, accuracy, FTE liberati); go/no-go per scale-up",
     "Riduzione MTTR verificata su impianti pilota; "
     "adozione sistema > 80% supervisori; report go/no-go scale-up",
     "3 mesi — [INFO RICHIESTA AL PM: confermare]"),
    ("Fase 6 – Scale-up\n(M9-M12)",
     "Estensione a tutti gli impianti idroelettrici in scope EGP TGX; "
     "formazione personale; messa a regime monitoraggio e supporto continuativo; "
     "[INFO: specificare il numero totale di impianti in scope per lo scale-up]",
     "Sistema in produzione su tutti gli impianti in scope; "
     "piano di miglioramento continuo definito",
     "3 mesi — [INFO RICHIESTA AL PM: confermare]"),
]
for i, row_data in enumerate(roadmap):
    if i + 1 < len(t9.rows):
        fill_row(t9, i + 1, list(row_data))


# ── TABLE 10 – KPI QUANTITATIVI ───────────────────────────────────────────────
t10 = doc.tables[10]
kpi_quant = [
    ("MTTR Manutenzione (giorni dalla rilevazione anomalia all'intervento completato)",
     "[INFO RICHIESTA AL BSN: richiedere il MTTR AS-IS medio per gli impianti idroelettrici]",
     "[INFO: definire target TO-BE – es. -30% rispetto AS-IS]",
     "Tracking timestamp: da alert SCADA a chiusura OdL su ERP"),
    ("Accuracy classificazione anomalie AI (% anomalie correttamente prioritizzate)",
     "[INFO RICHIESTA: richiedere stima di accuracy AS-IS dell'analisi manuale]",
     "> 88% [INFO: confermare target con il Business Owner]",
     "Confronto classificazione AI vs. esito reale (feedback loop post-intervento)"),
    ("Tempo medio analisi singola anomalia",
     "[INFO RICHIESTA AL BSN: indicare il tempo medio AS-IS per l'analisi manuale (ore/anomalia)]",
     "< 1 ora (AI + HITL) [INFO: confermare target]",
     "Timestamp apertura -> chiusura analisi su Dashboard HITL"),
    ("Volume anomalie gestite per FTE (anomalie/settimana/operatore)",
     "[INFO RICHIESTA: richiedere la capacita' AS-IS in anomalie/settimana/FTE]",
     "[INFO: definire target TO-BE atteso con AI]",
     "Conteggio anomalie processate su dashboard / FTE assegnati"),
    ("Lead time pianificazione intervento",
     "[INFO RICHIESTA AL BSN: indicare il lead time medio attuale dalla lista interventi al piano approvato]",
     "< 6 ore (AI-generated + HITL approval) [INFO: confermare target]",
     "Timestamp ricezione lista interventi -> piano approvato dal supervisore"),
    ("Disponibilita' impianto (% ore in produzione / ore totali)",
     "[INFO RICHIESTA: richiedere il dato AS-IS di disponibilita' degli impianti idroelettrici]",
     "[INFO: definire target miglioramento % – tipicamente +1-3% per impianti idro]",
     "Rapporto ore in produzione / ore totali nel periodo di riferimento"),
    ("Tasso fermate impreviste/anno per impianto",
     "[INFO RICHIESTA: richiedere storico fermate impreviste/anno per impianto]",
     "[INFO: definire target riduzione % fermate impreviste]",
     "Conteggio fermate impreviste anno post-vs.-pre deployment AI"),
    ("Copertura Knowledge Base (% documentazione tecnica indicizzata)",
     "0% (non strutturata)",
     "> 85%",
     "Ratio documenti indicizzati in KB / totale documenti tecnici disponibili"),
]
for i, row_data in enumerate(kpi_quant):
    if i + 1 < len(t10.rows):
        fill_row(t10, i + 1, list(row_data))


# ── TABLE 11 – KPI QUALITATIVI ────────────────────────────────────────────────
t11 = doc.tables[11]
kpi_qual = [
    ("Carico cognitivo operatori (self-reported)",
     "Riduzione carico analisi manuale; target: maggioranza operatori lo valuta ridotto "
     "(scala Likert, survey semestrale) [INFO: definire target score con il Business Owner]",
     "Survey semestrale a tecnici d'ufficio e supervisori di impianto"),
    ("Adozione sistema AI da parte dei supervisori",
     "Target: > 90% supervisori utilizza regolarmente la Dashboard HITL "
     "entro 6 mesi dal go-live [INFO: confermare target con EGP TGX]",
     "Log di accesso e utilizzo Dashboard HITL"),
    ("Qualita' dei piani AI (% piani approvati senza modifiche sostanziali)",
     "Target: > 70% dei piani generati dall'AI approvati senza modifiche sostanziali "
     "[INFO: confermare target con i supervisori coinvolti nel pilota]",
     "Tracking HITL: piani approvati as-is / totale piani generati"),
    ("Completezza verbali post-intervento digitali",
     "Target: > 95% verbali compilati in formato strutturato tramite Mobile App "
     "[INFO: adattare alla realta' operativa impianti remoti]",
     "Tracking invio verbali digitali / totale interventi eseguiti"),
    ("Qualita' knowledge retrieval (% query KB risolte senza escalation)",
     "Target: > 75% delle query Knowledge Agent risolte autonomamente "
     "[INFO: da calibrare dopo l'analisi della documentazione disponibile]",
     "Tracking query KB: risolte autonomamente / totale query"),
    ("Soddisfazione utenti sistema AI (NPS interno)",
     "Target: NPS > 30 entro 12 mesi dal go-live [INFO: confermare target]",
     "Survey NPS trimestrale a tutti gli utenti del sistema"),
    ("Tracciabilita' decisioni AI (% raccomandazioni con spiegazione leggibile)",
     "Target: 100% delle raccomandazioni AI presentate con motivazione comprensibile "
     "al supervisore (obbligatorio per conformita' HITL)",
     "Audit log Dashboard HITL: % decisioni con explanation presente"),
    ("[INFO RICHIESTA AL BSN: specificare ulteriori KPI qualitativi rilevanti per AIDRO]",
     "[INFO: definire target]",
     "[INFO: definire metodo di misurazione]"),
]
for i, row_data in enumerate(kpi_qual):
    if i + 1 < len(t11.rows):
        fill_row(t11, i + 1, list(row_data))


# ── SAVE ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"\n Blueprint AIDRO salvato in: {OUTPUT}")
print("\nSezioni completate con dati contestuali EGP TGX.")
print("SEZIONI CON [INFO RICHIESTA] DA COMPLETARE CON IL TEAM EGP TGX / BSN:")
print("  - Sistemi AS-IS: SCADA idroelettrico, sistema monitoraggio invasi, piattaforma predittiva")
print("  - Metriche AS-IS: MTTR attuale, FTE coinvolti, ore/giorno monitoraggio manuale")
print("  - Perimetro: confermare se ottimizzazione produzione (Water Resource Agent) e' in scope")
print("  - Ruoli: Business Owner, Data Owner, IT Owner, Product Owner del progetto GDS+OP-36")
print("  - Vincoli: normative specifiche applicabili, accessi API sistemi esistenti")
print("  - KPI: valori baseline AS-IS da richiedere al BSN")
print("  - Roadmap: durate e n. impianti pilota da confermare col PM")